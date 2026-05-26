"""FastAPI entrypoint for the ESG QLoRA extraction service."""

from __future__ import annotations

import os
import io
import time
import asyncio
import base64
import smtplib
import random
import queue
import secrets
import string
import threading
import uuid
import json
import re
import hashlib
import bcrypt
import jwt
import aiosqlite
from concurrent.futures import ThreadPoolExecutor
from email.message import EmailMessage
from email.utils import formataddr
from pathlib import Path
from datetime import datetime, timedelta, timezone
from typing import Any, Callable, Dict, List, Literal, Optional, Tuple
from urllib.parse import unquote

from PIL import Image, ImageDraw, ImageFilter, ImageFont
from dotenv import load_dotenv
from docx import Document as DocxDocument

from fastapi import FastAPI, File, Form, UploadFile, HTTPException, Depends, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, Field

load_dotenv(Path(__file__).resolve().parent / ".env")

# ── Auth config ──────────────────────────────────────────────────────────────
_DEFAULT_JWT_SECRET = "esg-demo-secret-change-in-prod"
_APP_ENV = os.getenv("APP_ENV", os.getenv("ENVIRONMENT", "development")).strip().lower()
_JWT_SECRET = os.getenv("JWT_SECRET", _DEFAULT_JWT_SECRET).strip() or _DEFAULT_JWT_SECRET
_JWT_ALGORITHM = "HS256"
_TOKEN_MINUTES = 60 * 24  # 1 day


def _resolve_auth_db_path() -> str:
    configured = os.getenv("AUTH_DB_PATH", "").strip()
    if configured:
        return configured

    fly_data_dir = Path("/data")
    if fly_data_dir.exists() and os.access(fly_data_dir, os.W_OK):
        return str(fly_data_dir / "auth.db")

    return str(Path(__file__).resolve().parent / "auth.db")


_DB_PATH = _resolve_auth_db_path()
_FEEDBACK_DB_PATH = os.path.join(os.path.dirname(__file__), "backend", "causalgraph.db")
_security = HTTPBearer(auto_error=False)
_CLEANUP_EXECUTOR = ThreadPoolExecutor(max_workers=1)
_GRAPH_CACHE_REFRESH_EXECUTOR = ThreadPoolExecutor(max_workers=1)
_ENTITY_TOKEN_PATTERN = re.compile(r"[A-Za-z][A-Za-z0-9&.-]*")
_ENTITY_OF_PATTERN = re.compile(
    r"\b(?:of|for|about|by|from|at|on)\s+([A-Z][A-Za-z0-9&.-]*(?:\s+[A-Z][A-Za-z0-9&.-]*){0,3})"
)
_ENTITY_POSSESSIVE_PATTERN = re.compile(r"\b([A-Z][A-Za-z0-9&.-]*(?:\s+[A-Z][A-Za-z0-9&.-]*){0,3})['’]s\b")
_ENTITY_ACRONYM_PATTERN = re.compile(r"\b[A-Z]{2,8}\b")
_ENTITY_TITLE_PHRASE_PATTERN = re.compile(r"\b([A-Z][A-Za-z0-9&.-]*(?:\s+[A-Z][A-Za-z0-9&.-]*){1,4})\b")
_ENTITY_SPLIT_PATTERN = re.compile(r"[^A-Za-z0-9&]+")
_ENTITY_ALIAS_MAP: Dict[str, Tuple[str, ...]] = {
    "american flight": ("american airlines", "american airline", "aa"),
    "american flights": ("american airlines", "american airline", "aa"),
    "american airline": ("american airlines", "aa"),
    "american airlines": ("american airline", "aa"),
}
_DOCUMENT_ENTITY_TERMS_CACHE: Dict[str, Tuple[float, set[str]]] = {}
_ENTITY_STOP_WORDS = {
    "a",
    "about",
    "and",
    "annual",
    "be",
    "can",
    "company",
    "across",
    "compare",
    "could",
    "document",
    "esg",
    "effects",
    "explain",
    "for",
    "from",
    "give",
    "governance",
    "i",
    "in",
    "inc",
    "is",
    "its",
    "know",
    "limited",
    "llc",
    "ltd",
    "me",
    "of",
    "on",
    "please",
    "price",
    "report",
    "reports",
    "responsibility",
    "share",
    "show",
    "something",
    "strategy",
    "summarise",
    "summarize",
    "sustainability",
    "tell",
    "that",
    "these",
    "this",
    "those",
    "the",
    "to",
    "want",
    "what",
    "would",
    "year",
}


def _is_production_like_env(value: Optional[str] = None) -> bool:
    env = str(value if value is not None else _APP_ENV).strip().lower()
    return env in {"prod", "production", "staging"}


def _validate_startup_security_config(*, app_env: Optional[str] = None, jwt_secret: Optional[str] = None) -> None:
    env = str(app_env if app_env is not None else _APP_ENV).strip().lower()
    secret = str(jwt_secret if jwt_secret is not None else _JWT_SECRET).strip()
    if not _is_production_like_env(env):
        if secret == _DEFAULT_JWT_SECRET:
            print("[startup] WARNING: using demo JWT_SECRET; set a strong JWT_SECRET before deployment.")
        return

    if not secret or secret == _DEFAULT_JWT_SECRET:
        raise RuntimeError("JWT_SECRET must be set to a strong non-default value when APP_ENV=production/staging.")
    if len(secret) < 32:
        raise RuntimeError("JWT_SECRET must be at least 32 characters when APP_ENV=production/staging.")


def _parse_cors_origins(raw: Optional[str]) -> List[str]:
    origins = [item.strip() for item in str(raw or "").split(",") if item.strip()]
    if origins:
        return origins
    return [
        "http://127.0.0.1:3000",
        "http://localhost:3000",
        "http://127.0.0.1:3001",
        "http://localhost:3001",
    ]


_CORS_ALLOW_ORIGINS = _parse_cors_origins(os.getenv("CORS_ALLOW_ORIGINS", ""))
_CORS_ALLOW_ORIGIN_REGEX = os.getenv(
    "CORS_ALLOW_ORIGIN_REGEX",
    r"https://.*\.ngrok-free\.app|https://.*\.ngrok\.app",
).strip() or None


def _env_flag(name: str, default: str = "false") -> bool:
    return os.getenv(name, default).strip().lower() in {"1", "true", "yes", "on"}


_MAIL_ENABLED = _env_flag("MAIL_ENABLED", "false")
_MAIL_SMTP_HOST = os.getenv("MAIL_SMTP_HOST", "").strip()
_MAIL_SMTP_PORT = int(os.getenv("MAIL_SMTP_PORT", "465"))
_MAIL_SMTP_SSL = _env_flag("MAIL_SMTP_SSL", "true")
_MAIL_SMTP_STARTTLS = _env_flag("MAIL_SMTP_STARTTLS", "false")
_MAIL_SMTP_USER = os.getenv("MAIL_SMTP_USER", "").strip()
_MAIL_SMTP_PASSWORD = os.getenv("MAIL_SMTP_PASSWORD", "").strip()
_MAIL_FROM = os.getenv("MAIL_FROM", _MAIL_SMTP_USER).strip()
_MAIL_FROM_NAME = os.getenv("MAIL_FROM_NAME", "CausalGraph AI").strip()
_EMAIL_CODE_TTL_SECONDS = max(60, int(os.getenv("EMAIL_CODE_TTL_SECONDS", "600")))
_EMAIL_CODE_RESEND_COOLDOWN_SECONDS = max(15, int(os.getenv("EMAIL_CODE_RESEND_COOLDOWN_SECONDS", "60")))
_EMAIL_CODE_MAX_ATTEMPTS = max(1, int(os.getenv("EMAIL_CODE_MAX_ATTEMPTS", "5")))
_EMAIL_CODE_LENGTH = max(4, int(os.getenv("EMAIL_CODE_LENGTH", "6")))
_RAG_RATE_LIMIT_ENABLED = _env_flag("RAG_RATE_LIMIT_ENABLED", "true")
_RAG_FREE_DAILY_POINTS = max(1, int(os.getenv("RAG_FREE_DAILY_POINTS", "30")))
_RAG_FLASH_POINT_COST = max(1, int(os.getenv("RAG_FLASH_POINT_COST", "1")))
_RAG_DEEP_POINT_COST = max(1, int(os.getenv("RAG_DEEP_POINT_COST", "5")))
_RAG_MIN_SECONDS_BETWEEN_REQUESTS = max(0, int(os.getenv("RAG_MIN_SECONDS_BETWEEN_REQUESTS", "20")))
_RAG_ANONYMOUS_ENABLED = _env_flag("RAG_ANONYMOUS_ENABLED", "false")
_DESKTOP_SCREENSHOT_SUMMARY_MODEL = os.getenv(
    "DESKTOP_SCREENSHOT_SUMMARY_MODEL",
    os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
).strip() or "gpt-4o-mini"
_DESKTOP_SCREENSHOT_SUMMARY_MAX_TOKENS = max(128, int(os.getenv("DESKTOP_SCREENSHOT_SUMMARY_MAX_TOKENS", "700")))
_DESKTOP_SCREENSHOT_MAX_IMAGE_BYTES = max(256_000, int(os.getenv("DESKTOP_SCREENSHOT_MAX_IMAGE_BYTES", "6000000")))
_DESKTOP_WORD_EDIT_MODEL = os.getenv(
    "DESKTOP_WORD_EDIT_MODEL",
    os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
).strip() or "gpt-4o-mini"
_DESKTOP_WORD_EDIT_MAX_BYTES = max(512_000, int(os.getenv("DESKTOP_WORD_EDIT_MAX_BYTES", "12000000")))
_DESKTOP_WORD_EDIT_MAX_PARAGRAPHS = max(4, int(os.getenv("DESKTOP_WORD_EDIT_MAX_PARAGRAPHS", "28")))
_DESKTOP_WORD_EDIT_MAX_CHARS = max(2_000, int(os.getenv("DESKTOP_WORD_EDIT_MAX_CHARS", "14000")))
_DESKTOP_WORD_EDIT_MAX_SUGGESTIONS = max(1, int(os.getenv("DESKTOP_WORD_EDIT_MAX_SUGGESTIONS", "8")))
_DESKTOP_WORD_EDIT_MAX_TOKENS = max(500, int(os.getenv("DESKTOP_WORD_EDIT_MAX_TOKENS", "1800")))
_DESKTOP_WORD_EDIT_EVIDENCE_TOP_K = max(0, int(os.getenv("DESKTOP_WORD_EDIT_EVIDENCE_TOP_K", "6")))
_WORD_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_PUBLIC_GRAPH_CACHE_VERSION = os.getenv("PUBLIC_GRAPH_CACHE_VERSION", "v1").strip() or "v1"
_PUBLIC_GRAPH_CACHE_TTL_SECONDS = max(60, int(os.getenv("PUBLIC_GRAPH_CACHE_TTL_SECONDS", "1800")))
_PUBLIC_GRAPH_CACHE_STALE_SECONDS = max(
    _PUBLIC_GRAPH_CACHE_TTL_SECONDS,
    int(os.getenv("PUBLIC_GRAPH_CACHE_STALE_SECONDS", str(24 * 60 * 60))),
)
_PUBLIC_GRAPH_CACHE_LOCK_SECONDS = max(10, int(os.getenv("PUBLIC_GRAPH_CACHE_LOCK_SECONDS", "180")))
_PUBLIC_GRAPH_CACHE_WAIT_SECONDS = max(1.0, float(os.getenv("PUBLIC_GRAPH_CACHE_WAIT_SECONDS", "12")))


async def _get_db():
    Path(_DB_PATH).parent.mkdir(parents=True, exist_ok=True)
    async with aiosqlite.connect(_DB_PATH) as db:
        yield db


async def _init_auth_db():
    Path(_DB_PATH).parent.mkdir(parents=True, exist_ok=True)
    async with aiosqlite.connect(_DB_PATH) as db:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id TEXT PRIMARY KEY,
                email TEXT UNIQUE NOT NULL,
                username TEXT NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL DEFAULT 'user',
                created_at TEXT NOT NULL
            )
        """)
        await db.execute("""
            CREATE TABLE IF NOT EXISTS admin_invite_codes (
                id TEXT PRIMARY KEY,
                code TEXT UNIQUE NOT NULL,
                created_by_user_id TEXT NOT NULL,
                created_at TEXT NOT NULL,
                expires_at TEXT NOT NULL,
                used_at TEXT,
                used_by_user_id TEXT
            )
        """)
        await db.execute("""
            CREATE TABLE IF NOT EXISTS email_verification_codes (
                id TEXT PRIMARY KEY,
                email TEXT NOT NULL,
                purpose TEXT NOT NULL,
                code_hash TEXT NOT NULL,
                created_at TEXT NOT NULL,
                expires_at TEXT NOT NULL,
                consumed_at TEXT,
                attempts INTEGER NOT NULL DEFAULT 0,
                last_sent_at TEXT NOT NULL
            )
        """)
        await db.execute("""
            CREATE TABLE IF NOT EXISTS rag_rate_usage (
                user_id TEXT NOT NULL,
                usage_date TEXT NOT NULL,
                points_used INTEGER NOT NULL DEFAULT 0,
                request_count INTEGER NOT NULL DEFAULT 0,
                flash_count INTEGER NOT NULL DEFAULT 0,
                deep_count INTEGER NOT NULL DEFAULT 0,
                last_request_at TEXT,
                PRIMARY KEY (user_id, usage_date)
            )
        """)
        await db.execute("""
            CREATE TABLE IF NOT EXISTS rag_unlimited_users (
                email TEXT PRIMARY KEY,
                note TEXT NOT NULL DEFAULT '',
                created_by_user_id TEXT NOT NULL,
                created_at TEXT NOT NULL
            )
        """)
        await db.execute("CREATE INDEX IF NOT EXISTS admin_invite_codes_expires_at_idx ON admin_invite_codes(expires_at)")
        await db.execute("CREATE INDEX IF NOT EXISTS admin_invite_codes_used_at_idx ON admin_invite_codes(used_at)")
        await db.execute("CREATE INDEX IF NOT EXISTS email_verification_codes_email_purpose_idx ON email_verification_codes(email, purpose, created_at)")
        await db.execute("CREATE INDEX IF NOT EXISTS email_verification_codes_expires_at_idx ON email_verification_codes(expires_at)")
        await db.execute("CREATE INDEX IF NOT EXISTS rag_rate_usage_date_idx ON rag_rate_usage(usage_date)")
        await db.execute("CREATE INDEX IF NOT EXISTS rag_unlimited_users_created_at_idx ON rag_unlimited_users(created_at)")
        await _ensure_column(db, "users", "role", "TEXT NOT NULL DEFAULT 'user'")
        await db.execute("UPDATE users SET role = 'user' WHERE role IS NULL OR lower(role) NOT IN ('admin', 'user')")
        await init_user_memory_db(db)
        admin_emails = sorted(_admin_email_set())
        if admin_emails:
            placeholders = ",".join("?" for _ in admin_emails)
            await db.execute(
                f"UPDATE users SET role = 'admin' WHERE lower(email) IN ({placeholders})",
                tuple(admin_emails),
            )
        await db.execute("DELETE FROM admin_invite_codes WHERE datetime(expires_at) <= datetime('now') OR used_at IS NOT NULL")
        await db.execute("DELETE FROM email_verification_codes WHERE datetime(expires_at) <= datetime('now') OR consumed_at IS NOT NULL")
        await db.commit()


async def _init_feedback_db():
    Path(_FEEDBACK_DB_PATH).parent.mkdir(parents=True, exist_ok=True)
    async with aiosqlite.connect(_FEEDBACK_DB_PATH) as db:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS answer_feedback (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id       TEXT NOT NULL,
                session_id    TEXT NOT NULL,
                message_id    TEXT NOT NULL,
                query         TEXT NOT NULL,
                answer        TEXT NOT NULL,
                rating        TEXT NOT NULL,
                reason_tags   TEXT,
                reason_text   TEXT,
                sources_json  TEXT,
                timings_json  TEXT,
                created_at    TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        await db.execute("CREATE INDEX IF NOT EXISTS idx_feedback_rating ON answer_feedback(rating, created_at)")
        await db.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_feedback_user_message ON answer_feedback(user_id, message_id)")
        await db.commit()


async def _ensure_column(db: aiosqlite.Connection, table_name: str, column_name: str, definition: str) -> None:
    cursor = await db.execute(f"PRAGMA table_info({table_name})")
    rows = await cursor.fetchall()
    existing = {str(row[1]) for row in rows}
    if column_name not in existing:
        await db.execute(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {definition}")


def _normalize_role(role: str) -> str:
    value = str(role or "").strip().lower()
    return value if value in {"admin", "user"} else "user"


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _utc_in_minutes_iso(minutes: int) -> str:
    return (datetime.now(timezone.utc) + timedelta(minutes=minutes)).isoformat()


def _utc_in_seconds_iso(seconds: int) -> str:
    return (datetime.now(timezone.utc) + timedelta(seconds=seconds)).isoformat()


def _parse_utc_iso(value: str) -> datetime:
    parsed = datetime.fromisoformat(str(value))
    if parsed.tzinfo is None:
        return parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(timezone.utc)


def _normalize_email(email: str) -> str:
    return str(email or "").strip().lower()


def _hash_pw(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()


def _check_pw(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed.encode())


def _hash_email_code(email: str, code: str) -> str:
    payload = f"{_normalize_email(email)}:{str(code).strip()}".encode()
    return bcrypt.hashpw(payload, bcrypt.gensalt()).decode()


def _check_email_code(email: str, code: str, hashed: str) -> bool:
    payload = f"{_normalize_email(email)}:{str(code).strip()}".encode()
    return bcrypt.checkpw(payload, hashed.encode())


def _rag_usage_date() -> str:
    return datetime.now(timezone.utc).date().isoformat()


def _rag_point_cost(reasoning_mode: Optional[str]) -> int:
    mode = str(reasoning_mode or "flash").strip().lower()
    return _RAG_DEEP_POINT_COST if mode == "deep" else _RAG_FLASH_POINT_COST


def _rag_limit_error(status_code: int, message: str, **extra: Any) -> HTTPException:
    detail = {"error": "rag_rate_limited", "message": message, **extra}
    return HTTPException(status_code=status_code, detail=detail)


async def _is_rag_unlimited_user(db: aiosqlite.Connection, current_user: Optional[Dict[str, Any]]) -> bool:
    email = _normalize_email(str((current_user or {}).get("email") or ""))
    if not email:
        return False
    cursor = await db.execute("SELECT 1 FROM rag_unlimited_users WHERE email = ?", (email,))
    return await cursor.fetchone() is not None


async def _enforce_rag_rate_limit(
    db: aiosqlite.Connection,
    current_user: Optional[Dict[str, Any]],
    reasoning_mode: Optional[str],
) -> Dict[str, Any]:
    if not _RAG_RATE_LIMIT_ENABLED:
        return {"bypassed": True, "reason": "disabled"}
    if _is_admin_user(current_user):
        return {"bypassed": True, "reason": "admin"}
    if await _is_rag_unlimited_user(db, current_user):
        return {"bypassed": True, "reason": "unlimited_user"}
    if not current_user:
        if not _RAG_ANONYMOUS_ENABLED:
            raise _rag_limit_error(401, "Please sign in to use the AI agent.")
        user_id = "anonymous"
    else:
        user_id = str(current_user.get("id") or "").strip()
    if not user_id:
        raise _rag_limit_error(401, "Please sign in to use the AI agent.")

    mode = "deep" if str(reasoning_mode or "flash").strip().lower() == "deep" else "flash"
    cost = _rag_point_cost(mode)
    usage_date = _rag_usage_date()
    now = datetime.now(timezone.utc)
    now_iso = now.isoformat()

    cursor = await db.execute(
        """
        SELECT points_used, request_count, flash_count, deep_count, last_request_at
        FROM rag_rate_usage
        WHERE user_id = ? AND usage_date = ?
        """,
        (user_id, usage_date),
    )
    row = await cursor.fetchone()
    points_used = int(row[0] or 0) if row else 0
    request_count = int(row[1] or 0) if row else 0
    flash_count = int(row[2] or 0) if row else 0
    deep_count = int(row[3] or 0) if row else 0
    last_request_at = str(row[4] or "") if row else ""

    if last_request_at and _RAG_MIN_SECONDS_BETWEEN_REQUESTS > 0:
        elapsed = (now - _parse_utc_iso(last_request_at)).total_seconds()
        if elapsed < _RAG_MIN_SECONDS_BETWEEN_REQUESTS:
            retry_after = max(1, int(_RAG_MIN_SECONDS_BETWEEN_REQUESTS - elapsed))
            raise _rag_limit_error(
                429,
                f"Please wait {retry_after} seconds before sending another message.",
                retry_after_seconds=retry_after,
                points_limit=_RAG_FREE_DAILY_POINTS,
                points_used=points_used,
                points_remaining=max(0, _RAG_FREE_DAILY_POINTS - points_used),
            )

    if points_used + cost > _RAG_FREE_DAILY_POINTS:
        raise _rag_limit_error(
            429,
            "Daily message limit reached. Please try again tomorrow.",
            points_limit=_RAG_FREE_DAILY_POINTS,
            points_used=points_used,
            points_remaining=max(0, _RAG_FREE_DAILY_POINTS - points_used),
            points_required=cost,
            reset_at=f"{usage_date}T23:59:59+00:00",
        )

    new_points = points_used + cost
    new_request_count = request_count + 1
    new_flash_count = flash_count + (1 if mode == "flash" else 0)
    new_deep_count = deep_count + (1 if mode == "deep" else 0)
    await db.execute(
        """
        INSERT INTO rag_rate_usage (
            user_id, usage_date, points_used, request_count, flash_count, deep_count, last_request_at
        ) VALUES (?,?,?,?,?,?,?)
        ON CONFLICT(user_id, usage_date) DO UPDATE SET
            points_used = excluded.points_used,
            request_count = excluded.request_count,
            flash_count = excluded.flash_count,
            deep_count = excluded.deep_count,
            last_request_at = excluded.last_request_at
        """,
        (user_id, usage_date, new_points, new_request_count, new_flash_count, new_deep_count, now_iso),
    )
    await db.commit()
    return {
        "bypassed": False,
        "mode": mode,
        "points_cost": cost,
        "points_limit": _RAG_FREE_DAILY_POINTS,
        "points_used": new_points,
        "points_remaining": max(0, _RAG_FREE_DAILY_POINTS - new_points),
        "request_count": new_request_count,
        "flash_count": new_flash_count,
        "deep_count": new_deep_count,
    }


def _make_token(user_id: str, email: str) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc) + timedelta(minutes=_TOKEN_MINUTES),
    }
    return jwt.encode(payload, _JWT_SECRET, algorithm=_JWT_ALGORITHM)


def _decode_token(token: str) -> dict:
    try:
        return jwt.decode(token, _JWT_SECRET, algorithms=[_JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid token")


async def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(_security),
    db: aiosqlite.Connection = Depends(_get_db),
):
    if credentials is None:
        raise HTTPException(status_code=401, detail="Not authenticated")
    payload = _decode_token(credentials.credentials)
    user_id = payload.get("sub")
    cursor = await db.execute(
        "SELECT id, email, username, role, created_at FROM users WHERE id = ?", (user_id,)
    )
    row = await cursor.fetchone()
    if not row:
        raise HTTPException(status_code=401, detail="User not found")
    return {"id": row[0], "email": row[1], "username": row[2], "role": _normalize_role(row[3]), "created_at": row[4]}


async def get_optional_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(_security),
    db: aiosqlite.Connection = Depends(_get_db),
):
    if credentials is None:
        return None
    try:
        payload = _decode_token(credentials.credentials)
    except HTTPException:
        return None
    user_id = payload.get("sub")
    cursor = await db.execute(
        "SELECT id, email, username, role, created_at FROM users WHERE id = ?",
        (user_id,),
    )
    row = await cursor.fetchone()
    if not row:
        return None
    return {"id": row[0], "email": row[1], "username": row[2], "role": _normalize_role(row[3]), "created_at": row[4]}


def _admin_email_set() -> set[str]:
    return {
        email.strip().lower()
        for email in os.getenv("ADMIN_EMAILS", "").split(",")
        if email.strip()
    }


def _is_local_request(request: Request) -> bool:
    if request.headers.get("x-forwarded-for") or request.headers.get("x-forwarded-host"):
        return False
    host_header = request.headers.get("host", "")
    if "ngrok" in host_header:
        return False
    host = request.client.host if request.client else ""
    return host in {"127.0.0.1", "::1", "localhost"}


async def require_admin(current_user: dict = Depends(get_current_user)):
    if _normalize_role(current_user.get("role", "")) == "admin":
        return current_user
    raise HTTPException(status_code=403, detail="Admin access required")


# ── Captcha ──────────────────────────────────────────────────────────────────
_CAPTCHA_STORE: Dict[str, Tuple[str, float]] = {}
_CAPTCHA_TTL = 300  # 5 minutes


def _gen_captcha_image(code: str) -> bytes:
    width, height = 160, 60
    img = Image.new("RGB", (width, height), color=(245, 245, 250))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 32)
    except Exception:
        font = ImageFont.load_default()
    for i, char in enumerate(code):
        x = 18 + i * 30 + random.randint(-3, 3)
        y = 12 + random.randint(-6, 6)
        color = (random.randint(20, 80), random.randint(20, 80), random.randint(80, 150))
        draw.text((x, y), char, fill=color, font=font)
    for _ in range(5):
        x1, y1 = random.randint(0, width), random.randint(0, height)
        x2, y2 = random.randint(0, width), random.randint(0, height)
        draw.line([(x1, y1), (x2, y2)], fill=(180, 180, 200), width=1)
    for _ in range(40):
        x, y = random.randint(0, width - 1), random.randint(0, height - 1)
        draw.point((x, y), fill=(150, 150, 180))
    img = img.filter(ImageFilter.SMOOTH)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _cleanup_captchas():
    now = time.time()
    for key in [k for k, (_, exp) in _CAPTCHA_STORE.items() if exp < now]:
        _CAPTCHA_STORE.pop(key, None)


def _validate_captcha(captcha_id: str, captcha_code: str, *, consume: bool) -> None:
    _cleanup_captchas()
    stored = _CAPTCHA_STORE.get(str(captcha_id or ""))
    if not stored or stored[0] != str(captcha_code or "").strip():
        raise HTTPException(status_code=400, detail="Invalid or expired captcha")
    if consume:
        _CAPTCHA_STORE.pop(str(captcha_id or ""), None)


def _generate_email_code() -> str:
    return "".join(secrets.choice(string.digits) for _ in range(_EMAIL_CODE_LENGTH))


async def _assert_email_code_send_allowed(db: aiosqlite.Connection, email: str, purpose: str) -> None:
    cursor = await db.execute(
        """
        SELECT last_sent_at
        FROM email_verification_codes
        WHERE email = ? AND purpose = ? AND consumed_at IS NULL
        ORDER BY datetime(last_sent_at) DESC
        LIMIT 1
        """,
        (_normalize_email(email), purpose),
    )
    row = await cursor.fetchone()
    if not row:
        return
    elapsed = (datetime.now(timezone.utc) - _parse_utc_iso(row[0])).total_seconds()
    if elapsed < _EMAIL_CODE_RESEND_COOLDOWN_SECONDS:
        retry_after = int(_EMAIL_CODE_RESEND_COOLDOWN_SECONDS - elapsed)
        raise HTTPException(
            status_code=429,
            detail=f"Email verification code was sent recently. Try again in {retry_after} seconds.",
        )


async def _store_email_code(db: aiosqlite.Connection, email: str, code: str, purpose: str) -> None:
    now = _utc_now_iso()
    await db.execute(
        """
        INSERT INTO email_verification_codes (
            id, email, purpose, code_hash, created_at, expires_at, consumed_at, attempts, last_sent_at
        ) VALUES (?,?,?,?,?,?,?,?,?)
        """,
        (
            str(uuid.uuid4()),
            _normalize_email(email),
            purpose,
            _hash_email_code(email, code),
            now,
            _utc_in_seconds_iso(_EMAIL_CODE_TTL_SECONDS),
            None,
            0,
            now,
        ),
    )


async def _verify_email_code(db: aiosqlite.Connection, email: str, code: Optional[str], purpose: str) -> None:
    clean_code = str(code or "").strip()
    if not clean_code:
        raise HTTPException(status_code=400, detail="Email verification code is required")
    cursor = await db.execute(
        """
        SELECT id, code_hash, expires_at, attempts
        FROM email_verification_codes
        WHERE email = ? AND purpose = ? AND consumed_at IS NULL
        ORDER BY datetime(created_at) DESC
        LIMIT 1
        """,
        (_normalize_email(email), purpose),
    )
    row = await cursor.fetchone()
    if not row:
        raise HTTPException(status_code=400, detail="Invalid or expired email verification code")
    if _parse_utc_iso(row[2]) <= datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="Invalid or expired email verification code")
    attempts = int(row[3] or 0)
    if attempts >= _EMAIL_CODE_MAX_ATTEMPTS:
        raise HTTPException(status_code=400, detail="Email verification code attempt limit exceeded")
    if not _check_email_code(email, clean_code, row[1]):
        await db.execute(
            "UPDATE email_verification_codes SET attempts = attempts + 1 WHERE id = ?",
            (row[0],),
        )
        raise HTTPException(status_code=400, detail="Invalid or expired email verification code")
    await db.execute(
        "UPDATE email_verification_codes SET consumed_at = ? WHERE id = ?",
        (_utc_now_iso(), row[0]),
    )


def _deliver_email_verification_code(*, email: str, code: str) -> None:
    if not _MAIL_ENABLED:
        if _is_production_like_env():
            raise HTTPException(status_code=503, detail="Email delivery is not configured")
        print(f"[email-verification] MAIL_ENABLED=false; code for {_normalize_email(email)}: {code}")
        return

    missing = [
        name
        for name, value in {
            "MAIL_SMTP_HOST": _MAIL_SMTP_HOST,
            "MAIL_SMTP_USER": _MAIL_SMTP_USER,
            "MAIL_SMTP_PASSWORD": _MAIL_SMTP_PASSWORD,
            "MAIL_FROM": _MAIL_FROM,
        }.items()
        if not value
    ]
    if missing:
        raise HTTPException(status_code=503, detail=f"Email delivery is missing configuration: {', '.join(missing)}")

    message = EmailMessage()
    message["Subject"] = "Your CausalGraph AI verification code"
    message["From"] = formataddr((_MAIL_FROM_NAME, _MAIL_FROM))
    message["To"] = _normalize_email(email)
    message.set_content(
        "\n".join(
            [
                "Your CausalGraph AI verification code is:",
                "",
                code,
                "",
                f"This code expires in {int(_EMAIL_CODE_TTL_SECONDS / 60)} minutes.",
                "If you did not request this code, you can ignore this email.",
            ]
        )
    )

    try:
        if _MAIL_SMTP_SSL:
            with smtplib.SMTP_SSL(_MAIL_SMTP_HOST, _MAIL_SMTP_PORT, timeout=15) as smtp:
                smtp.login(_MAIL_SMTP_USER, _MAIL_SMTP_PASSWORD)
                smtp.send_message(message)
        else:
            with smtplib.SMTP(_MAIL_SMTP_HOST, _MAIL_SMTP_PORT, timeout=15) as smtp:
                if _MAIL_SMTP_STARTTLS:
                    smtp.starttls()
                smtp.login(_MAIL_SMTP_USER, _MAIL_SMTP_PASSWORD)
                smtp.send_message(message)
    except smtplib.SMTPException as exc:
        raise HTTPException(status_code=502, detail="Email delivery failed") from exc


class EmailCodeSendRequest(BaseModel):
    email: str
    captcha_id: str
    captcha_code: str


class RegisterRequest(BaseModel):
    email: str
    username: str
    password: str
    captcha_id: str
    captcha_code: str
    email_code: Optional[str] = None
    role: str = "user"
    admin_invite_code: Optional[str] = None


class LoginRequest(BaseModel):
    email: str
    password: str


class AdminInviteCreateRequest(BaseModel):
    ttl_minutes: int = 5


class RagUnlimitedUserRequest(BaseModel):
    email: str
    note: Optional[str] = ""


async def _consume_admin_invite(code: str, user_id: str, db: aiosqlite.Connection) -> None:
    invite_code = str(code or "").strip()
    if not invite_code:
        raise HTTPException(status_code=400, detail="Admin registration requires an invitation code")
    cursor = await db.execute(
        """
        SELECT id, expires_at, used_at
        FROM admin_invite_codes
        WHERE code = ?
        """,
        (invite_code,),
    )
    row = await cursor.fetchone()
    if not row:
        raise HTTPException(status_code=400, detail="Invalid admin invitation code")
    if row[2]:
        raise HTTPException(status_code=400, detail="Admin invitation code has already been used")
    expires_at = datetime.fromisoformat(str(row[1]))
    if expires_at <= datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="Admin invitation code has expired")

    await db.execute(
        """
        UPDATE admin_invite_codes
        SET used_at = ?, used_by_user_id = ?
        WHERE id = ?
        """,
        (_utc_now_iso(), user_id, row[0]),
    )

from ai_service.extractor import extract_esg
from ai_service.schemas import EsgExtractionRequest, EsgExtractionResponse
from chat_memory_service import RedisUnavailableError, chat_memory_service
from configs.settings import CHUNK_DIR, DATA_DIR, EMBEDDING_FALLBACK_DIM, GRAPH_DIR, NEO4J_AUTO_SYNC, VECTOR_DIR, VECTOR_STORE_PROVIDER, neo4j_configured
from user_memory_service import (
    delete_user_memory,
    format_memories_for_prompt,
    get_memory_settings,
    get_relevant_user_memories,
    init_user_memory_db,
    list_user_memories,
    remember_exchange,
    update_memory_settings,
    update_user_memory,
)
import document_registry
from graph.causal_reasoning import CausalReasoner
from graph.neo4j_store import assert_neo4j_ready, get_neo4j_store, neo4j_sdk_available
from admin_audit import (
    admin_overview,
    get_upload,
    get_latest_upload_by_document_id,
    init_admin_db,
    list_latest_uploads_by_document_id,
    list_uploads,
    mark_upload_deleted,
    record_upload_cleanup,
    update_upload_metadata,
)
from ingestion_jobs import get_ingestion_job, start_ingestion_job
from pipeline_runtime import (
    delete_uploaded_document,
    ingest_uploaded_document,
    load_registered_document,
    rebuild_document_graph,
    summarize_registered_document,
)
from rag.bm25_index import warm_bm25_index
from rag.embeddings import embedding_backend_is_real, get_embedding_backend, get_embedding_model
from rag.openai_client import get_openai_client
from rag.openai_compat import chat_token_kwargs
from rag.rag_pipeline import answer_question, stream_answer_question
from rag.answer_intent import classify_answer_intent
from rag.query_rewriter import format_history
from rag.retriever import retrieve_context
from scripts.run_pdf_pipeline import run_pdf_pipeline


app = FastAPI(title="ESG QLoRA Extraction API", version="1.0.0")

_APP_ROOT = Path(__file__).resolve().parent
_KG_VIEW_TEMPLATE = _APP_ROOT / "kg_view" / "templates" / "index.html"
_KG_VIEW_STATIC = _APP_ROOT / "kg_view" / "static"
_TEXT_KG_WEB_ROOT = _APP_ROOT / "text-to-kg-esg" / "web"
_TEXT_KG_VIEW_TEMPLATE = _TEXT_KG_WEB_ROOT / "templates" / "index.html"
_TEXT_KG_VIEW_STATIC = _TEXT_KG_WEB_ROOT / "static"
_KG_VIEW_LLM_CLUSTER_LABELS = _env_flag("KG_VIEW_LLM_CLUSTER_LABELS", "true")
_KG_VIEW_CLUSTER_LABEL_MODEL = os.getenv(
    "KG_VIEW_CLUSTER_LABEL_MODEL",
    os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
).strip() or "gpt-4o-mini"
_KG_VIEW_CLUSTER_LABEL_CACHE: Dict[str, Dict[str, str]] = {}
_KG_VIEW_CACHE_ENABLED = _env_flag("KG_VIEW_CACHE_ENABLED", "true")
_KG_VIEW_CACHE_VERSION = os.getenv("KG_VIEW_CACHE_VERSION", f"{_PUBLIC_GRAPH_CACHE_VERSION}-kgv4").strip() or f"{_PUBLIC_GRAPH_CACHE_VERSION}-kgv4"
_KG_VIEW_CACHE_TTL_SECONDS = max(60, int(os.getenv("KG_VIEW_CACHE_TTL_SECONDS", str(_PUBLIC_GRAPH_CACHE_TTL_SECONDS))))
_KG_VIEW_CACHE_STALE_SECONDS = max(
    _KG_VIEW_CACHE_TTL_SECONDS,
    int(os.getenv("KG_VIEW_CACHE_STALE_SECONDS", str(_PUBLIC_GRAPH_CACHE_STALE_SECONDS))),
)
_KG_VIEW_CACHE_LOCK_SECONDS = max(10, int(os.getenv("KG_VIEW_CACHE_LOCK_SECONDS", str(_PUBLIC_GRAPH_CACHE_LOCK_SECONDS))))
_KG_VIEW_CACHE_WAIT_SECONDS = max(1.0, float(os.getenv("KG_VIEW_CACHE_WAIT_SECONDS", str(_PUBLIC_GRAPH_CACHE_WAIT_SECONDS))))
_KG_VIEW_CACHE_MEMORY_MAX = max(16, int(os.getenv("KG_VIEW_CACHE_MEMORY_MAX", "128")))
_KG_VIEW_DETAIL_RENDER_NODE_LIMIT = max(100, int(os.getenv("KG_VIEW_DETAIL_RENDER_NODE_LIMIT", "2000")))
_KG_VIEW_REDIS_RETRY_SECONDS = max(1.0, float(os.getenv("KG_VIEW_REDIS_RETRY_SECONDS", "15")))
_KG_VIEW_CACHE_MEMORY: Dict[str, Dict[str, Any]] = {}
_KG_VIEW_CACHE_MEMORY_LOCKS: Dict[str, threading.Lock] = {}
_KG_VIEW_CACHE_MEMORY_GUARD = threading.Lock()
_KG_VIEW_REDIS_DISABLED_UNTIL = 0.0
_KG_VIEW_TICKET_TTL_SECONDS = max(60, int(os.getenv("KG_VIEW_TICKET_TTL_SECONDS", "3600")))
_KG_VIEW_TICKETS: Dict[str, Dict[str, Any]] = {}
_KG_VIEW_TICKETS_LOCK = threading.Lock()


class KgViewTicketRequest(BaseModel):
    document_id: Optional[str] = ""


def _assert_embedding_dim_matches_pinecone() -> None:
    if VECTOR_STORE_PROVIDER != "pinecone":
        return
    get_embedding_model()
    if not embedding_backend_is_real():
        print(
            f"[startup] WARNING: embedding backend={get_embedding_backend()} "
            f"(dim={EMBEDDING_FALLBACK_DIM}) does NOT match Pinecone 1024-d index. "
            "Vector retrieval will be skipped; BM25-only fallback active."
        )


def _ensure_neo4j_schema_startup() -> None:
    try:
        from graph.neo4j_store import get_neo4j_store, neo4j_enabled

        if neo4j_enabled():
            store = get_neo4j_store()
            if store is not None:
                store.setup_schema()
                print("[startup] Neo4j schema (indexes + fulltext) ensured.")
    except Exception as exc:
        print(f"[startup] Neo4j schema setup skipped: {type(exc).__name__}: {exc}")

app.add_middleware(
    CORSMiddleware,
    allow_origins=_CORS_ALLOW_ORIGINS,
    allow_origin_regex=_CORS_ALLOW_ORIGIN_REGEX,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.add_middleware(GZipMiddleware, minimum_size=1024)

app.mount("/kg-static", StaticFiles(directory=str(_KG_VIEW_STATIC)), name="kg-static")
if _TEXT_KG_VIEW_STATIC.exists():
    app.mount("/static", StaticFiles(directory=str(_TEXT_KG_VIEW_STATIC)), name="text_to_kg_static")


@app.on_event("startup")
async def startup():
    _validate_startup_security_config()
    await _init_auth_db()
    await _init_feedback_db()
    init_admin_db()
    warm_bm25_index()
    _ensure_neo4j_schema_startup()
    _assert_embedding_dim_matches_pinecone()


@app.get("/auth/captcha")
async def get_captcha():
    _cleanup_captchas()
    code = "".join(random.choices(string.digits, k=4))
    captcha_id = str(uuid.uuid4())
    _CAPTCHA_STORE[captcha_id] = (code, time.time() + _CAPTCHA_TTL)
    img_b64 = base64.b64encode(_gen_captcha_image(code)).decode()
    return {"captcha_id": captcha_id, "image": f"data:image/png;base64,{img_b64}"}


@app.post("/auth/email-code/send")
async def send_email_code(req: EmailCodeSendRequest, db: aiosqlite.Connection = Depends(_get_db)):
    _validate_captcha(req.captcha_id, req.captcha_code, consume=False)
    email = _normalize_email(req.email)
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="A valid email address is required")
    cursor = await db.execute("SELECT id FROM users WHERE email = ?", (email,))
    if await cursor.fetchone():
        raise HTTPException(status_code=400, detail="Email already registered")
    await _assert_email_code_send_allowed(db, email, "register")
    code = _generate_email_code()
    await _store_email_code(db, email, code, "register")
    _deliver_email_verification_code(email=email, code=code)
    await db.commit()
    return {
        "sent": True,
        "ttl_seconds": _EMAIL_CODE_TTL_SECONDS,
        "cooldown_seconds": _EMAIL_CODE_RESEND_COOLDOWN_SECONDS,
    }


@app.post("/auth/register")
async def register(req: RegisterRequest, db: aiosqlite.Connection = Depends(_get_db)):
    _validate_captcha(req.captcha_id, req.captcha_code, consume=True)
    email = _normalize_email(req.email)
    cursor = await db.execute("SELECT id FROM users WHERE email = ?", (email,))
    if await cursor.fetchone():
        raise HTTPException(status_code=400, detail="Email already registered")
    await _verify_email_code(db, email, req.email_code, "register")
    user_id = str(uuid.uuid4())
    role = _normalize_role(req.role)
    if role == "admin":
        await _consume_admin_invite(req.admin_invite_code or "", user_id, db)
    await db.execute(
        "INSERT INTO users (id, email, username, password_hash, role, created_at) VALUES (?,?,?,?,?,?)",
        (user_id, email, req.username, _hash_pw(req.password), role, _utc_now_iso()),
    )
    await db.commit()
    token = _make_token(user_id, email)
    return {
        "token": token,
        "user": {
            "id": user_id,
            "email": email,
            "username": req.username,
            "role": role,
            "created_at": _utc_now_iso(),
        },
    }


@app.post("/auth/login")
async def login(req: LoginRequest, db: aiosqlite.Connection = Depends(_get_db)):
    cursor = await db.execute(
        "SELECT id, email, username, password_hash, role, created_at FROM users WHERE email = ?", (req.email,)
    )
    row = await cursor.fetchone()
    if not row or not _check_pw(req.password, row[3]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = _make_token(row[0], row[1])
    return {
        "token": token,
        "user": {
            "id": row[0],
            "email": row[1],
            "username": row[2],
            "role": _normalize_role(row[4]),
            "created_at": row[5],
        },
    }


@app.get("/auth/me")
async def me(current_user: dict = Depends(get_current_user)):
    return current_user


@app.get("/admin/overview")
async def admin_dashboard_overview(
    days: int = Query(default=14, ge=1, le=90),
    current_user: dict = Depends(require_admin),
):
    return JSONResponse(content=admin_overview(days=days))


@app.get("/admin/uploads")
async def admin_uploads(
    limit: int = Query(default=100, ge=1, le=500),
    offset: int = Query(default=0, ge=0),
    current_user: dict = Depends(require_admin),
):
    return JSONResponse(content={"uploads": list_uploads(limit=limit, offset=offset)})


@app.post("/admin/invite-codes")
async def create_admin_invite_code(
    request: AdminInviteCreateRequest,
    current_user: dict = Depends(require_admin),
    db: aiosqlite.Connection = Depends(_get_db),
):
    ttl_minutes = max(1, min(int(request.ttl_minutes), 5))
    code = f"ADM-{secrets.token_urlsafe(8).replace('-', '').replace('_', '')[:10].upper()}"
    invite_id = str(uuid.uuid4())
    now = _utc_now_iso()
    expires_at = _utc_in_minutes_iso(ttl_minutes)
    await db.execute(
        """
        INSERT INTO admin_invite_codes (
            id, code, created_by_user_id, created_at, expires_at, used_at, used_by_user_id
        ) VALUES (?, ?, ?, ?, ?, NULL, NULL)
        """,
        (invite_id, code, str(current_user.get("id") or ""), now, expires_at),
    )
    await db.commit()
    return {
        "invite_code": code,
        "expires_at": expires_at,
        "ttl_minutes": ttl_minutes,
        "single_use": True,
    }


def _rag_unlimited_user_payload(row: Tuple[Any, ...]) -> Dict[str, Any]:
    return {
        "email": row[0],
        "note": row[1],
        "created_by_user_id": row[2],
        "created_at": row[3],
    }


@app.get("/admin/rag-unlimited-users")
async def list_rag_unlimited_users(
    current_user: dict = Depends(require_admin),
    db: aiosqlite.Connection = Depends(_get_db),
):
    cursor = await db.execute(
        """
        SELECT email, note, created_by_user_id, created_at
        FROM rag_unlimited_users
        ORDER BY lower(email)
        """
    )
    rows = await cursor.fetchall()
    return {"users": [_rag_unlimited_user_payload(row) for row in rows]}


@app.post("/admin/rag-unlimited-users")
async def add_rag_unlimited_user(
    request: RagUnlimitedUserRequest,
    current_user: dict = Depends(require_admin),
    db: aiosqlite.Connection = Depends(_get_db),
):
    email = _normalize_email(request.email)
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="A valid email address is required")
    note = str(request.note or "").strip()[:240]
    await db.execute(
        """
        INSERT INTO rag_unlimited_users (email, note, created_by_user_id, created_at)
        VALUES (?, ?, ?, ?)
        ON CONFLICT(email) DO UPDATE SET
            note = excluded.note,
            created_by_user_id = excluded.created_by_user_id,
            created_at = excluded.created_at
        """,
        (email, note, str(current_user.get("id") or ""), _utc_now_iso()),
    )
    await db.commit()
    cursor = await db.execute(
        """
        SELECT email, note, created_by_user_id, created_at
        FROM rag_unlimited_users
        WHERE email = ?
        """,
        (email,),
    )
    row = await cursor.fetchone()
    return {"user": _rag_unlimited_user_payload(row)}


@app.delete("/admin/rag-unlimited-users/{email}")
async def delete_rag_unlimited_user(
    email: str,
    current_user: dict = Depends(require_admin),
    db: aiosqlite.Connection = Depends(_get_db),
):
    normalized_email = _normalize_email(unquote(email))
    cursor = await db.execute("DELETE FROM rag_unlimited_users WHERE email = ?", (normalized_email,))
    await db.commit()
    return {"deleted": cursor.rowcount > 0, "email": normalized_email}


class AdminUploadUpdateRequest(BaseModel):
    title: Optional[str] = None
    domain: Optional[str] = None
    source_type: Optional[str] = None
    source: Optional[str] = None


@app.patch("/admin/uploads/{job_id}")
async def admin_update_upload(
    job_id: str,
    request: AdminUploadUpdateRequest,
    current_user: dict = Depends(require_admin),
):
    updates = request.dict(exclude_unset=True)
    updated = update_upload_metadata(job_id, updates)
    if updated is None:
        raise HTTPException(status_code=404, detail="Upload record not found")
    document_id = str(updated.get("document_id") or "").strip()
    if document_id:
        document_registry.update_metadata(document_id, updates)
    return JSONResponse(content={"upload": updated})


@app.delete("/admin/uploads/{job_id}")
async def admin_delete_upload(
    job_id: str,
    reason: str = Query(default=""),
    current_user: dict = Depends(require_admin),
):
    upload = get_upload(job_id)
    if upload is None:
        raise HTTPException(status_code=404, detail="Upload record not found")
    status = str(upload.get("status") or "")
    if status in {"queued", "running"}:
        raise HTTPException(status_code=409, detail="Running or queued uploads cannot be deleted until processing finishes")
    deleted = mark_upload_deleted(
        job_id,
        deleted_by=str(current_user.get("email") or current_user.get("username") or ""),
        reason=reason,
        status="deleted",
        cleanup_status="cleanup_skipped" if status == "rejected" else "cleanup_pending",
        cleanup_detail="Rejected upload has no indexed resources." if status == "rejected" else "Cleanup queued.",
    )
    if deleted is None:
        raise HTTPException(status_code=404, detail="Upload record not found")
    if status != "rejected":
        _CLEANUP_EXECUTOR.submit(
            _cleanup_deleted_upload,
            job_id,
            upload,
            str(current_user.get("email") or current_user.get("username") or ""),
            reason,
        )
    return JSONResponse(
        content={
            "upload": deleted,
            "cleanup": {
                "queued": status != "rejected",
                "warnings": [],
                "neo4j": {"enabled": False, "deleted": False, "reason": "background_cleanup_queued"},
            },
        }
    )


def _cleanup_deleted_upload(job_id: str, upload: Dict[str, Any], deleted_by: str, reason: str) -> None:
    try:
        cleanup = delete_uploaded_document(upload)
        document_id = str(upload.get("document_id") or "").strip()
        if document_id:
            document_registry.remove(document_id)
        warnings = list(cleanup.get("warnings") or [])
        neo4j_result = cleanup.get("neo4j") or {}
        if neo4j_result.get("enabled") and neo4j_result.get("deleted") is False and neo4j_result.get("reason") not in {
            "missing_document_id",
        }:
            warnings.append(f"neo4j: {neo4j_result.get('reason') or 'delete_not_confirmed'}")
        if warnings:
            record_upload_cleanup(
                job_id,
                cleanup_status="cleanup_failed",
                cleanup_detail=f"Cleanup warnings: {'; '.join(warnings[:5])}",
                status="deleted_with_warnings",
            )
        else:
            deleted_count = len(cleanup.get("deleted_paths") or [])
            record_upload_cleanup(
                job_id,
                cleanup_status="cleanup_completed",
                cleanup_detail=f"Deleted {deleted_count} local path(s). Neo4j: {neo4j_result.get('reason') or neo4j_result.get('deleted')}.",
                status="deleted",
            )
    except Exception as exc:
        record_upload_cleanup(
            job_id,
            cleanup_status="cleanup_failed",
            cleanup_detail=f"Cleanup failed: {type(exc).__name__}: {exc}",
            status="deleted_with_warnings",
        )


class RagAskRequest(BaseModel):
    question: str
    top_k: int = 5
    history: List[Dict[str, Any]] = []
    session_id: Optional[str] = None
    document_ids: List[str] = []
    preferred_document_id: Optional[str] = None
    document_group: Optional[str] = None
    source_type: Optional[str] = None
    domain: Optional[str] = None
    # Deprecated: legacy intent selector (ask | predict | graph). The pipeline
    # now drives behavior entirely off `reasoning_mode` (flash | deep). Kept on
    # the schema so older clients don't fail validation; the value is ignored.
    mode: Optional[str] = "ask"
    reasoning_mode: Optional[str] = "flash"


class DesktopScreenshotSummaryRequest(BaseModel):
    image_data_url: str
    prompt: Optional[str] = Field(
        default="Summarize the visible information and point out anything that needs attention.",
        max_length=1200,
    )


class DesktopWordEditExportRequest(BaseModel):
    session_id: str = Field(..., min_length=8, max_length=80)
    accepted_suggestion_ids: List[str] = []


def _validate_desktop_image_data_url(value: str) -> str:
    data_url = str(value or "").strip()
    if not data_url.startswith("data:image/"):
        raise HTTPException(status_code=400, detail="Screenshot must be an image data URL")
    try:
        header, encoded = data_url.split(",", 1)
    except ValueError:
        raise HTTPException(status_code=400, detail="Screenshot data URL is malformed")
    header_lower = header.lower()
    if ";base64" not in header_lower:
        raise HTTPException(status_code=400, detail="Screenshot image must be base64 encoded")
    mime = header_lower.removeprefix("data:").split(";", 1)[0]
    if mime not in {"image/png", "image/jpeg", "image/jpg", "image/webp"}:
        raise HTTPException(status_code=400, detail="Unsupported screenshot image type")
    try:
        raw = base64.b64decode(encoded, validate=True)
    except Exception:
        raise HTTPException(status_code=400, detail="Screenshot image payload is not valid base64")
    if not raw:
        raise HTTPException(status_code=400, detail="Screenshot image is empty. Capture the screen again after granting screen recording permission.")
    if len(raw) > _DESKTOP_SCREENSHOT_MAX_IMAGE_BYTES:
        raise HTTPException(status_code=413, detail="Screenshot image is too large")
    normalized_mime = "image/jpeg" if mime == "image/jpg" else mime
    return f"data:{normalized_mime};base64,{base64.b64encode(raw).decode('ascii')}"


def _extract_chat_completion_text(response: Any) -> str:
    choice = response.choices[0]
    content = getattr(getattr(choice, "message", None), "content", "")
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict) and item.get("type") == "text":
                parts.append(str(item.get("text") or ""))
        return "\n".join(parts).strip()
    return str(content or "").strip()


def _safe_word_filename(filename: Optional[str]) -> str:
    name = Path(unquote(str(filename or "document.docx"))).name
    name = re.sub(r"[^A-Za-z0-9._ ()-]+", "_", name).strip(" ._") or "document.docx"
    if not name.lower().endswith(".docx"):
        name = f"{Path(name).stem or 'document'}.docx"
    return name


def _word_user_key(current_user: Dict[str, Any]) -> str:
    raw = str(current_user.get("id") or current_user.get("email") or "user").strip()
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", raw) or "user"


def _word_edit_user_root(current_user: Dict[str, Any]) -> Path:
    root = DATA_DIR / "word_edits" / _word_user_key(current_user)
    root.mkdir(parents=True, exist_ok=True)
    return root


def _word_edit_session_dir(current_user: Dict[str, Any], session_id: str) -> Path:
    normalized = re.sub(r"[^A-Za-z0-9_.-]+", "", str(session_id or ""))
    if len(normalized) < 8:
        raise HTTPException(status_code=400, detail="Invalid Word edit session")
    return _word_edit_user_root(current_user) / normalized


def _word_session_json_path(current_user: Dict[str, Any], session_id: str) -> Path:
    return _word_edit_session_dir(current_user, session_id) / "session.json"


def _parse_docx_paragraphs(file_bytes: bytes) -> Tuple[Any, List[Dict[str, Any]]]:
    document = DocxDocument(io.BytesIO(file_bytes))
    paragraphs: List[Dict[str, Any]] = []
    visible_index = 0
    for docx_index, paragraph in enumerate(document.paragraphs):
        text = re.sub(r"\s+", " ", str(paragraph.text or "")).strip()
        if not text:
            continue
        visible_index += 1
        paragraphs.append({
            "id": f"p_{visible_index:03d}",
            "docx_index": docx_index,
            "text": text,
        })
    if not paragraphs:
        raise HTTPException(status_code=400, detail="The Word document does not contain extractable paragraph text.")
    return document, paragraphs


def _select_word_paragraphs_for_review(paragraphs: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    selected: List[Dict[str, str]] = []
    total_chars = 0
    for paragraph in paragraphs:
        text = str(paragraph.get("text") or "").strip()
        if len(text) < 35 and len(selected) >= 3:
            continue
        clipped = text[:1400]
        if selected and total_chars + len(clipped) > _DESKTOP_WORD_EDIT_MAX_CHARS:
            break
        selected.append({"id": str(paragraph.get("id") or ""), "text": clipped})
        total_chars += len(clipped)
        if len(selected) >= _DESKTOP_WORD_EDIT_MAX_PARAGRAPHS:
            break
    return selected or [{"id": str(paragraphs[0]["id"]), "text": str(paragraphs[0]["text"])[:1400]}]


def _parse_json_object_text(raw: str) -> Dict[str, Any]:
    text = str(raw or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s*```$", "", text)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start < 0 or end <= start:
            raise
        parsed = json.loads(text[start:end + 1])
    if not isinstance(parsed, dict):
        raise ValueError("Model response was not a JSON object")
    return parsed


_WORD_EDIT_CATEGORY_LABELS = {
    "clarity": "Clarity",
    "logic": "Logic",
    "evidence": "Evidence",
    "structure": "Structure",
    "tone": "Tone",
    "esg_concept": "ESG concept",
}

_WORD_EDIT_GOAL_LABELS = {
    "academic": "Academic analysis",
    "business": "Business report",
    "executive": "Executive summary",
    "esg": "ESG analysis",
}

_WORD_EDIT_TEMPLATE_LABELS = {
    "general": "General review",
    "esg_report": "ESG analysis report",
    "company_comparison": "Company comparison report",
    "risk_assessment": "Risk assessment",
    "sustainability_strategy": "Sustainability strategy analysis",
    "causal_impact": "Causal impact analysis",
    "literature_paragraph": "Literature-style paragraph",
    "executive_summary": "Executive summary",
}

_WORD_EDIT_EVIDENCE_GAP_LABELS = {
    "metric": "metric",
    "source": "source",
    "comparison": "comparison",
    "causal_link": "causal link",
    "citation": "citation",
    "concept_definition": "concept definition",
}


def _normalize_word_edit_goal(value: str) -> str:
    goal = str(value or "").strip().lower()
    return goal if goal in _WORD_EDIT_GOAL_LABELS else "academic"


def _normalize_word_edit_template(value: str) -> str:
    template = str(value or "").strip().lower()
    return template if template in _WORD_EDIT_TEMPLATE_LABELS else "general"


def _word_edit_goal_instruction(goal: str) -> str:
    normalized = _normalize_word_edit_goal(goal)
    instructions = {
        "academic": "Prioritize academic clarity, cautious claims, argument flow, and explicit evidence gaps.",
        "business": "Prioritize concise business reporting, decision relevance, metrics, risks, and management implications.",
        "executive": "Prioritize executive brevity, readable recommendations, business impact, and action-oriented wording.",
        "esg": "Prioritize ESG terminology, standards-aware phrasing, materiality, metrics, risk, governance, and evidence traceability.",
    }
    return instructions[normalized]


def _word_edit_template_instruction(template: str) -> str:
    normalized = _normalize_word_edit_template(template)
    instructions = {
        "general": "Review the draft without forcing a specific document structure.",
        "esg_report": "Check whether the draft has ESG context, material issues, metrics, evidence, risks, and implications.",
        "company_comparison": "Check whether the draft compares companies using consistent metrics, evidence, and like-for-like logic.",
        "risk_assessment": "Check whether the draft identifies risk drivers, likelihood, impact, controls, and evidence.",
        "sustainability_strategy": "Check whether the draft links sustainability initiatives to strategy, resources, metrics, and trade-offs.",
        "causal_impact": "Check whether causal claims distinguish evidence, correlation, mechanism, counterfactuals, and uncertainty.",
        "literature_paragraph": "Check whether the paragraph has a clear claim, prior literature positioning, synthesis, and citation gaps.",
        "executive_summary": "Check whether the draft is concise, decision-oriented, and explicit about business implications.",
    }
    return instructions[normalized]


def _build_word_evidence_query(paragraphs: List[Dict[str, Any]], instruction: str, goal: str, template: str) -> str:
    selected = _select_word_paragraphs_for_review(paragraphs)
    text = "\n".join(str(item.get("text") or "") for item in selected[:8])
    return "\n".join(
        part
        for part in [
            _WORD_EDIT_GOAL_LABELS.get(_normalize_word_edit_goal(goal), "Academic analysis"),
            _WORD_EDIT_TEMPLATE_LABELS.get(_normalize_word_edit_template(template), "General review"),
            str(instruction or "").strip(),
            text[:4000],
        ]
        if part
    )


def _serialize_word_evidence_sources(sources: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    serialized: List[Dict[str, Any]] = []
    for index, source in enumerate(sources[:_DESKTOP_WORD_EDIT_EVIDENCE_TOP_K], start=1):
        text = re.sub(r"\s+", " ", str(source.get("text") or "")).strip()
        if not text:
            continue
        title = str(source.get("title") or source.get("document_title") or source.get("source") or source.get("document_id") or "Evidence").strip()
        chunk_id = str(source.get("chunk_id") or source.get("id") or "").strip()
        serialized.append({
            "id": f"E{index}",
            "title": title,
            "document_id": str(source.get("document_id") or ""),
            "chunk_id": chunk_id,
            "text": text[:700],
            "score": float(source.get("score") or 0.0),
        })
    return serialized


def _retrieve_word_review_evidence(
    *,
    current_user: Dict[str, Any],
    paragraphs: List[Dict[str, Any]],
    instruction: str,
    goal: str,
    template: str,
) -> List[Dict[str, Any]]:
    if _DESKTOP_WORD_EDIT_EVIDENCE_TOP_K <= 0:
        return []
    filters: Dict[str, Any] = {}
    if current_user and not _is_admin_user(current_user):
        filters["owner_user_id"] = str(current_user.get("id") or "")
    query = _build_word_evidence_query(paragraphs, instruction, goal, template)
    try:
        return _serialize_word_evidence_sources(
            retrieve_context(query=query, top_k=_DESKTOP_WORD_EDIT_EVIDENCE_TOP_K, filters=filters)
        )
    except Exception as exc:
        print(f"[desktop.word] evidence retrieval skipped: {type(exc).__name__}: {exc}")
        return []


def _normalize_word_edit_suggestions(
    parsed: Dict[str, Any],
    paragraph_lookup: Dict[str, Dict[str, Any]],
    evidence_sources: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, Any]]:
    raw_suggestions = parsed.get("suggestions")
    if not isinstance(raw_suggestions, list):
        return []
    allowed_evidence_ids = {str(item.get("id") or "") for item in (evidence_sources or [])}
    suggestions: List[Dict[str, Any]] = []
    seen_paragraphs: set[str] = set()
    for raw_item in raw_suggestions:
        if not isinstance(raw_item, dict):
            continue
        paragraph_id = str(raw_item.get("paragraph_id") or raw_item.get("paragraph") or "").strip()
        if paragraph_id not in paragraph_lookup or paragraph_id in seen_paragraphs:
            continue
        replacement = re.sub(r"\s+", " ", str(raw_item.get("replacement") or "")).strip()
        if not replacement:
            continue
        original = str(paragraph_lookup[paragraph_id].get("text") or "").strip()
        if replacement == original:
            continue
        reason = re.sub(r"\s+", " ", str(raw_item.get("reason") or "Improves clarity and analytical quality.")).strip()
        problem = re.sub(
            r"\s+",
            " ",
            str(raw_item.get("problem") or raw_item.get("problem_solved") or "Improves draft quality.").strip(),
        )
        category = str(raw_item.get("category") or "clarity").strip().lower()
        if category not in _WORD_EDIT_CATEGORY_LABELS:
            category = "clarity"
        severity = str(raw_item.get("severity") or "").strip().lower()
        if severity not in {"low", "medium", "high"}:
            severity = "medium" if category in {"logic", "evidence", "esg_concept"} else "low"
        raw_refs = raw_item.get("evidence_refs") if isinstance(raw_item.get("evidence_refs"), list) else []
        evidence_refs = [str(ref) for ref in raw_refs if str(ref) in allowed_evidence_ids]
        raw_gaps = raw_item.get("evidence_gap_types") if isinstance(raw_item.get("evidence_gap_types"), list) else []
        evidence_gap_types = [
            str(gap).strip().lower()
            for gap in raw_gaps
            if str(gap).strip().lower() in _WORD_EDIT_EVIDENCE_GAP_LABELS
        ]
        evidence_needed = bool(raw_item.get("evidence_needed")) or (category == "evidence" and not evidence_refs)
        if evidence_needed and not evidence_gap_types:
            evidence_gap_types = ["source"]
        suggestions.append({
            "id": f"s_{len(suggestions) + 1:03d}",
            "paragraph_id": paragraph_id,
            "operation": "replace",
            "category": category,
            "category_label": _WORD_EDIT_CATEGORY_LABELS[category],
            "severity": severity,
            "original": original,
            "replacement": replacement,
            "problem": problem[:260],
            "reason": reason[:500],
            "evidence_refs": evidence_refs,
            "evidence_needed": evidence_needed,
            "evidence_gap_types": evidence_gap_types,
            "sources": [paragraph_id],
        })
        seen_paragraphs.add(paragraph_id)
        if len(suggestions) >= _DESKTOP_WORD_EDIT_MAX_SUGGESTIONS:
            break
    return suggestions


def _fallback_word_edit_suggestions(paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    suggestions: List[Dict[str, Any]] = []
    for paragraph in paragraphs:
        text = str(paragraph.get("text") or "").strip()
        normalized = re.sub(r"\s+", " ", text).strip()
        if normalized and normalized != text:
            suggestions.append({
                "id": f"s_{len(suggestions) + 1:03d}",
                "paragraph_id": str(paragraph.get("id") or ""),
                "operation": "replace",
                "category": "clarity",
                "category_label": _WORD_EDIT_CATEGORY_LABELS["clarity"],
                "severity": "low",
                "original": text,
                "replacement": normalized,
                "problem": "The paragraph contains inconsistent spacing that makes the draft less polished.",
                "reason": "Normalizes spacing without changing the meaning.",
                "evidence_refs": [],
                "evidence_needed": False,
                "evidence_gap_types": [],
                "sources": [str(paragraph.get("id") or "")],
            })
        if len(suggestions) >= 3:
            break
    return suggestions


def _generate_word_edit_suggestions(
    instruction: str,
    paragraphs: List[Dict[str, Any]],
    *,
    goal: str = "academic",
    template: str = "general",
    evidence_sources: Optional[List[Dict[str, Any]]] = None,
) -> Tuple[List[Dict[str, Any]], str]:
    client = get_openai_client()
    if client is None:
        return _fallback_word_edit_suggestions(paragraphs), "fallback_no_openai"

    paragraph_lookup = {str(item.get("id") or ""): item for item in paragraphs}
    review_paragraphs = _select_word_paragraphs_for_review(paragraphs)
    normalized_goal = _normalize_word_edit_goal(goal)
    normalized_template = _normalize_word_edit_template(template)
    prompt = {
        "writing_goal": _WORD_EDIT_GOAL_LABELS[normalized_goal],
        "goal_instruction": _word_edit_goal_instruction(normalized_goal),
        "writing_template": _WORD_EDIT_TEMPLATE_LABELS[normalized_template],
        "template_instruction": _word_edit_template_instruction(normalized_template),
        "task": str(instruction or "").strip()
        or "Improve this Word document for academic or business analysis while preserving factual meaning.",
        "rules": [
            "Return JSON only.",
            "Suggest paragraph-level replacements only; do not invent facts, data, or citations.",
            "Preserve the user's meaning and named entities.",
            "Prioritize clarity, structure, analytical strength, ESG/business terminology, and evidence-aware phrasing.",
            "Each suggestion must include category: clarity, logic, evidence, structure, tone, or esg_concept.",
            "Each suggestion must state the concrete problem it solves for the writer.",
            "Each suggestion must include severity: low, medium, or high.",
            "Use evidence_refs only when an evidence source directly supports the change.",
            "If a paragraph needs evidence but no source supports it, set evidence_needed true and explain the gap.",
            "Use evidence_gap_types to flag missing metric, source, comparison, causal_link, citation, or concept_definition.",
            f"Return at most {_DESKTOP_WORD_EDIT_MAX_SUGGESTIONS} suggestions.",
        ],
        "schema": {
            "suggestions": [
                {
                    "paragraph_id": "p_001",
                    "category": "clarity|logic|evidence|structure|tone|esg_concept",
                    "severity": "low|medium|high",
                    "problem": "The concrete weakness this edit fixes.",
                    "replacement": "Full replacement paragraph text.",
                    "reason": "Why this edit helps.",
                    "evidence_refs": ["E1"],
                    "evidence_needed": False,
                    "evidence_gap_types": ["metric", "source", "comparison", "causal_link", "citation", "concept_definition"],
                }
            ]
        },
        "evidence_sources": evidence_sources or [],
        "paragraphs": review_paragraphs,
    }
    messages = [
        {
            "role": "system",
            "content": (
                "You are a careful Word document editor for business, ESG, finance, and academic writing. "
                "You help users improve draft quality, but you must not fabricate evidence."
            ),
        },
        {"role": "user", "content": json.dumps(prompt, ensure_ascii=False)},
    ]

    def _create_completion(use_response_format: bool) -> Any:
        kwargs: Dict[str, Any] = {
            "model": _DESKTOP_WORD_EDIT_MODEL,
            "messages": messages,
            "temperature": 0.2,
            **chat_token_kwargs(_DESKTOP_WORD_EDIT_MODEL, _DESKTOP_WORD_EDIT_MAX_TOKENS),
        }
        if use_response_format:
            kwargs["response_format"] = {"type": "json_object"}
        return client.chat.completions.create(**kwargs)

    try:
        try:
            response = _create_completion(True)
        except Exception as exc:
            if "response_format" not in str(exc).lower() and "json" not in str(exc).lower():
                raise
            response = _create_completion(False)
        parsed = _parse_json_object_text(_extract_chat_completion_text(response))
        suggestions = _normalize_word_edit_suggestions(parsed, paragraph_lookup, evidence_sources=evidence_sources)
        return suggestions, _DESKTOP_WORD_EDIT_MODEL
    except Exception as exc:
        print(f"[desktop.word] suggestion generation fell back: {type(exc).__name__}: {exc}")
        return _fallback_word_edit_suggestions(paragraphs), "fallback_parse_error"


def _load_word_edit_session(current_user: Dict[str, Any], session_id: str) -> Dict[str, Any]:
    path = _word_session_json_path(current_user, session_id)
    if not path.is_file():
        raise HTTPException(status_code=404, detail="Word edit session was not found.")
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise HTTPException(status_code=500, detail="Word edit session is unreadable.") from exc
    if str(payload.get("user_id") or "") != str(current_user.get("id") or ""):
        raise HTTPException(status_code=403, detail="You do not have access to this Word edit session.")
    return payload


def _replace_paragraph_text(paragraph: Any, replacement: str) -> None:
    text = str(replacement or "")
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(text)


def _unique_output_name(file_name: str) -> str:
    stem = Path(file_name).stem or "document"
    return _safe_word_filename(f"{stem}.edited.docx")


def _normalize_entity_token(value: str) -> str:
    return str(value or "").strip(" .,!?:;()[]{}\"'’").lower()


def _filtered_entity_tokens(value: str) -> List[str]:
    normalized = _ENTITY_SPLIT_PATTERN.sub(" ", str(value or ""))
    tokens = [_normalize_entity_token(match.group(0)) for match in _ENTITY_TOKEN_PATTERN.finditer(normalized)]
    return [token for token in tokens if len(token) >= 2 and token not in _ENTITY_STOP_WORDS]


def _append_entity_term(terms: List[str], value: str) -> None:
    normalized = " ".join(_filtered_entity_tokens(value))
    if normalized and normalized not in terms:
        terms.append(normalized)


def _alias_terms_for_entity(tokens: List[str], joined: str) -> List[str]:
    keys = {joined, *tokens}
    if "american" in tokens and any(token in {"flight", "flights", "airline", "airlines"} for token in tokens):
        keys.add("american flight")
    aliases: List[str] = []
    for key in keys:
        aliases.extend(_ENTITY_ALIAS_MAP.get(key, ()))
    return aliases


def _terms_from_document_value(value: str) -> set[str]:
    tokens = _filtered_entity_tokens(value)
    terms = set(tokens)
    if 1 < len(tokens) <= 16:
        max_size = min(4, len(tokens))
        for size in range(2, max_size + 1):
            for index in range(0, len(tokens) - size + 1):
                terms.add(" ".join(tokens[index:index + size]))
    return terms


def _graph_entity_terms(graph_path: str) -> set[str]:
    if not graph_path:
        return set()
    path = Path(graph_path)
    if not path.exists() or not path.is_file():
        return set()

    try:
        mtime = path.stat().st_mtime
    except OSError:
        return set()

    cache_key = str(path.resolve())
    cached = _DOCUMENT_ENTITY_TERMS_CACHE.get(cache_key)
    if cached and cached[0] == mtime:
        return set(cached[1])

    terms: set[str] = set()
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        _DOCUMENT_ENTITY_TERMS_CACHE[cache_key] = (mtime, terms)
        return terms

    nodes = payload.get("nodes") if isinstance(payload, dict) else []
    if not isinstance(nodes, list):
        nodes = []
    for node in nodes[:500]:
        if not isinstance(node, dict):
            continue
        for key in ("id", "name", "label", "title"):
            terms.update(_terms_from_document_value(str(node.get(key) or "")))
        properties = node.get("properties") or node.get("metadata") or {}
        if isinstance(properties, dict):
            for key in ("name", "display_name", "label", "title"):
                terms.update(_terms_from_document_value(str(properties.get(key) or "")))

    _DOCUMENT_ENTITY_TERMS_CACHE[cache_key] = (mtime, terms)
    return set(terms)


def _extract_query_entity_terms(question: str) -> List[str]:
    text = str(question or "")
    terms: List[str] = []

    def add_phrase(phrase: str) -> None:
        tokens = _filtered_entity_tokens(phrase)
        if not tokens:
            return
        # Keep both phrase-level and token-level keys so "Apple Inc." can match
        # either "apple inc" or an uploaded file named "apple-report.pdf".
        joined = " ".join(tokens)
        candidates = [joined, *tokens, *_alias_terms_for_entity(tokens, joined)]
        for value in candidates:
            if value and value not in terms:
                _append_entity_term(terms, value)

    for pattern in (_ENTITY_OF_PATTERN, _ENTITY_POSSESSIVE_PATTERN, _ENTITY_TITLE_PHRASE_PATTERN):
        for match in pattern.finditer(text):
            add_phrase(match.group(1))

    lowered = text.lower()
    for alias in _ENTITY_ALIAS_MAP:
        if re.search(rf"\b{re.escape(alias)}\b", lowered):
            add_phrase(alias)

    for match in _ENTITY_ACRONYM_PATTERN.finditer(text):
        token = _normalize_entity_token(match.group(0))
        if token and token not in _ENTITY_STOP_WORDS and token not in terms:
            terms.append(token)

    return terms


def _document_entity_terms(entry: Dict[str, Any]) -> set[str]:
    values = [
        entry.get("document_id", ""),
        entry.get("title", ""),
        entry.get("source", ""),
    ]
    paths = entry.get("paths") or {}
    values.extend([paths.get("processed_text", ""), paths.get("chunks", ""), paths.get("graph", ""), paths.get("vector_store", "")])

    terms: set[str] = set()
    for value in values:
        terms.update(_terms_from_document_value(str(value or "")))
    terms.update(_graph_entity_terms(str(paths.get("graph") or "")))
    return terms


def _scope_document_ids_for_query(question: str, entries: List[Dict[str, Any]]) -> Tuple[List[str], List[str]]:
    query_terms = _extract_query_entity_terms(question)
    if not query_terms:
        return [], []

    phrase_terms = {term for term in query_terms if " " in term}
    alias_terms = {
        term for term in query_terms
        if " " not in term and any(term in aliases for aliases in _ENTITY_ALIAS_MAP.values())
    }
    phrase_matched_ids: List[str] = []
    alias_matched_ids: List[str] = []
    fallback_matched_ids: List[str] = []
    for entry in entries:
        document_id = str(entry.get("document_id") or "").strip()
        if not document_id:
            continue
        document_terms = _document_entity_terms(entry)
        if phrase_terms and any(term in document_terms for term in phrase_terms):
            phrase_matched_ids.append(document_id)
        elif alias_terms and any(term in document_terms for term in alias_terms):
            alias_matched_ids.append(document_id)
        elif any(term in document_terms for term in query_terms):
            fallback_matched_ids.append(document_id)

    return phrase_matched_ids or alias_matched_ids or fallback_matched_ids, query_terms


def _question_with_recent_user_context(question: str, history: Optional[List[Dict[str, Any]]]) -> str:
    recent_user_turns: List[str] = []
    for item in reversed(history or []):
        if str(item.get("role") or "").strip().lower() != "user":
            continue
        content = str(item.get("content") or "").strip()
        if not content:
            continue
        recent_user_turns.append(content)
        if len(recent_user_turns) >= 2:
            break
    parts = [*reversed(recent_user_turns), str(question or "").strip()]
    return "\n".join(part for part in parts if part)


def _no_accessible_documents_response() -> Dict[str, Any]:
    return {"error_response": JSONResponse(
        status_code=403,
        content={"answer": "", "sources": [], "error": "no_accessible_documents", "message": "No accessible documents for this account."},
    )}


def _resolve_rag_request_context(request: RagAskRequest, current_user: Optional[dict]) -> Dict[str, Any]:
    effective_document_ids = [str(item).strip() for item in (request.document_ids or []) if str(item).strip()]
    preferred_document_id = str(request.preferred_document_id or "").strip() or None
    entity_scope_miss = False
    entity_scope_terms: List[str] = []
    history = request.history
    memory_backend = "disabled"
    if request.session_id and current_user:
        try:
            history = chat_memory_service.get_recent_history(
                user_id=str(current_user["id"]),
                session_id=str(request.session_id),
            )
            memory_backend = "redis"
        except RedisUnavailableError:
            memory_backend = "disabled"
    scope_question = _question_with_recent_user_context(request.question, history)
    if current_user and not _is_admin_user(current_user):
        retrievable_entries = _retrievable_registry_entries(current_user)
        allowed_ids = {str(entry.get("document_id") or "").strip() for entry in retrievable_entries if str(entry.get("document_id") or "").strip()}
        broad_retrievable_scope = False
        if effective_document_ids or preferred_document_id:
            scoped_ids, scoped_terms = _scope_document_ids_for_query(scope_question, retrievable_entries)
            scoped_allowed_ids = sorted(set(scoped_ids) & allowed_ids)
            requested_ids = set(effective_document_ids)
            if preferred_document_id:
                requested_ids.add(preferred_document_id)
            if scoped_allowed_ids and requested_ids and requested_ids.isdisjoint(scoped_allowed_ids):
                effective_document_ids = scoped_allowed_ids
                preferred_document_id = None
                entity_scope_terms = scoped_terms
        if preferred_document_id:
            if preferred_document_id not in allowed_ids:
                return _no_accessible_documents_response()
            if not effective_document_ids:
                effective_document_ids = [preferred_document_id]
        if effective_document_ids:
            effective_document_ids = [doc_id for doc_id in effective_document_ids if doc_id in allowed_ids]
        else:
            scoped_ids, entity_scope_terms = _scope_document_ids_for_query(scope_question, retrievable_entries)
            if scoped_ids:
                effective_document_ids = sorted(set(scoped_ids) & allowed_ids)
            elif entity_scope_terms and not preferred_document_id:
                global_ids = {
                    str(entry.get("document_id") or "").strip()
                    for entry in retrievable_entries
                    if _is_global_entry(entry) and str(entry.get("document_id") or "").strip()
                }
                if global_ids:
                    # Global KB should remain searchable even when the target entity
                    # is not encoded in the document title/source metadata.
                    effective_document_ids = sorted(global_ids & allowed_ids)
                else:
                    entity_scope_miss = True
                    effective_document_ids = []
            else:
                # Broad user search is already safely constrained by owner/global
                # metadata filters in the vector store. Do not enumerate every
                # accessible document ID here; large $in filters can make Pinecone
                # slow or exceed request limits.
                broad_retrievable_scope = bool(allowed_ids)
                effective_document_ids = []
        if not effective_document_ids:
            if not entity_scope_miss and not broad_retrievable_scope:
                return _no_accessible_documents_response()
    elif current_user and _is_admin_user(current_user) and not effective_document_ids and not preferred_document_id:
        scoped_ids, entity_scope_terms = _scope_document_ids_for_query(scope_question, _retrievable_registry_entries(current_user))
        if scoped_ids:
            effective_document_ids = scoped_ids
        elif entity_scope_terms:
            entity_scope_miss = True
    elif not current_user:
        public_entries = [entry for entry in _collect_document_entries() if _can_retrieve_entry(None, entry)]
        public_ids = {str(entry.get("document_id") or "").strip() for entry in public_entries if str(entry.get("document_id") or "").strip()}
        if preferred_document_id:
            if preferred_document_id not in public_ids:
                return _no_accessible_documents_response()
            effective_document_ids = [preferred_document_id]
        elif effective_document_ids:
            effective_document_ids = [doc_id for doc_id in effective_document_ids if doc_id in public_ids]
            if not effective_document_ids:
                return _no_accessible_documents_response()
        else:
            scoped_ids, entity_scope_terms = _scope_document_ids_for_query(scope_question, public_entries)
            if scoped_ids:
                effective_document_ids = scoped_ids
            elif entity_scope_terms:
                entity_scope_miss = True
            else:
                # Anonymous requests do not carry owner_user_id, so public/global
                # access must be enforced by explicit document IDs.
                effective_document_ids = sorted(public_ids)
        if not effective_document_ids and not entity_scope_miss:
            return _no_accessible_documents_response()

    filters = {
        "document_ids": effective_document_ids,
        "preferred_document_id": preferred_document_id,
        "document_group": request.document_group,
        "source_type": request.source_type,
        "domain": request.domain,
    }
    if current_user and not _is_admin_user(current_user) and not effective_document_ids:
        filters["owner_user_id"] = str(current_user.get("id") or "")
    if entity_scope_miss:
        filters["entity_scope_miss"] = True
        filters["entity_scope_terms"] = entity_scope_terms

    return {
        "filters": filters,
        "history": history,
        "memory_backend": memory_backend,
        "user_id": str(current_user["id"]) if current_user else None,
        "error_response": None,
    }


def _load_request_history_for_rag(request: RagAskRequest, current_user: Optional[dict]) -> Tuple[List[Dict[str, Any]], str]:
    history = list(request.history or [])
    memory_backend = "disabled"
    if request.session_id and current_user:
        try:
            history = chat_memory_service.get_recent_history(
                user_id=str(current_user["id"]),
                session_id=str(request.session_id),
            )
            memory_backend = "redis"
        except RedisUnavailableError:
            memory_backend = "disabled"
    return history, memory_backend


def _resolve_general_rag_request_context(
    request: RagAskRequest,
    current_user: Optional[dict],
    history: List[Dict[str, Any]],
    memory_backend: str,
) -> Dict[str, Any]:
    filters: Dict[str, Any] = {
        "document_ids": [],
        "preferred_document_id": None,
        "document_group": request.document_group,
        "source_type": request.source_type,
        "domain": request.domain,
        "answer_mode": "general",
    }
    if current_user and not _is_admin_user(current_user):
        filters["owner_user_id"] = str(current_user.get("id") or "")
    return {
        "filters": filters,
        "history": history,
        "memory_backend": memory_backend,
        "user_id": str(current_user["id"]) if current_user else None,
        "error_response": None,
    }


async def _load_long_term_memory_context(
    db: aiosqlite.Connection,
    current_user: Optional[dict],
    query: str,
) -> Tuple[str, List[Dict[str, Any]]]:
    if not current_user:
        return "", []
    user_id = str(current_user.get("id") or "").strip()
    if not user_id:
        return "", []
    try:
        memories = await get_relevant_user_memories(db, user_id, query)
    except Exception as exc:
        print(f"[memory] retrieval skipped: {type(exc).__name__}: {exc}")
        return "", []
    return format_memories_for_prompt(memories), memories


def _history_with_long_term_memory(history: Optional[List[Dict[str, Any]]], memory_context: str) -> List[Dict[str, Any]]:
    normalized = list(history or [])
    if not memory_context:
        return normalized
    return [{"role": "assistant", "content": memory_context}, *normalized]


def _remember_exchange_later(
    *,
    user_id: Optional[str],
    user_message: str,
    assistant_message: str,
    source: str,
) -> None:
    normalized_user_id = str(user_id or "").strip()
    if not normalized_user_id or not str(user_message or "").strip() or not str(assistant_message or "").strip():
        return

    def _worker() -> None:
        async def _run() -> None:
            async with aiosqlite.connect(_DB_PATH) as memory_db:
                await init_user_memory_db(memory_db)
                await remember_exchange(
                    memory_db,
                    user_id=normalized_user_id,
                    user_message=user_message,
                    assistant_message=assistant_message,
                    source=source,
                )

        try:
            asyncio.run(_run())
        except Exception as exc:
            print(f"[memory] background store skipped: {type(exc).__name__}: {exc}")

    threading.Thread(target=_worker, daemon=True).start()


class PipelinePdfRequest(BaseModel):
    pdf_path: str
    name: str


class ManualDocumentRequest(BaseModel):
    title: str
    content: str
    domain: str = "general"
    source: str = ""
    source_type: str = ""


class RebuildDocumentGraphRequest(BaseModel):
    id: str
    title: str
    domain: str = "general"
    source: str = ""
    document_group: str = "user_upload"
    source_type: str = ""
    processed_text_path: str = ""
    chunks_path: str = ""
    extractions_path: str = ""
    graph_path: str = ""
    vector_store_path: str = ""


class ChatSessionCreateRequest(BaseModel):
    title: str = ""
    selected_document_id: str = ""
    mode: str = "ask"


class ChatSessionUpdateRequest(BaseModel):
    title: Optional[str] = None
    selected_document_id: Optional[str] = None
    mode: Optional[str] = None


class ChatSessionMessageRequest(BaseModel):
    role: str
    content: str
    timestamp: Optional[str] = None
    data: Optional[Dict[str, Any]] = None


class MemorySettingsUpdateRequest(BaseModel):
    enabled: Optional[bool] = None
    auto_extract: Optional[bool] = None
    raw_retention_days: Optional[int] = Field(default=None, ge=1, le=365)


class MemoryUpdateRequest(BaseModel):
    category: Optional[str] = None
    content: Optional[str] = Field(default=None, min_length=1, max_length=420)
    sensitivity: Optional[Literal["normal", "sensitive"]] = None
    confidence: Optional[float] = Field(default=None, ge=0.0, le=1.0)


_FEEDBACK_REASON_TAGS = {"missing_evidence", "wrong_citation", "hallucination", "irrelevant", "other"}


class FeedbackRequest(BaseModel):
    session_id: str
    message_id: str
    query: str
    answer: str
    rating: Literal["up", "down"]
    reason_tags: List[str] = Field(default_factory=list)
    reason_text: Optional[str] = ""
    sources: List[Dict[str, Any]] = Field(default_factory=list)
    timings_ms: Optional[Dict[str, Any]] = None


@app.post("/feedback")
async def submit_answer_feedback(
    request: FeedbackRequest,
    current_user: dict = Depends(get_current_user),
):
    session_id = str(request.session_id or "").strip()
    message_id = str(request.message_id or "").strip()
    query = str(request.query or "").strip()
    answer = str(request.answer or "").strip()
    reason_text = str(request.reason_text or "").strip()
    reason_tags = [
        str(tag).strip().lower()
        for tag in (request.reason_tags or [])
        if str(tag).strip().lower() in _FEEDBACK_REASON_TAGS
    ]
    reason_tags = list(dict.fromkeys(reason_tags))

    if not session_id or not message_id or not query or not answer:
        raise HTTPException(status_code=422, detail="session_id, message_id, query, and answer are required")
    if request.rating == "down" and not reason_tags and not reason_text:
        raise HTTPException(status_code=422, detail="Downvote feedback requires at least one reason tag or free-text reason")

    try:
        async with aiosqlite.connect(_FEEDBACK_DB_PATH) as db:
            cursor = await db.execute(
                """
                INSERT INTO answer_feedback (
                    user_id, session_id, message_id, query, answer, rating,
                    reason_tags, reason_text, sources_json, timings_json
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    str(current_user["id"]),
                    session_id,
                    message_id,
                    query,
                    answer,
                    request.rating,
                    json.dumps(reason_tags, ensure_ascii=False),
                    reason_text,
                    json.dumps(request.sources or [], ensure_ascii=False),
                    json.dumps(request.timings_ms or {}, ensure_ascii=False),
                ),
            )
            await db.commit()
            return JSONResponse(content={"ok": True, "id": cursor.lastrowid})
    except aiosqlite.IntegrityError:
        return JSONResponse(
            status_code=409,
            content={"ok": False, "error": "feedback_duplicate", "message": "Feedback already submitted for this answer."},
        )


@app.get("/admin/feedback/recent")
async def admin_recent_answer_feedback(
    rating: Optional[Literal["up", "down"]] = None,
    limit: int = Query(50, ge=1, le=200),
    current_user: dict = Depends(require_admin),
):
    _ = current_user
    clauses: List[str] = []
    params: List[Any] = []
    if rating:
        clauses.append("rating = ?")
        params.append(rating)
    where_sql = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    params.append(limit)

    async with aiosqlite.connect(_FEEDBACK_DB_PATH) as db:
        cursor = await db.execute(
            f"""
            SELECT id, user_id, session_id, message_id, query, answer, rating,
                   reason_tags, reason_text, sources_json, timings_json, created_at
            FROM answer_feedback
            {where_sql}
            ORDER BY datetime(created_at) DESC, id DESC
            LIMIT ?
            """,
            tuple(params),
        )
        rows = await cursor.fetchall()

    feedback = []
    for row in rows:
        try:
            reason_tags = json.loads(row[7] or "[]")
        except json.JSONDecodeError:
            reason_tags = []
        try:
            sources = json.loads(row[9] or "[]")
        except json.JSONDecodeError:
            sources = []
        try:
            timings_ms = json.loads(row[10] or "{}")
        except json.JSONDecodeError:
            timings_ms = {}
        feedback.append({
            "id": row[0],
            "user_id": row[1],
            "session_id": row[2],
            "message_id": row[3],
            "query": row[4],
            "answer": row[5],
            "rating": row[6],
            "reason_tags": reason_tags,
            "reason_text": row[8],
            "sources": sources,
            "timings_ms": timings_ms,
            "created_at": row[11],
        })

    return JSONResponse(content={"feedback": feedback})


@app.get("/chat/sessions")
async def list_chat_sessions(current_user: dict = Depends(get_current_user)):
    try:
        sessions = chat_memory_service.list_sessions(user_id=str(current_user["id"]))
        return JSONResponse(content={"sessions": sessions, "memory_backend": "redis"})
    except RedisUnavailableError as exc:
        return JSONResponse(
            content={"sessions": [], "warning": str(exc), "memory_backend": "disabled"},
        )


@app.post("/chat/sessions")
async def create_chat_session(
    request: ChatSessionCreateRequest,
    current_user: dict = Depends(get_current_user),
):
    try:
        session = chat_memory_service.create_session(
            user_id=str(current_user["id"]),
            title=request.title,
            selected_document_id=request.selected_document_id,
            mode=request.mode,
        )
        return JSONResponse(content={"session": session, "memory_backend": "redis"})
    except RedisUnavailableError as exc:
        return JSONResponse(
            status_code=503,
            content={"error": "chat_memory_unavailable", "message": str(exc), "memory_backend": "disabled"},
        )


@app.get("/chat/sessions/{session_id}")
async def get_chat_session(session_id: str, current_user: dict = Depends(get_current_user)):
    try:
        payload = chat_memory_service.get_session(user_id=str(current_user["id"]), session_id=session_id, include_messages=True)
        return JSONResponse(content={**payload, "memory_backend": "redis"})
    except RedisUnavailableError as exc:
        status = 404 if "not found" in str(exc).lower() else 503
        return JSONResponse(
            status_code=status,
            content={"error": "chat_session_unavailable", "message": str(exc), "memory_backend": "disabled"},
        )


@app.patch("/chat/sessions/{session_id}")
async def update_chat_session(
    session_id: str,
    request: ChatSessionUpdateRequest,
    current_user: dict = Depends(get_current_user),
):
    try:
        session = chat_memory_service.update_session(
            user_id=str(current_user["id"]),
            session_id=session_id,
            title=request.title,
            selected_document_id=request.selected_document_id,
            mode=request.mode,
        )
        return JSONResponse(content={"session": session, "memory_backend": "redis"})
    except RedisUnavailableError as exc:
        status = 404 if "not found" in str(exc).lower() else 503
        return JSONResponse(
            status_code=status,
            content={"error": "chat_session_update_failed", "message": str(exc), "memory_backend": "disabled"},
        )


@app.post("/chat/sessions/{session_id}/messages")
async def append_chat_message(
    session_id: str,
    request: ChatSessionMessageRequest,
    current_user: dict = Depends(get_current_user),
):
    try:
        session = chat_memory_service.append_message(
            user_id=str(current_user["id"]),
            session_id=session_id,
            message=request.model_dump(),
        )
        return JSONResponse(content={"session": session, "memory_backend": "redis"})
    except RedisUnavailableError as exc:
        status = 404 if "not found" in str(exc).lower() else 503
        return JSONResponse(
            status_code=status,
            content={"error": "chat_message_append_failed", "message": str(exc), "memory_backend": "disabled"},
        )


@app.delete("/chat/sessions/{session_id}")
async def delete_chat_session(session_id: str, current_user: dict = Depends(get_current_user)):
    try:
        chat_memory_service.delete_session(user_id=str(current_user["id"]), session_id=session_id)
        return JSONResponse(content={"deleted": True, "memory_backend": "redis"})
    except RedisUnavailableError as exc:
        status = 404 if "not found" in str(exc).lower() else 503
        return JSONResponse(
            status_code=status,
            content={"error": "chat_session_delete_failed", "message": str(exc), "memory_backend": "disabled"},
        )


@app.get("/memory/settings")
async def get_long_term_memory_settings(
    current_user: dict = Depends(get_current_user),
    db: aiosqlite.Connection = Depends(_get_db),
):
    settings = await get_memory_settings(db, str(current_user["id"]))
    return JSONResponse(content={"settings": settings, "memory_backend": "sqlite+vector"})


@app.patch("/memory/settings")
async def update_long_term_memory_settings(
    request: MemorySettingsUpdateRequest,
    current_user: dict = Depends(get_current_user),
    db: aiosqlite.Connection = Depends(_get_db),
):
    settings = await update_memory_settings(
        db,
        str(current_user["id"]),
        enabled=request.enabled,
        auto_extract=request.auto_extract,
        raw_retention_days=request.raw_retention_days,
    )
    return JSONResponse(content={"settings": settings, "memory_backend": "sqlite+vector"})


@app.get("/memory")
async def get_long_term_memories(
    category: Optional[str] = None,
    include_deleted: bool = False,
    limit: int = Query(80, ge=1, le=300),
    current_user: dict = Depends(get_current_user),
    db: aiosqlite.Connection = Depends(_get_db),
):
    memories = await list_user_memories(
        db,
        str(current_user["id"]),
        category=category,
        include_deleted=include_deleted,
        limit=limit,
    )
    settings = await get_memory_settings(db, str(current_user["id"]))
    return JSONResponse(content={"memories": memories, "settings": settings, "memory_backend": "sqlite+vector"})


@app.patch("/memory/{memory_id}")
async def patch_long_term_memory(
    memory_id: str,
    request: MemoryUpdateRequest,
    current_user: dict = Depends(get_current_user),
    db: aiosqlite.Connection = Depends(_get_db),
):
    memory = await update_user_memory(
        db,
        str(current_user["id"]),
        memory_id,
        category=request.category,
        content=request.content,
        sensitivity=request.sensitivity,
        confidence=request.confidence,
    )
    if memory is None:
        return JSONResponse(
            status_code=404,
            content={"error": "memory_not_found", "message": "No active memory found for this account."},
        )
    return JSONResponse(content={"memory": memory, "memory_backend": "sqlite+vector"})


@app.delete("/memory/{memory_id}")
async def remove_long_term_memory(
    memory_id: str,
    current_user: dict = Depends(get_current_user),
    db: aiosqlite.Connection = Depends(_get_db),
):
    deleted = await delete_user_memory(db, str(current_user["id"]), memory_id)
    if not deleted:
        return JSONResponse(
            status_code=404,
            content={"error": "memory_not_found", "message": "No active memory found for this account."},
        )
    return JSONResponse(content={"deleted": True, "memory_backend": "sqlite+vector"})


def _sort_registry_entries(entries: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return sorted(
        entries,
        key=lambda entry: str(entry.get("ingested_at") or ""),
        reverse=True,
    )


def _entry_from_upload(upload: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    document_id = str(upload.get("document_id") or "").strip()
    if not document_id:
        return None
    paths = upload.get("paths") or {}
    graph_path = str(paths.get("graph_path") or "").strip()
    chunks_path = str(paths.get("chunks_path") or "").strip()
    vector_store_path = str(paths.get("vector_store_path") or "").strip()
    if not graph_path or not chunks_path or not vector_store_path:
        return None
    owner_user_id = str((upload.get("uploader") or {}).get("id") or "")
    fallback_group = upload.get("document_group") or "global_kb"
    fallback_visibility = "private" if fallback_group == "user_private" else "global"
    return {
        "document_id": document_id,
        "title": upload.get("title", ""),
        "domain": upload.get("domain", "") or "general",
        "source": upload.get("source", "") or upload.get("filename", ""),
        "source_type": upload.get("source_type", ""),
        "document_group": fallback_group,
        "owner_user_id": owner_user_id,
        "visibility_scope": fallback_visibility,
        "ingested_at": upload.get("created_at", ""),
        "paths": {
            "processed_text": str(paths.get("processed_text_path") or ""),
            "chunks": chunks_path,
            "extractions": str(paths.get("extractions_path") or ""),
            "graph": graph_path,
            "vector_store": vector_store_path,
        },
        "neo4j_sync": {},
    }


def _entry_from_chunk_file(chunks_path: Path) -> Optional[Dict[str, Any]]:
    if chunks_path.name.startswith("."):
        return None
    try:
        with chunks_path.open("r", encoding="utf-8") as handle:
            first_line = handle.readline().strip()
    except (OSError, UnicodeDecodeError):
        return None
    if not first_line:
        return None

    try:
        first_chunk = json.loads(first_line)
    except json.JSONDecodeError:
        return None
    if not isinstance(first_chunk, dict):
        return None

    document_id = str(first_chunk.get("document_id") or chunks_path.stem.replace("_chunks", "")).strip()
    if not document_id:
        return None

    graph_path = GRAPH_DIR / f"{document_id}_graph.json"
    vector_store_path = VECTOR_DIR / document_id
    document_group = str(first_chunk.get("document_group") or "user_upload").strip() or "user_upload"
    owner_user_id = str(first_chunk.get("owner_user_id") or first_chunk.get("uploaded_by") or "").strip()
    visibility_scope = str(first_chunk.get("visibility_scope") or "").strip()
    if not visibility_scope and document_group.lower() in {"global", "global_kb", "shared"}:
        visibility_scope = "global"
    elif not visibility_scope and document_group.lower() in {"private", "user_private"}:
        visibility_scope = "private"

    return {
        "document_id": document_id,
        "title": first_chunk.get("document_title") or document_id,
        "domain": first_chunk.get("domain") or "general",
        "source": first_chunk.get("source") or chunks_path.name,
        "source_type": first_chunk.get("source_type") or "",
        "document_group": document_group,
        "owner_user_id": owner_user_id,
        "visibility_scope": visibility_scope,
        "ingested_at": first_chunk.get("ingested_at") or "",
        "paths": {
            "processed_text": "",
            "chunks": str(chunks_path),
            "extractions": "",
            "graph": str(graph_path) if graph_path.exists() else "",
            "vector_store": str(vector_store_path) if vector_store_path.exists() else "",
        },
        "neo4j_sync": {},
    }


def _collect_orphan_chunk_entries(existing_document_ids: set[str]) -> List[Dict[str, Any]]:
    entries: List[Dict[str, Any]] = []
    if not CHUNK_DIR.exists():
        return entries

    for chunks_path in sorted(CHUNK_DIR.glob("*_chunks.jsonl")):
        entry = _entry_from_chunk_file(chunks_path)
        if not entry:
            continue
        document_id = str(entry.get("document_id") or "").strip()
        if document_id and document_id not in existing_document_ids:
            entries.append(entry)
    return entries


def _collect_document_entries() -> List[Dict[str, Any]]:
    audit_index = list_latest_uploads_by_document_id()
    merged: Dict[str, Dict[str, Any]] = {}

    for entry in document_registry.list_entries(valid_only=True):
        document_id = str(entry.get("document_id") or "").strip()
        if document_id:
            merged[document_id] = entry

    for document_id, upload in audit_index.items():
        if str(upload.get("status") or "") in {"deleted", "deleted_with_warnings"}:
            continue
        if document_id in merged:
            continue
        fallback = _entry_from_upload(upload)
        if fallback is not None:
            merged[document_id] = fallback

    for entry in _collect_orphan_chunk_entries(set(merged.keys())):
        document_id = str(entry.get("document_id") or "").strip()
        if document_id:
            merged[document_id] = entry

    return _sort_registry_entries(list(merged.values()))


def _is_admin_user(user: Optional[Dict[str, Any]]) -> bool:
    return _normalize_role((user or {}).get("role", "")) == "admin"


def _is_global_entry(entry: Dict[str, Any]) -> bool:
    group = str(entry.get("document_group") or "").strip().lower()
    visibility = str(entry.get("visibility_scope") or "").strip().lower()
    return group in {"global_kb", "global", "shared"} or visibility == "global"


def _is_legacy_unowned_user_upload(entry: Dict[str, Any]) -> bool:
    group = str(entry.get("document_group") or "").strip().lower()
    owner_user_id = str(entry.get("owner_user_id") or "").strip()
    visibility = str(entry.get("visibility_scope") or "").strip().lower()
    return group == "user_upload" and not owner_user_id and not visibility


def _can_access_entry(user: Dict[str, Any], entry: Dict[str, Any]) -> bool:
    if _is_admin_user(user):
        return True
    if _is_legacy_unowned_user_upload(entry):
        return bool(str(user.get("id") or "").strip())
    group = str(entry.get("document_group") or "").strip().lower()
    owner_user_id = str(entry.get("owner_user_id") or "").strip()
    current_user_id = str(user.get("id") or "").strip()
    if group in {"user_private", "private"}:
        return bool(current_user_id) and current_user_id == owner_user_id
    return bool(current_user_id) and owner_user_id == current_user_id


def _can_retrieve_entry(user: Optional[Dict[str, Any]], entry: Dict[str, Any]) -> bool:
    if _is_admin_user(user):
        return True
    if _is_global_entry(entry):
        return True
    if not user:
        return False
    if _is_legacy_unowned_user_upload(entry):
        return bool(str(user.get("id") or "").strip())
    owner_user_id = str(entry.get("owner_user_id") or "").strip()
    current_user_id = str(user.get("id") or "").strip()
    return bool(current_user_id) and owner_user_id == current_user_id


def _accessible_registry_entries(user: Dict[str, Any], include_invalid: bool = False) -> List[Dict[str, Any]]:
    entries = document_registry.list_entries(valid_only=not include_invalid)
    return [entry for entry in entries if _can_access_entry(user, entry)]


def _retrievable_registry_entries(user: Optional[Dict[str, Any]], include_invalid: bool = False) -> List[Dict[str, Any]]:
    entries = document_registry.list_entries(valid_only=False) if include_invalid else _collect_document_entries()
    return [entry for entry in entries if _can_retrieve_entry(user, entry)]


def _accessible_document_ids(user: Dict[str, Any]) -> List[str]:
    output = []
    for entry in _accessible_registry_entries(user):
        document_id = str(entry.get("document_id") or "").strip()
        if document_id:
            output.append(document_id)
    return output


class Neo4jQuestionRequest(BaseModel):
    question: str
    hops: int = 2
    limit: int = 10


_PUBLIC_GRAPH_MAX_NODES = 30000
_PUBLIC_GRAPH_MAX_EDGES = 40000


def _bounded_graph_limit(value: int, default: int, maximum: int, minimum: int = 1) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        parsed = default
    return max(minimum, min(parsed, maximum))


def _resolve_project_file(value: str) -> Path:
    path = Path(str(value or "")).expanduser()
    if path.is_absolute():
        return path
    return (_APP_ROOT / path).resolve()


def _safe_graph_float(value: object, fallback: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return fallback


def _graph_text_domain(*values: object) -> str:
    text = " ".join(str(value or "") for value in values).lower()
    direct = text.strip()
    if direct in {"environmental", "environment", "e"}:
        return "environmental"
    if direct in {"social", "s"}:
        return "social"
    if direct in {"governance", "g"}:
        return "governance"
    if direct in {"ai", "artificial intelligence"}:
        return "ai"
    if re.search(r"\b(ai|llm|machine learning|model|algorithm|retrieval|rag|embedding|vector|parser|parsing|summary|summarization|prediction)\b", text):
        return "ai"
    if re.search(r"\b(climate|carbon|emission|scope\s?[123]|ghg|greenhouse|net zero|renewable|energy|water|waste|recycl|circular|biodiversity|pollution)\b", text):
        return "environmental"
    if re.search(r"\b(employee|workforce|worker|diversity|inclusion|dei|safety|health|supplier|supply chain|community|human rights|labor|labour|training|customer)\b", text):
        return "social"
    if re.search(r"\b(governance|board|committee|audit|assurance|compliance|ethic|risk|control|policy|shareholder|remuneration|compensation|anti[- ]?bribery|corruption)\b", text):
        return "governance"
    return "general"


def _local_graph_node_payload(node: Dict[str, Any], entry: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    node_id = str(node.get("id") or "").strip()
    if not node_id:
        return None
    properties = node.get("properties") if isinstance(node.get("properties"), dict) else {}
    label = str(properties.get("display_name") or properties.get("name") or node_id).strip() or node_id
    description = str(properties.get("description") or "")
    direct_domain = str(properties.get("esg_domain") or properties.get("domain") or entry.get("domain") or "")
    domain = _graph_text_domain(direct_domain, node.get("type"), label, description)
    return {
        "id": node_id,
        "label": label,
        "domain": domain,
        "type": str(node.get("type") or properties.get("type") or "Entity"),
        "confidence": _safe_graph_float(properties.get("confidence"), 0.8),
        "description": description,
        "company": str(properties.get("company") or ""),
        "year": str(properties.get("year") or ""),
        "normalizedName": str(properties.get("normalized_name") or node_id),
        "metadata": {
            "documentIds": [str(entry.get("document_id") or "")],
            "documentTitles": [str(entry.get("title") or "")],
            "frequency": 1,
        },
    }


def _merge_local_node(existing: Dict[str, Any], candidate: Dict[str, Any]) -> None:
    existing["confidence"] = max(float(existing.get("confidence") or 0), float(candidate.get("confidence") or 0))
    if not existing.get("description") and candidate.get("description"):
        existing["description"] = candidate["description"]
    if existing.get("domain") == "general" and candidate.get("domain") != "general":
        existing["domain"] = candidate["domain"]
    metadata = existing.setdefault("metadata", {})
    candidate_metadata = candidate.get("metadata") or {}
    for key in ("documentIds", "documentTitles"):
        values = list(metadata.get(key) or [])
        for value in candidate_metadata.get(key) or []:
            if value and value not in values:
                values.append(value)
        metadata[key] = values[:12]
    metadata["frequency"] = int(metadata.get("frequency") or 1) + 1


def _local_graph_edge_payload(edge: Dict[str, Any], entry: Dict[str, Any], node_domains: Dict[str, str]) -> Optional[Dict[str, Any]]:
    source = str(edge.get("source") or "").strip()
    target = str(edge.get("target") or "").strip()
    if not source or not target:
        return None
    properties = edge.get("properties") if isinstance(edge.get("properties"), dict) else {}
    relation = str(edge.get("relation") or edge.get("relationship_type") or properties.get("relation_type") or "RELATED_TO")
    evidence = str(properties.get("evidence") or properties.get("context") or "")
    direct_domain = str(properties.get("domain") or properties.get("esg_domain") or entry.get("domain") or "")
    inferred = _graph_text_domain(direct_domain, relation, evidence, node_domains.get(source), node_domains.get(target))
    return {
        "source": source,
        "target": target,
        "relationship_type": relation,
        "confidence": _safe_graph_float(properties.get("confidence"), 0.75),
        "evidence": evidence,
        "domain": inferred,
        "relationship_action": str(properties.get("action") or ""),
        "relationship_nature": str(properties.get("nature") or ""),
        "documentId": str(entry.get("document_id") or ""),
        "chunkId": str(properties.get("chunk_id") or ""),
    }


def _merge_local_knowledge_graph(entries: List[Dict[str, Any]], node_limit: int, edge_limit: int) -> Dict[str, Any]:
    nodes_by_id: Dict[str, Dict[str, Any]] = {}
    raw_edges: List[Dict[str, Any]] = []
    degree: Dict[str, int] = {}
    loaded_document_ids: List[str] = []

    for entry in entries:
        paths = entry.get("paths") or {}
        graph_path = _resolve_project_file(str(paths.get("graph") or ""))
        if not graph_path.is_file():
            continue
        try:
            graph = json.loads(graph_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError, UnicodeDecodeError):
            continue
        document_id = str(entry.get("document_id") or "")
        if document_id:
            loaded_document_ids.append(document_id)
        for node in graph.get("nodes") or []:
            if not isinstance(node, dict):
                continue
            payload = _local_graph_node_payload(node, entry)
            if not payload:
                continue
            existing = nodes_by_id.get(payload["id"])
            if existing:
                _merge_local_node(existing, payload)
            else:
                nodes_by_id[payload["id"]] = payload
        for edge in graph.get("edges") or []:
            if not isinstance(edge, dict):
                continue
            source = str(edge.get("source") or "").strip()
            target = str(edge.get("target") or "").strip()
            if not source or not target:
                continue
            degree[source] = degree.get(source, 0) + 1
            degree[target] = degree.get(target, 0) + 1
            raw_edges.append({"edge": edge, "entry": entry})

    sorted_nodes = sorted(
        nodes_by_id.values(),
        key=lambda node: (
            degree.get(str(node.get("id") or ""), 0),
            int((node.get("metadata") or {}).get("frequency") or 0),
            float(node.get("confidence") or 0),
            str(node.get("label") or ""),
        ),
        reverse=True,
    )
    selected_nodes = sorted_nodes[:node_limit]
    selected_ids = {str(node.get("id") or "") for node in selected_nodes}
    node_domains = {str(node.get("id") or ""): str(node.get("domain") or "general") for node in selected_nodes}

    edges: List[Dict[str, Any]] = []
    seen_edges: set[str] = set()
    for item in raw_edges:
        if len(edges) >= edge_limit:
            break
        edge = _local_graph_edge_payload(item["edge"], item["entry"], node_domains)
        if not edge or edge["source"] not in selected_ids or edge["target"] not in selected_ids:
            continue
        edge_key = f"{edge['source']}|{edge['relationship_type']}|{edge['target']}|{edge.get('documentId', '')}"
        if edge_key in seen_edges:
            continue
        seen_edges.add(edge_key)
        edges.append(edge)

    return {
        "nodes": selected_nodes,
        "edges": edges,
        "metadata": {
            "node_count": len(selected_nodes),
            "edge_count": len(edges),
            "total_node_count": len(nodes_by_id),
            "total_edge_count": len(raw_edges),
            "document_count": len(set(loaded_document_ids)),
            "source": "local_graph_json",
            "is_directed": True,
            "is_acyclic": False,
        },
    }


def _neo4j_visualization_to_graph_data(payload: Dict[str, Any], edge_limit: int) -> Dict[str, Any]:
    nodes = []
    for row in payload.get("nodes") or []:
        if not isinstance(row, dict):
            continue
        node_id = str(row.get("id") or "").strip()
        if not node_id:
            continue
        domain = _graph_text_domain(row.get("domain") or row.get("esg_domain"), row.get("type"), row.get("label"), row.get("description"))
        nodes.append(
            {
                "id": node_id,
                "label": str(row.get("label") or row.get("name") or node_id),
                "domain": domain,
                "type": str(row.get("type") or "Entity"),
                "confidence": _safe_graph_float(row.get("confidence"), 0.75),
                "description": str(row.get("description") or ""),
                "company": str(row.get("company") or ""),
                "year": str(row.get("year") or ""),
                "normalizedName": str(row.get("normalized_name") or node_id),
                "metadata": {"frequency": int(row.get("frequency") or 0)},
            }
        )
    node_ids = {node["id"] for node in nodes}
    edges = []
    for row in (payload.get("edges") or [])[:edge_limit]:
        if not isinstance(row, dict):
            continue
        source = str(row.get("source") or "").strip()
        target = str(row.get("target") or "").strip()
        if source not in node_ids or target not in node_ids:
            continue
        edges.append(
            {
                "source": source,
                "target": target,
                "relationship_type": str(row.get("relationship_type") or row.get("type") or "RELATED_TO"),
                "confidence": _safe_graph_float(row.get("confidence"), 0.75),
                "evidence": str(row.get("evidence") or ""),
                "domain": _graph_text_domain(row.get("domain"), row.get("category"), row.get("relationship_type"), row.get("evidence")),
                "relationship_action": str(row.get("relationship_action") or ""),
                "relationship_nature": str(row.get("relationship_nature") or ""),
                "documentId": str(row.get("document_id") or ""),
                "chunkId": str(row.get("chunk_id") or ""),
            }
        )
    return {
        "nodes": nodes,
        "edges": edges,
        "metadata": {
            "node_count": len(nodes),
            "edge_count": len(edges),
            "source": "neo4j",
            "is_directed": True,
            "is_acyclic": False,
        },
    }


def _load_real_knowledge_graph(current_user: Optional[Dict[str, Any]], node_limit: int, edge_limit: int) -> Dict[str, Any]:
    entries = _retrievable_registry_entries(current_user)
    document_ids = [
        str(entry.get("document_id") or "").strip()
        for entry in entries
        if str(entry.get("document_id") or "").strip()
    ]

    try:
        store = get_neo4j_store()
        if store is not None and (_is_admin_user(current_user) or document_ids):
            raw = store.get_visualization_graph(
                limit=node_limit,
                document_ids=None if _is_admin_user(current_user) else document_ids,
            )
            graph = _neo4j_visualization_to_graph_data(raw, edge_limit=edge_limit)
            if graph["nodes"]:
                return graph
    except Exception as exc:
        print(f"[public-knowledge-graph] Neo4j graph unavailable, using local graph files: {type(exc).__name__}: {exc}")

    return _merge_local_knowledge_graph(entries, node_limit=node_limit, edge_limit=edge_limit)


def _public_graph_cache_enabled() -> bool:
    return bool(getattr(chat_memory_service, "enabled", False))


def _public_graph_cache_scope(current_user: Optional[Dict[str, Any]]) -> str:
    if _is_admin_user(current_user):
        return "admin"
    user_id = str((current_user or {}).get("id") or "").strip()
    if user_id:
        return f"user:{user_id}"
    return "public"


def _public_graph_cache_key(current_user: Optional[Dict[str, Any]], node_limit: int, edge_limit: int) -> str:
    raw = "|".join(
        [
            _PUBLIC_GRAPH_CACHE_VERSION,
            _public_graph_cache_scope(current_user),
            str(node_limit),
            str(edge_limit),
        ]
    )
    digest = hashlib.sha256(raw.encode("utf-8")).hexdigest()[:32]
    return f"public:knowledge_graph:{_PUBLIC_GRAPH_CACHE_VERSION}:{digest}"


def _public_graph_cache_lock_key() -> str:
    return f"public:knowledge_graph:{_PUBLIC_GRAPH_CACHE_VERSION}:refresh_lock"


def _public_graph_cache_get(cache_key: str) -> Optional[Dict[str, Any]]:
    try:
        raw = chat_memory_service.client.execute("GET", cache_key)
    except RedisUnavailableError:
        return None
    if not raw:
        return None
    try:
        payload = json.loads(str(raw))
    except json.JSONDecodeError:
        return None
    if not isinstance(payload, dict) or not isinstance(payload.get("graph"), dict):
        return None
    return payload


def _public_graph_cache_store(cache_key: str, graph: Dict[str, Any]) -> Dict[str, Any]:
    now = time.time()
    payload = {
        "graph": graph,
        "created_at": now,
        "expires_at": now + _PUBLIC_GRAPH_CACHE_TTL_SECONDS,
        "cache_version": _PUBLIC_GRAPH_CACHE_VERSION,
    }
    raw = json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
    chat_memory_service.client.execute("SET", cache_key, raw, "EX", _PUBLIC_GRAPH_CACHE_STALE_SECONDS)
    return payload


def _public_graph_cache_is_fresh(payload: Dict[str, Any]) -> bool:
    try:
        return float(payload.get("expires_at") or 0) > time.time()
    except (TypeError, ValueError):
        return False


def _public_graph_cache_acquire_lock() -> str:
    token = str(uuid.uuid4())
    try:
        result = chat_memory_service.client.execute(
            "SET",
            _public_graph_cache_lock_key(),
            token,
            "NX",
            "EX",
            _PUBLIC_GRAPH_CACHE_LOCK_SECONDS,
        )
    except RedisUnavailableError:
        return ""
    return token if str(result or "").upper() == "OK" else ""


def _public_graph_cache_release_lock(token: str) -> None:
    if not token:
        return
    try:
        current = chat_memory_service.client.execute("GET", _public_graph_cache_lock_key())
        if current == token:
            chat_memory_service.client.execute("DEL", _public_graph_cache_lock_key())
    except RedisUnavailableError:
        return


def _load_and_store_public_graph_cache(
    cache_key: str,
    current_user: Optional[Dict[str, Any]],
    node_limit: int,
    edge_limit: int,
    lock_token: str,
) -> Dict[str, Any]:
    try:
        graph = _load_real_knowledge_graph(
            current_user=current_user,
            node_limit=node_limit,
            edge_limit=edge_limit,
        )
        _public_graph_cache_store(cache_key, graph)
        return graph
    finally:
        _public_graph_cache_release_lock(lock_token)


def _refresh_public_graph_cache_in_background(
    cache_key: str,
    current_user: Optional[Dict[str, Any]],
    node_limit: int,
    edge_limit: int,
    lock_token: str,
) -> None:
    try:
        _load_and_store_public_graph_cache(
            cache_key=cache_key,
            current_user=dict(current_user or {}),
            node_limit=node_limit,
            edge_limit=edge_limit,
            lock_token=lock_token,
        )
    except Exception as exc:
        print(f"[public-knowledge-graph-cache] refresh failed: {type(exc).__name__}: {exc}")


async def _wait_for_public_graph_cache(cache_key: str) -> Optional[Dict[str, Any]]:
    deadline = time.time() + _PUBLIC_GRAPH_CACHE_WAIT_SECONDS
    while time.time() < deadline:
        await asyncio.sleep(0.25)
        payload = _public_graph_cache_get(cache_key)
        if payload is not None:
            return payload
    return None


def _kg_view_cache_scope(request: Request, current_user: Optional[Dict[str, Any]]) -> str:
    effective_user = _kg_view_effective_user(request, current_user)
    document_id = _kg_view_request_document_id(request)
    if document_id and _kg_view_can_use_document_scope(request, effective_user, document_id):
        return f"document:{document_id}"
    return _public_graph_cache_scope(effective_user)


def _kg_view_cache_key(
    request: Request,
    current_user: Optional[Dict[str, Any]],
    name: str,
    params: Optional[Dict[str, Any]] = None,
) -> str:
    raw = json.dumps(
        {
            "version": _KG_VIEW_CACHE_VERSION,
            "scope": _kg_view_cache_scope(request, current_user),
            "name": name,
            "params": params or {},
        },
        ensure_ascii=False,
        sort_keys=True,
        default=str,
    )
    digest = hashlib.sha256(raw.encode("utf-8")).hexdigest()[:40]
    return f"kg-view:{_KG_VIEW_CACHE_VERSION}:{digest}"


def _kg_view_cache_lock_key(cache_key: str) -> str:
    return f"{cache_key}:lock"


def _kg_view_cache_is_fresh(payload: Dict[str, Any]) -> bool:
    try:
        return float(payload.get("expires_at") or 0) > time.time()
    except (TypeError, ValueError):
        return False


def _kg_view_cache_memory_set_payload(cache_key: str, payload: Dict[str, Any]) -> None:
    with _KG_VIEW_CACHE_MEMORY_GUARD:
        _KG_VIEW_CACHE_MEMORY[cache_key] = payload
        if len(_KG_VIEW_CACHE_MEMORY) <= _KG_VIEW_CACHE_MEMORY_MAX:
            return
        oldest_keys = sorted(
            _KG_VIEW_CACHE_MEMORY,
            key=lambda key: float(_KG_VIEW_CACHE_MEMORY.get(key, {}).get("created_at") or 0),
        )
        for old_key in oldest_keys[: max(1, len(_KG_VIEW_CACHE_MEMORY) - _KG_VIEW_CACHE_MEMORY_MAX)]:
            _KG_VIEW_CACHE_MEMORY.pop(old_key, None)


def _kg_view_cache_memory_get(cache_key: str) -> Optional[Dict[str, Any]]:
    now = time.time()
    with _KG_VIEW_CACHE_MEMORY_GUARD:
        payload = _KG_VIEW_CACHE_MEMORY.get(cache_key)
        if payload is None:
            return None
        try:
            stale_expires_at = float(payload.get("stale_expires_at") or 0)
        except (TypeError, ValueError):
            stale_expires_at = 0
        if stale_expires_at <= now:
            _KG_VIEW_CACHE_MEMORY.pop(cache_key, None)
            return None
        return payload


def _kg_view_redis_available_for_cache() -> bool:
    return chat_memory_service.enabled and time.time() >= _KG_VIEW_REDIS_DISABLED_UNTIL


def _kg_view_note_redis_failure() -> None:
    global _KG_VIEW_REDIS_DISABLED_UNTIL
    _KG_VIEW_REDIS_DISABLED_UNTIL = time.time() + _KG_VIEW_REDIS_RETRY_SECONDS


def _kg_view_cache_get(cache_key: str) -> Optional[Dict[str, Any]]:
    if not _KG_VIEW_CACHE_ENABLED:
        return None
    memory_payload = _kg_view_cache_memory_get(cache_key)
    if memory_payload is not None:
        return memory_payload
    if _kg_view_redis_available_for_cache():
        try:
            raw = chat_memory_service.client.execute("GET", cache_key)
        except RedisUnavailableError:
            _kg_view_note_redis_failure()
            raw = None
        if raw:
            try:
                payload = json.loads(str(raw))
            except json.JSONDecodeError:
                payload = None
            if isinstance(payload, dict) and "payload" in payload:
                _kg_view_cache_memory_set_payload(cache_key, payload)
                return payload
    return None


def _kg_view_cache_store(cache_key: str, value: Any) -> Dict[str, Any]:
    now = time.time()
    payload = {
        "payload": value,
        "created_at": now,
        "expires_at": now + _KG_VIEW_CACHE_TTL_SECONDS,
        "stale_expires_at": now + _KG_VIEW_CACHE_STALE_SECONDS,
        "cache_version": _KG_VIEW_CACHE_VERSION,
    }
    _kg_view_cache_memory_set_payload(cache_key, payload)
    if _kg_view_redis_available_for_cache():
        raw = json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
        try:
            chat_memory_service.client.execute("SET", cache_key, raw, "EX", _KG_VIEW_CACHE_STALE_SECONDS)
        except RedisUnavailableError:
            _kg_view_note_redis_failure()
            pass
    return payload


def _kg_view_cache_memory_lock(cache_key: str) -> threading.Lock:
    with _KG_VIEW_CACHE_MEMORY_GUARD:
        lock = _KG_VIEW_CACHE_MEMORY_LOCKS.get(cache_key)
        if lock is None:
            lock = threading.Lock()
            _KG_VIEW_CACHE_MEMORY_LOCKS[cache_key] = lock
        return lock


def _kg_view_cache_acquire_lock(cache_key: str) -> str:
    token = str(uuid.uuid4())
    if _kg_view_redis_available_for_cache():
        try:
            result = chat_memory_service.client.execute(
                "SET",
                _kg_view_cache_lock_key(cache_key),
                token,
                "NX",
                "EX",
                _KG_VIEW_CACHE_LOCK_SECONDS,
            )
            if str(result or "").upper() == "OK":
                return f"redis:{token}"
        except RedisUnavailableError:
            _kg_view_note_redis_failure()
            pass
    lock = _kg_view_cache_memory_lock(cache_key)
    if lock.acquire(blocking=False):
        return f"memory:{token}"
    return ""


def _kg_view_cache_release_lock(cache_key: str, token: str) -> None:
    if not token:
        return
    if token.startswith("redis:"):
        raw_token = token.split(":", 1)[1]
        try:
            current = chat_memory_service.client.execute("GET", _kg_view_cache_lock_key(cache_key))
            if current == raw_token:
                chat_memory_service.client.execute("DEL", _kg_view_cache_lock_key(cache_key))
        except RedisUnavailableError:
            _kg_view_note_redis_failure()
            pass
        return
    lock = _kg_view_cache_memory_lock(cache_key)
    try:
        lock.release()
    except RuntimeError:
        pass


def _kg_view_cache_build_and_store(cache_key: str, lock_token: str, builder: Callable[[], Any]) -> Any:
    try:
        value = builder()
        _kg_view_cache_store(cache_key, value)
        return value
    finally:
        _kg_view_cache_release_lock(cache_key, lock_token)


def _kg_view_cache_refresh_background(cache_key: str, lock_token: str, builder: Callable[[], Any]) -> None:
    try:
        _kg_view_cache_build_and_store(cache_key, lock_token, builder)
    except Exception as exc:
        print(f"[kg-view-cache] refresh failed: {type(exc).__name__}: {exc}")


async def _kg_view_cache_wait(cache_key: str) -> Optional[Dict[str, Any]]:
    deadline = time.time() + _KG_VIEW_CACHE_WAIT_SECONDS
    while time.time() < deadline:
        await asyncio.sleep(0.25)
        payload = _kg_view_cache_get(cache_key)
        if payload is not None:
            return payload
    return None


async def _kg_view_cached_json(
    request: Request,
    current_user: Optional[Dict[str, Any]],
    name: str,
    params: Dict[str, Any],
    builder: Callable[[], Dict[str, Any]],
) -> JSONResponse:
    if not _KG_VIEW_CACHE_ENABLED:
        return JSONResponse(content=builder(), headers={"X-KG-View-Cache": "bypass"})

    cache_key = _kg_view_cache_key(request, current_user, name, params)
    cached = _kg_view_cache_get(cache_key)
    if cached is not None:
        payload = cached.get("payload")
        if isinstance(payload, dict):
            if _kg_view_cache_is_fresh(cached):
                return JSONResponse(content=payload, headers={"X-KG-View-Cache": "hit"})
            lock_token = _kg_view_cache_acquire_lock(cache_key)
            if lock_token:
                _GRAPH_CACHE_REFRESH_EXECUTOR.submit(
                    _kg_view_cache_refresh_background,
                    cache_key,
                    lock_token,
                    builder,
                )
            return JSONResponse(content=payload, headers={"X-KG-View-Cache": "stale"})

    lock_token = _kg_view_cache_acquire_lock(cache_key)
    if lock_token:
        loop = asyncio.get_running_loop()
        payload = await loop.run_in_executor(
            _GRAPH_CACHE_REFRESH_EXECUTOR,
            _kg_view_cache_build_and_store,
            cache_key,
            lock_token,
            builder,
        )
        return JSONResponse(content=payload, headers={"X-KG-View-Cache": "miss"})

    warmed = await _kg_view_cache_wait(cache_key)
    if warmed is not None and isinstance(warmed.get("payload"), dict):
        return JSONResponse(content=warmed["payload"], headers={"X-KG-View-Cache": "wait-hit"})

    return JSONResponse(
        status_code=503,
        content={
            "error": "kg_view_cache_warming",
            "message": "Knowledge graph view is warming. Please retry shortly.",
        },
        headers={"Retry-After": "3", "X-KG-View-Cache": "warming"},
    )


def _kg_view_param_list(values: Optional[List[str]]) -> List[str]:
    return sorted({str(item).strip() for item in (values or []) if str(item).strip()})


def _kg_view_text(value: Any, fallback: str = "") -> str:
    text = str(value if value is not None else fallback).strip()
    return text or fallback


def _kg_view_ticket_key(ticket: str) -> str:
    return f"kg-view-ticket:{ticket}"


def _kg_view_public_user_payload(user: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not user:
        return {}
    user_id = _kg_view_text(user.get("id"))
    if not user_id:
        return {}
    return {
        "id": user_id,
        "email": _kg_view_text(user.get("email")),
        "username": _kg_view_text(user.get("username")),
        "role": _normalize_role(_kg_view_text(user.get("role"), "user")),
    }


def _kg_view_ticket_value(request: Request) -> str:
    return _kg_view_text(request.query_params.get("ticket") or request.cookies.get("kg_view_ticket"))


def _kg_view_ticket_payload_from_value(value: Any) -> Optional[Dict[str, Any]]:
    if not value:
        return None
    if isinstance(value, bytes):
        raw = value.decode("utf-8", errors="ignore")
    elif isinstance(value, dict):
        payload = value
        raw = ""
    else:
        raw = str(value)
    if not isinstance(value, dict):
        try:
            payload = json.loads(raw)
        except (TypeError, json.JSONDecodeError):
            return None
    if not isinstance(payload, dict):
        return None
    try:
        expires_at = float(payload.get("expires_at") or 0)
    except (TypeError, ValueError):
        return None
    if expires_at <= time.time():
        return None
    user = _kg_view_public_user_payload(payload.get("user") if isinstance(payload.get("user"), dict) else {})
    if not user:
        return None
    return {
        "user": user,
        "document_id": _kg_view_text(payload.get("document_id")),
        "expires_at": expires_at,
    }


def _kg_view_store_ticket(ticket: str, user: Dict[str, Any], document_id: str = "") -> Dict[str, Any]:
    payload = {
        "user": _kg_view_public_user_payload(user),
        "document_id": _kg_view_text(document_id),
        "expires_at": time.time() + _KG_VIEW_TICKET_TTL_SECONDS,
    }
    with _KG_VIEW_TICKETS_LOCK:
        _KG_VIEW_TICKETS[ticket] = payload
        now = time.time()
        expired = [key for key, item in _KG_VIEW_TICKETS.items() if float(item.get("expires_at") or 0) <= now]
        for key in expired[:128]:
            _KG_VIEW_TICKETS.pop(key, None)
    if _kg_view_redis_available_for_cache():
        try:
            chat_memory_service.client.execute(
                "SET",
                _kg_view_ticket_key(ticket),
                json.dumps(payload, ensure_ascii=False, separators=(",", ":")),
                "EX",
                _KG_VIEW_TICKET_TTL_SECONDS,
            )
        except RedisUnavailableError:
            _kg_view_note_redis_failure()
    return payload


def _kg_view_get_ticket_payload(request: Request) -> Optional[Dict[str, Any]]:
    if hasattr(request.state, "kg_view_ticket_payload"):
        return request.state.kg_view_ticket_payload
    ticket = _kg_view_ticket_value(request)
    payload: Optional[Dict[str, Any]] = None
    if ticket:
        with _KG_VIEW_TICKETS_LOCK:
            payload = _kg_view_ticket_payload_from_value(_KG_VIEW_TICKETS.get(ticket))
            if payload is None:
                _KG_VIEW_TICKETS.pop(ticket, None)
        if payload is None and _kg_view_redis_available_for_cache():
            try:
                payload = _kg_view_ticket_payload_from_value(
                    chat_memory_service.client.execute("GET", _kg_view_ticket_key(ticket))
                )
            except RedisUnavailableError:
                _kg_view_note_redis_failure()
                payload = None
            if payload is not None:
                with _KG_VIEW_TICKETS_LOCK:
                    _KG_VIEW_TICKETS[ticket] = payload
    request.state.kg_view_ticket_payload = payload
    return payload


def _kg_view_ticket_user(request: Request) -> Optional[Dict[str, Any]]:
    payload = _kg_view_get_ticket_payload(request)
    user = payload.get("user") if isinstance(payload, dict) else None
    return user if isinstance(user, dict) else None


def _kg_view_ticket_document_id(request: Request) -> str:
    payload = _kg_view_get_ticket_payload(request)
    if not isinstance(payload, dict):
        return ""
    return _kg_view_text(payload.get("document_id"))


def _kg_view_humanize(value: Any) -> str:
    text = _kg_view_text(value)
    text = re.sub(r"[_-]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text.title() if text else "Entity"


def _kg_view_domain(value: Any, *hints: Any) -> str:
    raw = " ".join(_kg_view_text(item) for item in (value, *hints)).lower()
    if "environment" in raw or "emission" in raw or "climate" in raw or "carbon" in raw:
        return "environmental"
    if "social" in raw or "employee" in raw or "workforce" in raw or "supplier" in raw or "community" in raw:
        return "social"
    if "govern" in raw or "board" in raw or "audit" in raw or "ethic" in raw or "compliance" in raw:
        return "governance"
    if raw.strip() == "ai" or "artificial" in raw or "model" in raw or "data center" in raw:
        return "ai"
    normalized = _kg_view_text(value, "general").lower()
    return normalized if normalized in {"environmental", "social", "governance", "ai", "general"} else "general"


def _kg_view_requested_scope(request: Request) -> str:
    raw = _kg_view_text(
        request.query_params.get("scope")
        or request.query_params.get("graph_scope")
        or request.cookies.get("kg_view_scope")
    ).lower()
    if raw in {"all", "global", "public"}:
        return "all"
    if raw in {"document", "current", "doc"}:
        return "document"
    return ""


def _kg_view_raw_document_id(request: Request) -> str:
    value = (
        request.query_params.get("document_id")
        or request.cookies.get("kg_document_id")
        or _kg_view_ticket_document_id(request)
        or ""
    )
    return _kg_view_text(value)


def _kg_view_request_document_id(request: Request) -> str:
    if _kg_view_requested_scope(request) == "all":
        return ""
    return _kg_view_raw_document_id(request)


def _kg_view_effective_user(request: Request, current_user: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    if current_user is not None:
        return current_user
    ticket_user = _kg_view_ticket_user(request)
    if ticket_user is not None:
        return ticket_user
    if _is_local_request(request):
        return {"id": "local-admin", "email": "local", "username": "Local Admin", "role": "admin"}
    return None


def _kg_view_can_use_document_scope(
    request: Request,
    current_user: Optional[Dict[str, Any]],
    document_id: str,
) -> bool:
    if not document_id:
        return False
    effective_user = _kg_view_effective_user(request, current_user)
    if _is_admin_user(effective_user) or _is_local_request(request):
        return True
    try:
        entry = document_registry.get_entry(document_id, valid_only=True)
        if entry is None:
            audit = get_latest_upload_by_document_id(document_id)
            if audit is not None:
                entry = _entry_from_upload(audit)
        if entry is None:
            return False
        if effective_user:
            return _can_access_entry(effective_user, entry)
        return _is_global_entry(entry)
    except Exception:
        return False



def _kg_view_active_scope(
    request: Request,
    current_user: Optional[Dict[str, Any]],
    document_id: str,
) -> str:
    requested = _kg_view_requested_scope(request)
    if requested == "all":
        return "all"
    if _kg_view_can_use_document_scope(request, current_user, document_id):
        return "document"
    return "all"


def _kg_view_document_graph_from_registry(
    document_id: str,
    current_user: Optional[Dict[str, Any]],
) -> Dict[str, Any]:
    document_id = _kg_view_text(document_id)
    if not document_id:
        return {"nodes": [], "edges": []}
    try:
        entry = document_registry.get_entry(document_id, valid_only=True)
        audit = get_latest_upload_by_document_id(document_id)
        if entry is None and audit is not None:
            entry = _entry_from_upload(audit)
        if entry is None:
            return {"nodes": [], "edges": []}
        if current_user and not _can_access_entry(current_user, entry):
            return {"nodes": [], "edges": []}
        if not current_user and not _is_global_entry(entry):
            return {"nodes": [], "edges": []}
        if audit and str(audit.get("status") or "") in {"deleted", "deleted_with_warnings"}:
            return {"nodes": [], "edges": []}
        document = load_registered_document(entry, audit=audit)
        graph = document.get("graph") if isinstance(document, dict) else None
        if isinstance(graph, dict):
            return {
                "nodes": list(graph.get("nodes") or []),
                "edges": list(graph.get("edges") or []),
                "metadata": dict(graph.get("metadata") or {}),
            }
    except Exception as exc:
        print(f"[kg-view] registry document graph fallback failed: {type(exc).__name__}: {exc}")
    return {"nodes": [], "edges": []}


def _kg_view_source_graph_uncached(
    request: Request,
    current_user: Optional[Dict[str, Any]],
    *,
    node_limit: int = 25000,
    edge_limit: int = 30000,
) -> Dict[str, Any]:
    effective_user = _kg_view_effective_user(request, current_user)
    document_id = _kg_view_request_document_id(request)
    if document_id and _kg_view_can_use_document_scope(request, effective_user, document_id):
        try:
            store = get_neo4j_store()
            if store is not None:
                raw = store.get_visualization_graph(limit=node_limit, document_id=document_id)
                graph = _neo4j_visualization_to_graph_data(raw, edge_limit=edge_limit)
                if graph.get("nodes") or graph.get("edges"):
                    return graph
        except Exception as exc:
            print(f"[kg-view] document graph fallback: {type(exc).__name__}: {exc}")
        registry_graph = _kg_view_document_graph_from_registry(document_id, effective_user)
        if registry_graph.get("nodes") or registry_graph.get("edges"):
            return registry_graph

    return _load_real_knowledge_graph(
        current_user=effective_user,
        node_limit=node_limit,
        edge_limit=edge_limit,
    )


def _kg_view_source_graph(
    request: Request,
    current_user: Optional[Dict[str, Any]],
    *,
    node_limit: int = 25000,
    edge_limit: int = 30000,
) -> Dict[str, Any]:
    node_limit = max(10, min(int(node_limit or 25000), _PUBLIC_GRAPH_MAX_NODES))
    edge_limit = max(0, min(int(edge_limit or 30000), _PUBLIC_GRAPH_MAX_EDGES))
    if not _KG_VIEW_CACHE_ENABLED:
        return _kg_view_source_graph_uncached(
            request,
            current_user,
            node_limit=node_limit,
            edge_limit=edge_limit,
        )

    cache_key = _kg_view_cache_key(
        request,
        current_user,
        "source-graph",
        {"node_limit": node_limit, "edge_limit": edge_limit},
    )
    builder = lambda: _kg_view_source_graph_uncached(
        request,
        current_user,
        node_limit=node_limit,
        edge_limit=edge_limit,
    )

    cached = _kg_view_cache_get(cache_key)
    if cached is not None and isinstance(cached.get("payload"), dict):
        graph = cached["payload"]
        if not _kg_view_cache_is_fresh(cached):
            lock_token = _kg_view_cache_acquire_lock(cache_key)
            if lock_token:
                _GRAPH_CACHE_REFRESH_EXECUTOR.submit(
                    _kg_view_cache_refresh_background,
                    cache_key,
                    lock_token,
                    builder,
                )
        return graph

    lock_token = _kg_view_cache_acquire_lock(cache_key)
    if lock_token:
        return _kg_view_cache_build_and_store(cache_key, lock_token, builder)

    deadline = time.time() + _KG_VIEW_CACHE_WAIT_SECONDS
    while time.time() < deadline:
        time.sleep(0.25)
        warmed = _kg_view_cache_get(cache_key)
        if warmed is not None and isinstance(warmed.get("payload"), dict):
            return warmed["payload"]

    raise HTTPException(
        status_code=503,
        detail="Knowledge graph view is warming. Please retry shortly.",
    )


def _kg_view_filtered_text_graph(
    graph: Dict[str, Any],
    *,
    domains: Optional[List[str]] = None,
    companies: Optional[List[str]] = None,
    years: Optional[List[str]] = None,
    limit: int = 5000,
) -> Dict[str, Any]:
    domain_set = {str(item).lower() for item in (domains or []) if str(item).strip()}
    company_set = {str(item) for item in (companies or []) if str(item).strip()}
    year_set = {str(item) for item in (years or []) if str(item).strip()}
    if year_set & {"All", "all", "__all__"}:
        year_set = set()
    bounded_limit = max(10, min(int(limit or 5000), 30000))

    nodes: List[Dict[str, Any]] = []
    source_nodes = graph.get("nodes") or []
    for row in source_nodes:
        if not isinstance(row, dict):
            continue
        node_id = _kg_view_text(row.get("id"))
        if not node_id:
            continue
        label = _kg_view_text(row.get("label") or row.get("name") or row.get("normalizedName"), node_id)
        domain = _kg_view_domain(row.get("domain") or row.get("esg_domain"), row.get("type"), label, row.get("description"))
        company = _kg_view_text(row.get("company") or row.get("metadata", {}).get("company")) or "Unknown"
        year = _kg_view_text(row.get("year") or row.get("metadata", {}).get("year"))
        if domain_set and domain not in domain_set:
            continue
        if company_set and company not in company_set:
            continue
        if year_set and year not in year_set:
            continue
        nodes.append(
            {
                "id": node_id,
                "label": label,
                "text": _kg_view_text(row.get("description") or label, label),
                "type": _kg_view_text(row.get("type"), "Entity").upper(),
                "esg_domain": domain,
                "year": int(year) if year.isdigit() else year,
                "company": company,
            }
        )
        if len(nodes) >= bounded_limit:
            break

    node_ids = {node["id"] for node in nodes}
    edges: List[Dict[str, Any]] = []
    for index, row in enumerate(graph.get("edges") or []):
        if not isinstance(row, dict):
            continue
        source = _kg_view_text(row.get("source"))
        target = _kg_view_text(row.get("target"))
        if source not in node_ids or target not in node_ids:
            continue
        rel_type = _kg_view_text(row.get("relationship_type") or row.get("type"), "RELATED_TO")
        confidence = _safe_graph_float(row.get("confidence"), 0.75)
        edges.append(
            {
                "id": _kg_view_text(row.get("id"), f"edge_{index}"),
                "source": source,
                "target": target,
                "type": rel_type,
                "action": _kg_view_text(row.get("relationship_action") or row.get("action")),
                "category": _kg_view_text(row.get("category") or row.get("domain") or rel_type),
                "nature": _kg_view_text(row.get("relationship_nature") or row.get("nature")),
                "evidence": _kg_view_text(row.get("evidence")),
                "direction": _kg_view_text(row.get("direction"), "e1_to_e2"),
                "credibility_score": round(max(0.0, min(confidence * 5.0, 5.0)), 2),
                "sentiment": _kg_view_text(row.get("sentiment"), "neutral"),
            }
        )
    return {"nodes": nodes, "edges": edges}


def _kg_view_filter_values(graph: Dict[str, Any]) -> Dict[str, List[Any]]:
    text_graph = _kg_view_filtered_text_graph(graph, limit=30000)
    companies = sorted({node["company"] for node in text_graph["nodes"] if node.get("company") and node.get("company") != "Unknown"})
    years = sorted({str(node.get("year") or "").strip() for node in text_graph["nodes"] if str(node.get("year") or "").strip()})
    domains = sorted({node["esg_domain"] for node in text_graph["nodes"] if node.get("esg_domain")})
    return {
        "companies": companies or ["Unknown"],
        "years": years or ["All"],
        "domains": domains or ["environmental", "social", "governance", "ai"],
    }


def _kg_view_llm_labels(domain: str, groups: List[Dict[str, Any]]) -> Dict[str, str]:
    if not _KG_VIEW_LLM_CLUSTER_LABELS or not groups:
        return {}
    client = get_openai_client()
    if client is None:
        return {}
    payload = [
        {
            "id": str(group["community_id"]),
            "fallback": group["label"],
            "members": group.get("members", [])[:16],
            "types": group.get("types", [])[:6],
        }
        for group in groups
        if group.get("community_id") != -1
    ][:8]
    if not payload:
        return {}
    cache_source = json.dumps({"domain": domain, "groups": payload}, ensure_ascii=False, sort_keys=True)
    cache_key = hashlib.sha256(cache_source.encode("utf-8")).hexdigest()[:24]
    if cache_key in _KG_VIEW_CLUSTER_LABEL_CACHE:
        return _KG_VIEW_CLUSTER_LABEL_CACHE[cache_key]
    try:
        response = client.chat.completions.create(
            model=_KG_VIEW_CLUSTER_LABEL_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You label ESG knowledge graph communities. Return only JSON: "
                        "{\"labels\":[{\"id\":\"...\",\"label\":\"2-5 word label\"}]}."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Domain: {domain}\n"
                        "Create concise business-readable labels for these graph communities. "
                        "Avoid generic labels like 'Miscellaneous' unless unavoidable.\n"
                        f"{json.dumps(payload, ensure_ascii=False)}"
                    ),
                },
            ],
            temperature=0,
            **chat_token_kwargs(_KG_VIEW_CLUSTER_LABEL_MODEL, 500),
        )
        content = response.choices[0].message.content or ""
        match = re.search(r"\{.*\}", content, flags=re.S)
        parsed = json.loads(match.group(0) if match else content)
        labels = {
            _kg_view_text(item.get("id")): _kg_view_text(item.get("label"))[:48]
            for item in parsed.get("labels", [])
            if isinstance(item, dict) and _kg_view_text(item.get("id")) and _kg_view_text(item.get("label"))
        }
        _KG_VIEW_CLUSTER_LABEL_CACHE[cache_key] = labels
        return labels
    except Exception as exc:
        print(f"[kg-view] LLM cluster labels skipped: {type(exc).__name__}: {exc}")
        return {}


def _kg_view_domain_grouping(
    graph: Dict[str, Any],
    domain: str,
    *,
    companies: Optional[List[str]] = None,
    years: Optional[List[str]] = None,
) -> Dict[str, Any]:
    from collections import Counter, defaultdict

    domain = _kg_view_domain(domain)
    detail_graph = _kg_view_filtered_text_graph(
        graph,
        domains=[domain],
        companies=companies,
        years=years,
        limit=5000,
    )
    nodes = detail_graph["nodes"]
    edges = detail_graph["edges"]
    if not nodes:
        return {"nodes": [], "edges": [], "cluster": domain, "total_triples": 0, "_assignments": {}, "_detail": detail_graph}

    node_by_id = {node["id"]: node for node in nodes}
    adjacency: Dict[str, Counter] = defaultdict(Counter)
    for edge in edges:
        source = edge["source"]
        target = edge["target"]
        if source in node_by_id and target in node_by_id and source != target:
            adjacency[source][target] += 1
            adjacency[target][source] += 1

    communities: List[List[str]] = []
    try:
        import networkx as nx

        graph_obj = nx.Graph()
        for node in nodes:
            graph_obj.add_node(node["id"])
        for source, targets in adjacency.items():
            for target, weight in targets.items():
                if source < target:
                    graph_obj.add_edge(source, target, weight=weight)
        if graph_obj.number_of_edges() > 0 and graph_obj.number_of_nodes() >= 3:
            communities = [sorted(group) for group in nx.algorithms.community.greedy_modularity_communities(graph_obj, weight="weight")]
    except Exception as exc:
        print(f"[kg-view] community detection fallback: {type(exc).__name__}: {exc}")

    if not communities:
        grouped: Dict[str, List[str]] = defaultdict(list)
        for node in nodes:
            grouped[_kg_view_text(node.get("type"), "Entity")].append(node["id"])
        communities = sorted(grouped.values(), key=len, reverse=True)

    communities = sorted(communities, key=len, reverse=True)
    named_groups = [group for group in communities if len(group) >= 2][:8]
    named_ids = {node_id for group in named_groups for node_id in group}
    other_group = [node["id"] for node in nodes if node["id"] not in named_ids]

    assignments: Dict[str, int] = {}
    response_nodes: List[Dict[str, Any]] = []
    domain_color = {
        "environmental": "#3fb950",
        "social": "#58a6ff",
        "governance": "#bc8cff",
        "ai": "#f0883e",
    }.get(domain, "#8b949e")

    def build_group_payload(community_id: int, ids: List[str], *, is_other: bool = False) -> Dict[str, Any]:
        members = [node_by_id[node_id]["label"] for node_id in ids if node_id in node_by_id]
        type_counts = Counter(node_by_id[node_id]["type"] for node_id in ids if node_id in node_by_id)
        top_type = type_counts.most_common(1)[0][0] if type_counts else "Entity"
        degree_rank = sorted(ids, key=lambda node_id: sum(adjacency.get(node_id, {}).values()), reverse=True)
        hub_label = node_by_id.get(degree_rank[0], {}).get("label") if degree_rank else ""
        fallback_label = "Other" if is_other else (_kg_view_text(hub_label) or _kg_view_humanize(top_type)).title()
        return {
            "id": "comm_other" if is_other else f"comm_{community_id}",
            "label": fallback_label[:48],
            "color": "#8b949e" if is_other else domain_color,
            "triple_count": sum(1 for edge in edges if edge["source"] in ids or edge["target"] in ids),
            "concept_count": len(members),
            "size": 20 + min(len(members) * 3, 45),
            "community_id": -1 if is_other else community_id,
            "members": members,
            "types": [item for item, _ in type_counts.most_common(6)],
            "is_other": is_other,
        }

    for community_id, group in enumerate(named_groups):
        for node_id in group:
            assignments[node_id] = community_id
        response_nodes.append(build_group_payload(community_id, group))

    if other_group:
        for node_id in other_group:
            assignments[node_id] = -1
        response_nodes.append(build_group_payload(-1, other_group, is_other=True))

    llm_labels = _kg_view_llm_labels(domain, response_nodes)
    for node in response_nodes:
        label = llm_labels.get(str(node["community_id"]))
        if label:
            node["label"] = label

    cross_edges: Counter = Counter()
    for edge in edges:
        source_group = assignments.get(edge["source"], -1)
        target_group = assignments.get(edge["target"], -1)
        if source_group == target_group:
            continue
        cross_edges[(source_group, target_group)] += 1

    def group_node_id(community_id: int) -> str:
        return "comm_other" if community_id == -1 else f"comm_{community_id}"

    response_edges = [
        {
            "source": group_node_id(source),
            "target": group_node_id(target),
            "weight": count,
            "width": max(1, min(count // 2, 12)),
        }
        for (source, target), count in cross_edges.most_common()
    ]

    return {
        "nodes": response_nodes,
        "edges": response_edges,
        "cluster": domain,
        "total_triples": len(edges),
        "_assignments": assignments,
        "_detail": detail_graph,
    }


def _kg_view_public_grouping_payload(grouping: Dict[str, Any]) -> Dict[str, Any]:
    return {key: value for key, value in grouping.items() if not key.startswith("_")}


def _kg_view_inject_scope_switcher(html: str, *, active_scope: str, has_document_scope: bool) -> str:
    if "kg-scope-switcher" in html:
        return html
    active_scope_json = json.dumps(active_scope)
    has_document_json = json.dumps(bool(has_document_scope))
    snippet = f"""
<style>
  #kg-scope-switcher {{
    position: fixed;
    top: 14px;
    right: 18px;
    z-index: 10000;
    display: flex;
    align-items: center;
    gap: 8px;
    padding: 8px;
    border: 1px solid #30363d;
    border-radius: 8px;
    background: rgba(13, 17, 23, 0.92);
    box-shadow: 0 12px 30px rgba(0, 0, 0, 0.28);
    color: #c9d1d9;
    font: 12px/1.2 -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
  }}
  #kg-scope-switcher .kg-scope-label {{
    color: #8b949e;
    font-weight: 700;
    letter-spacing: 0.06em;
    text-transform: uppercase;
  }}
  #kg-scope-switcher button {{
    border: 1px solid #30363d;
    border-radius: 6px;
    background: #161b22;
    color: #c9d1d9;
    cursor: pointer;
    font: inherit;
    font-weight: 700;
    padding: 7px 10px;
  }}
  #kg-scope-switcher button.active {{
    border-color: #58a6ff;
    background: #1f6feb;
    color: #fff;
  }}
  #kg-scope-switcher button:disabled {{
    cursor: not-allowed;
    opacity: 0.45;
  }}
</style>
<script>
(function() {{
  var activeScope = {active_scope_json};
  var hasDocumentScope = {has_document_json};
  try {{
    var cleanUrl = new URL(window.location.href);
    if (cleanUrl.searchParams.has('ticket')) {{
      cleanUrl.searchParams.delete('ticket');
      window.history.replaceState({{}}, document.title, cleanUrl.toString());
    }}
  }} catch (err) {{}}
  function setCookie(name, value) {{
    document.cookie = name + '=' + encodeURIComponent(value) + '; path=/; max-age=3600; SameSite=Lax';
  }}
  function setScope(scope) {{
    if (scope === 'document' && !hasDocumentScope) return;
    setCookie('kg_view_scope', scope);
    var url = new URL(window.location.href);
    url.searchParams.set('scope', scope);
    window.location.href = url.toString();
  }}
  function addSwitcher() {{
    if (document.getElementById('kg-scope-switcher')) return;
    var el = document.createElement('div');
    el.id = 'kg-scope-switcher';
    el.innerHTML =
      '<span class="kg-scope-label">Graph</span>' +
      '<button type="button" data-scope="document">Current document</button>' +
      '<button type="button" data-scope="all">All documents</button>';
    document.body.appendChild(el);
    var buttons = el.querySelectorAll('button[data-scope]');
    buttons.forEach(function(button) {{
      var scope = button.getAttribute('data-scope');
      if (scope === activeScope) button.classList.add('active');
      if (scope === 'document' && !hasDocumentScope) button.disabled = true;
      button.addEventListener('click', function() {{ setScope(scope); }});
    }});
  }}
  if (document.readyState === 'loading') {{
    document.addEventListener('DOMContentLoaded', addSwitcher);
  }} else {{
    addSwitcher();
  }}
}})();
</script>
"""
    marker = "</body>"
    if marker in html:
        return html.replace(marker, snippet + "\n" + marker, 1)
    return html + snippet


@app.get("/health")
async def health() -> JSONResponse:
    return JSONResponse(content={"status": "ok"})


@app.get("/healthz")
async def healthz() -> JSONResponse:
    return JSONResponse(content={"status": "ok"})


@app.post("/kg-view/ticket")
async def create_kg_view_ticket(
    payload: KgViewTicketRequest,
    request: Request,
    current_user: dict = Depends(get_current_user),
) -> JSONResponse:
    user_payload = _kg_view_public_user_payload(current_user)
    if not user_payload:
        raise HTTPException(status_code=401, detail="Not authenticated")
    document_id = _kg_view_text(payload.document_id)
    if document_id and not _kg_view_can_use_document_scope(request, current_user, document_id):
        raise HTTPException(status_code=403, detail="Document is not accessible")
    ticket = secrets.token_urlsafe(32)
    _kg_view_store_ticket(ticket, current_user, document_id)
    return JSONResponse(content={"ticket": ticket, "expires_in": _KG_VIEW_TICKET_TTL_SECONDS})


@app.get("/kg-view", response_class=HTMLResponse)
async def kg_view(
    request: Request,
    document_id: Optional[str] = None,
    current_user: Optional[dict] = Depends(get_optional_current_user),
) -> HTMLResponse:
    template = _TEXT_KG_VIEW_TEMPLATE if _TEXT_KG_VIEW_TEMPLATE.exists() else _KG_VIEW_TEMPLATE
    if not template.exists():
        raise HTTPException(status_code=404, detail="KG view template not found")
    ticket_payload = _kg_view_get_ticket_payload(request)
    effective_user = _kg_view_effective_user(request, current_user)
    raw_document_id = _kg_view_text(
        document_id
        or request.cookies.get("kg_document_id")
        or _kg_view_ticket_document_id(request)
    )
    has_document_scope = _kg_view_can_use_document_scope(request, effective_user, raw_document_id)
    active_scope = _kg_view_active_scope(request, effective_user, raw_document_id)
    html = _kg_view_inject_scope_switcher(
        template.read_text(encoding="utf-8"),
        active_scope=active_scope,
        has_document_scope=has_document_scope,
    )
    response = HTMLResponse(
        html,
        headers={"Cache-Control": "no-store, max-age=0"},
    )
    ticket = _kg_view_ticket_value(request)
    if ticket and ticket_payload:
        try:
            max_age = max(1, int(float(ticket_payload.get("expires_at") or 0) - time.time()))
        except (TypeError, ValueError):
            max_age = _KG_VIEW_TICKET_TTL_SECONDS
        response.set_cookie("kg_view_ticket", ticket, max_age=max_age, httponly=True, samesite="lax")
    elif request.cookies.get("kg_view_ticket") and ticket_payload is None:
        response.delete_cookie("kg_view_ticket")
    if raw_document_id and has_document_scope:
        response.set_cookie("kg_document_id", raw_document_id, max_age=3600, httponly=True, samesite="lax")
    elif not raw_document_id:
        response.delete_cookie("kg_document_id")
    response.set_cookie("kg_view_scope", active_scope, max_age=3600, httponly=False, samesite="lax")
    return response


@app.get("/api/filters")
async def text_kg_filters(
    request: Request,
    companies: List[str] = Query(default=[]),
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    selected_companies = _kg_view_param_list(companies)

    def build_payload() -> Dict[str, Any]:
        graph = _kg_view_source_graph(request, current_user)
        all_values = _kg_view_filter_values(graph)
        active_graph = _kg_view_filtered_text_graph(
            graph,
            companies=selected_companies or None,
            limit=30000,
        )
        active_years = sorted({str(node.get("year") or "").strip() for node in active_graph["nodes"] if str(node.get("year") or "").strip()})
        active_domains = sorted({node["esg_domain"] for node in active_graph["nodes"] if node.get("esg_domain")})
        return {
            "companies": all_values["companies"],
            "years": all_values["years"],
            "domains": all_values["domains"],
            "active_years": active_years or all_values["years"],
            "active_domains": active_domains or all_values["domains"],
        }

    return await _kg_view_cached_json(
        request,
        current_user,
        "filters",
        {"companies": selected_companies},
        build_payload,
    )


@app.get("/api/total_count")
async def text_kg_total_count(
    request: Request,
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    def build_payload() -> Dict[str, Any]:
        graph = _kg_view_source_graph(request, current_user)
        total_nodes = len(graph.get("nodes") or [])
        return {
            "total_nodes": min(total_nodes, _KG_VIEW_DETAIL_RENDER_NODE_LIMIT),
            "available_nodes": total_nodes,
        }

    return await _kg_view_cached_json(
        request,
        current_user,
        "total-count",
        {},
        build_payload,
    )


@app.get("/api/graph")
async def text_kg_graph(
    request: Request,
    years: List[str] = Query(default=[]),
    companies: List[str] = Query(default=[]),
    domains: List[str] = Query(default=[]),
    limit: int = Query(default=500, ge=10, le=30000),
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    normalized_years = _kg_view_param_list(years)
    normalized_companies = _kg_view_param_list(companies)
    normalized_domains = _kg_view_param_list(domains)
    effective_limit = min(limit, _KG_VIEW_DETAIL_RENDER_NODE_LIMIT)

    def build_payload() -> Dict[str, Any]:
        graph = _kg_view_source_graph(request, current_user, node_limit=max(effective_limit, 5000), edge_limit=30000)
        return _kg_view_filtered_text_graph(
            graph,
            domains=normalized_domains or None,
            companies=normalized_companies or None,
            years=normalized_years or None,
            limit=effective_limit,
        )

    return await _kg_view_cached_json(
        request,
        current_user,
        "graph",
        {
            "years": normalized_years,
            "companies": normalized_companies,
            "domains": normalized_domains,
            "limit": effective_limit,
        },
        build_payload,
    )


@app.get("/api/cluster-graph")
async def text_kg_cluster_graph(
    request: Request,
    years: List[str] = Query(default=[]),
    companies: List[str] = Query(default=[]),
    domains: List[str] = Query(default=[]),
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    normalized_years = _kg_view_param_list(years)
    normalized_companies = _kg_view_param_list(companies)
    normalized_domains = _kg_view_param_list(domains)

    def build_payload() -> Dict[str, Any]:
        from collections import Counter, defaultdict

        graph = _kg_view_source_graph(request, current_user)
        data = _kg_view_filtered_text_graph(
            graph,
            domains=normalized_domains or None,
            companies=normalized_companies or None,
            years=normalized_years or None,
            limit=30000,
        )
        colors = {
            "environmental": "#3fb950",
            "social": "#58a6ff",
            "governance": "#bc8cff",
            "ai": "#f0883e",
            "general": "#8b949e",
        }
        labels = {
            "environmental": "Environmental",
            "social": "Social",
            "governance": "Governance",
            "ai": "AI",
            "general": "General",
        }
        node_domain = {node["id"]: node["esg_domain"] for node in data["nodes"]}
        concept_sets: Dict[str, set] = defaultdict(set)
        triple_counts: Counter = Counter()
        edge_counts: Counter = Counter()
        edge_relations: Dict[Tuple[str, str], Counter] = defaultdict(Counter)
        for node in data["nodes"]:
            concept_sets[node["esg_domain"]].add(node["label"])
        for edge in data["edges"]:
            d1 = node_domain.get(edge["source"], "general")
            d2 = node_domain.get(edge["target"], "general")
            triple_counts[d1] += 1
            triple_counts[d2] += 1
            key = (min(d1, d2), max(d1, d2))
            edge_counts[key] += 1
            edge_relations[key][edge.get("type") or "RELATED_TO"] += 1

        domain_order = [domain for domain in ["environmental", "social", "governance", "ai", "general"] if domain in colors]
        visible_domains = set(node_domain.values()) | set(normalized_domains)
        nodes = []
        for domain in domain_order:
            if domain == "general" and domain not in visible_domains:
                continue
            nodes.append(
                {
                    "id": labels[domain],
                    "label": labels[domain],
                    "domain": domain,
                    "color": colors[domain],
                    "triple_count": triple_counts.get(domain, 0),
                    "concept_count": len(concept_sets.get(domain, set())),
                    "size": 30 + min(triple_counts.get(domain, 0) // 10, 40),
                }
            )
        edges = [
            {
                "source": labels.get(src, _kg_view_humanize(src)),
                "target": labels.get(tgt, _kg_view_humanize(tgt)),
                "weight": count,
                "top_relations": [{"type": rel, "count": rel_count} for rel, rel_count in edge_relations[(src, tgt)].most_common(3)],
                "width": max(2, min(count / 5, 15)),
            }
            for (src, tgt), count in edge_counts.most_common()
            if src in labels and tgt in labels
        ]
        return {"nodes": nodes, "edges": edges, "total_triples": len(data["edges"])}

    return await _kg_view_cached_json(
        request,
        current_user,
        "cluster-graph",
        {
            "years": normalized_years,
            "companies": normalized_companies,
            "domains": normalized_domains,
        },
        build_payload,
    )


@app.get("/api/cluster-detail")
async def text_kg_cluster_detail(
    request: Request,
    cluster: str,
    years: List[str] = Query(default=[]),
    companies: List[str] = Query(default=[]),
    limit: int = Query(default=500, ge=10, le=30000),
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    domain = _kg_view_domain(cluster)
    normalized_years = _kg_view_param_list(years)
    normalized_companies = _kg_view_param_list(companies)

    def build_payload() -> Dict[str, Any]:
        graph = _kg_view_source_graph(request, current_user)
        return _kg_view_filtered_text_graph(
            graph,
            domains=[domain],
            companies=normalized_companies or None,
            years=normalized_years or None,
            limit=limit,
        )

    return await _kg_view_cached_json(
        request,
        current_user,
        "cluster-detail",
        {
            "cluster": domain,
            "years": normalized_years,
            "companies": normalized_companies,
            "limit": limit,
        },
        build_payload,
    )


@app.get("/api/cluster-subgraph")
async def text_kg_cluster_subgraph(
    request: Request,
    cluster: str,
    years: List[str] = Query(default=[]),
    companies: List[str] = Query(default=[]),
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    domain = _kg_view_domain(cluster)
    normalized_years = _kg_view_param_list(years)
    normalized_companies = _kg_view_param_list(companies)

    def build_payload() -> Dict[str, Any]:
        graph = _kg_view_source_graph(request, current_user)
        grouping = _kg_view_domain_grouping(
            graph,
            domain,
            companies=normalized_companies or None,
            years=normalized_years or None,
        )
        return _kg_view_public_grouping_payload(grouping)

    return await _kg_view_cached_json(
        request,
        current_user,
        "cluster-subgraph",
        {
            "cluster": domain,
            "years": normalized_years,
            "companies": normalized_companies,
        },
        build_payload,
    )


@app.get("/api/cross-domain-communities")
async def text_kg_cross_domain_communities(
    request: Request,
    domain1: str,
    domain2: str,
    years: List[str] = Query(default=[]),
    companies: List[str] = Query(default=[]),
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    d1 = _kg_view_domain(domain1)
    d2 = _kg_view_domain(domain2)
    normalized_years = _kg_view_param_list(years)
    normalized_companies = _kg_view_param_list(companies)

    def build_payload() -> Dict[str, Any]:
        from collections import Counter

        graph = _kg_view_source_graph(request, current_user)
        g1 = _kg_view_domain_grouping(graph, d1, companies=normalized_companies or None, years=normalized_years or None)
        g2 = _kg_view_domain_grouping(graph, d2, companies=normalized_companies or None, years=normalized_years or None)
        data = _kg_view_filtered_text_graph(
            graph,
            domains=[d1, d2],
            companies=normalized_companies or None,
            years=normalized_years or None,
            limit=30000,
        )
        node_domain = {node["id"]: node["esg_domain"] for node in data["nodes"]}
        edge_counts: Counter = Counter()
        d1_counts: Counter = Counter()
        d2_counts: Counter = Counter()
        for edge in data["edges"]:
            source_domain = node_domain.get(edge["source"])
            target_domain = node_domain.get(edge["target"])
            if {source_domain, target_domain} != {d1, d2}:
                continue
            source_id = edge["source"] if source_domain == d1 else edge["target"]
            target_id = edge["target"] if target_domain == d2 else edge["source"]
            c1 = g1["_assignments"].get(source_id, -1)
            c2 = g2["_assignments"].get(target_id, -1)
            if c1 == -1 or c2 == -1:
                continue
            edge_counts[(c1, c2)] += 1
            d1_counts[c1] += 1
            d2_counts[c2] += 1

        def make_node(domain: str, group_node: Dict[str, Any], count: int) -> Dict[str, Any]:
            cid = group_node["community_id"]
            return {
                "id": f"{domain}_{cid}",
                "label": group_node["label"],
                "domain": domain,
                "color": group_node["color"],
                "concept_count": count,
                "same_domain_count": group_node["concept_count"],
                "members": group_node.get("members", []),
                "cross_members": group_node.get("members", []),
                "community_id": cid,
            }

        g1_by_cid = {node["community_id"]: node for node in g1["nodes"]}
        g2_by_cid = {node["community_id"]: node for node in g2["nodes"]}
        nodes_d1 = [make_node(d1, g1_by_cid[cid], count) for cid, count in d1_counts.items() if cid in g1_by_cid]
        nodes_d2 = [make_node(d2, g2_by_cid[cid], count) for cid, count in d2_counts.items() if cid in g2_by_cid]
        edges = [
            {"source": f"{d1}_{c1}", "target": f"{d2}_{c2}", "weight": weight}
            for (c1, c2), weight in edge_counts.most_common()
        ]
        return {
            "domain1": d1,
            "domain2": d2,
            "nodes_d1": nodes_d1,
            "nodes_d2": nodes_d2,
            "edges": edges,
            "cross_triple_count": sum(edge_counts.values()),
        }

    return await _kg_view_cached_json(
        request,
        current_user,
        "cross-domain-communities",
        {
            "domain1": d1,
            "domain2": d2,
            "years": normalized_years,
            "companies": normalized_companies,
        },
        build_payload,
    )


@app.get("/api/cross-cluster-detail")
async def text_kg_cross_cluster_detail(
    request: Request,
    cluster1: str,
    cluster2: str,
    years: List[str] = Query(default=[]),
    companies: List[str] = Query(default=[]),
    limit: int = Query(default=500, ge=10, le=30000),
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    d1 = _kg_view_domain(cluster1)
    d2 = _kg_view_domain(cluster2)
    normalized_years = _kg_view_param_list(years)
    normalized_companies = _kg_view_param_list(companies)

    def build_payload() -> Dict[str, Any]:
        graph = _kg_view_source_graph(request, current_user)
        data = _kg_view_filtered_text_graph(
            graph,
            domains=[d1, d2],
            companies=normalized_companies or None,
            years=normalized_years or None,
            limit=limit,
        )
        node_domain = {node["id"]: node["esg_domain"] for node in data["nodes"]}
        data["edges"] = [
            edge for edge in data["edges"]
            if {node_domain.get(edge["source"]), node_domain.get(edge["target"])} == {d1, d2}
        ]
        return data

    return await _kg_view_cached_json(
        request,
        current_user,
        "cross-cluster-detail",
        {
            "cluster1": d1,
            "cluster2": d2,
            "years": normalized_years,
            "companies": normalized_companies,
            "limit": limit,
        },
        build_payload,
    )


@app.get("/api/stats")
async def text_kg_stats(
    request: Request,
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    def build_payload() -> Dict[str, Any]:
        from collections import Counter

        graph = _kg_view_source_graph(request, current_user)
        data = _kg_view_filtered_text_graph(graph, limit=30000)
        return {
            "by_domain": dict(Counter(node["esg_domain"] for node in data["nodes"])),
            "by_type": dict(Counter(node["type"] for node in data["nodes"]).most_common(20)),
            "by_relationship": dict(Counter(edge["type"] for edge in data["edges"]).most_common(20)),
        }

    return await _kg_view_cached_json(
        request,
        current_user,
        "stats",
        {},
        build_payload,
    )


@app.get("/api/greenwashing")
async def text_kg_greenwashing(
    request: Request,
    company: Optional[str] = None,
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    normalized_company = _kg_view_text(company)

    def build_payload() -> Dict[str, Any]:
        graph = _kg_view_source_graph(request, current_user)
        data = _kg_view_filtered_text_graph(graph, companies=[normalized_company] if normalized_company else None, limit=30000)
        scores = [float(edge.get("credibility_score") or 0) for edge in data["edges"]]
        quantitative = sum(1 for edge in data["edges"] if re.search(r"\d", edge.get("evidence") or ""))
        total = len(data["edges"])
        avg = round(sum(scores) / len(scores), 2) if scores else 0
        weak = sum(1 for score in scores if score < 2)
        index = round((weak / max(total, 1)) * 100, 1)
        risk = "HIGH" if index >= 45 else "MEDIUM" if index >= 20 else "LOW"
        return {
            "company": company or "All",
            "total_triples": total,
            "greenwashing_index": index,
            "risk_level": risk,
            "credibility_avg": avg,
            "credibility_high": sum(1 for score in scores if score >= 3),
            "credibility_low": weak,
            "quantitative_ratio": round((quantitative / max(total, 1)) * 100, 1),
        }

    return await _kg_view_cached_json(
        request,
        current_user,
        "greenwashing",
        {"company": normalized_company},
        build_payload,
    )


@app.get("/public/knowledge-graph")
async def public_knowledge_graph(
    limit: int = Query(default=25000, ge=1, le=_PUBLIC_GRAPH_MAX_NODES),
    edge_limit: int = Query(default=30000, ge=0, le=_PUBLIC_GRAPH_MAX_EDGES),
    current_user: Optional[dict] = Depends(get_optional_current_user),
):
    node_limit = _bounded_graph_limit(limit, 25000, _PUBLIC_GRAPH_MAX_NODES)
    bounded_edge_limit = _bounded_graph_limit(edge_limit, 30000, _PUBLIC_GRAPH_MAX_EDGES, minimum=0)
    try:
        if not _public_graph_cache_enabled() or not chat_memory_service.is_available():
            graph = _load_real_knowledge_graph(
                current_user=current_user,
                node_limit=node_limit,
                edge_limit=bounded_edge_limit,
            )
            return JSONResponse(content=graph, headers={"X-Graph-Cache": "bypass"})

        cache_key = _public_graph_cache_key(
            current_user=current_user,
            node_limit=node_limit,
            edge_limit=bounded_edge_limit,
        )
        cached = _public_graph_cache_get(cache_key)
        if cached is not None:
            graph = cached["graph"]
            if _public_graph_cache_is_fresh(cached):
                return JSONResponse(content=graph, headers={"X-Graph-Cache": "hit"})

            lock_token = _public_graph_cache_acquire_lock()
            if lock_token:
                _GRAPH_CACHE_REFRESH_EXECUTOR.submit(
                    _refresh_public_graph_cache_in_background,
                    cache_key,
                    dict(current_user or {}),
                    node_limit,
                    bounded_edge_limit,
                    lock_token,
                )
            return JSONResponse(content=graph, headers={"X-Graph-Cache": "stale"})

        lock_token = _public_graph_cache_acquire_lock()
        if lock_token:
            loop = asyncio.get_running_loop()
            graph = await loop.run_in_executor(
                _GRAPH_CACHE_REFRESH_EXECUTOR,
                _load_and_store_public_graph_cache,
                cache_key,
                dict(current_user or {}),
                node_limit,
                bounded_edge_limit,
                lock_token,
            )
            return JSONResponse(content=graph, headers={"X-Graph-Cache": "miss"})

        warmed = await _wait_for_public_graph_cache(cache_key)
        if warmed is not None:
            return JSONResponse(content=warmed["graph"], headers={"X-Graph-Cache": "wait-hit"})

        return JSONResponse(
            status_code=503,
            content={
                "error": "knowledge_graph_cache_warming",
                "message": "Knowledge graph cache is warming. Please retry shortly.",
            },
            headers={"Retry-After": "3", "X-Graph-Cache": "warming"},
        )
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "knowledge_graph_failed", "message": str(exc)},
        )


@app.get("/kg-api/filters")
async def kg_filters(document_id: Optional[str] = None, current_user: dict = Depends(require_admin)):
    try:
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(
                status_code=400,
                content={"error": "neo4j_unavailable", "message": "Neo4j is not configured or unavailable."},
            )
        return JSONResponse(content=store.get_visualization_filters(document_id=document_id))
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "kg_filters_failed", "message": str(exc)},
        )


@app.get("/kg-api/graph")
async def kg_graph(
    document_id: Optional[str] = None,
    years: List[str] = Query(default=[]),
    companies: List[str] = Query(default=[]),
    esg_domains: List[str] = Query(default=[]),
    limit: int = 1000,
    current_user: dict = Depends(require_admin),
):
    try:
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(
                status_code=400,
                content={"error": "neo4j_unavailable", "message": "Neo4j is not configured or unavailable."},
            )
        return JSONResponse(
            content=store.get_visualization_graph(
                years=years or None,
                companies=companies or None,
                esg_domains=esg_domains or None,
                limit=limit,
                document_id=document_id,
            )
        )
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "kg_graph_failed", "message": str(exc)},
        )


@app.get("/kg-api/stats")
async def kg_stats(document_id: Optional[str] = None, current_user: dict = Depends(require_admin)):
    try:
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(
                status_code=400,
                content={"error": "neo4j_unavailable", "message": "Neo4j is not configured or unavailable."},
            )
        return JSONResponse(content=store.get_visualization_stats(document_id=document_id))
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "kg_stats_failed", "message": str(exc)},
        )


@app.get("/graph/neo4j/status")
async def neo4j_status(current_user: dict = Depends(require_admin)):
    if not neo4j_sdk_available():
        return JSONResponse(content={"enabled": False, "connected": False, "reason": "neo4j_sdk_missing"})
    if not neo4j_configured():
        return JSONResponse(content={"enabled": False, "connected": False, "reason": "neo4j_not_configured"})
    try:
        assert_neo4j_ready()
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(content={"enabled": False, "connected": False, "reason": "neo4j_unavailable"})
        status = store.ping()
        status["auto_sync"] = NEO4J_AUTO_SYNC
        status["stats"] = store.get_stats()
        return JSONResponse(content=status)
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"enabled": True, "connected": False, "reason": "neo4j_error", "message": str(exc)},
        )


@app.get("/graph/neo4j/entity/{entity_name}")
async def neo4j_entity(entity_name: str, limit: int = 20, current_user: dict = Depends(require_admin)):
    try:
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(
                status_code=400,
                content={"error": "neo4j_unavailable", "message": "Neo4j is not configured or the SDK is missing."},
            )
        return JSONResponse(content=store.get_entity(entity_name, limit=limit))
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "neo4j_entity_failed", "message": str(exc)},
        )


@app.get("/graph/neo4j/subgraph")
async def neo4j_subgraph(entity: str, hops: int = 2, limit: int = 50, current_user: dict = Depends(require_admin)):
    try:
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(
                status_code=400,
                content={"error": "neo4j_unavailable", "message": "Neo4j is not configured or the SDK is missing."},
            )
        return JSONResponse(content=store.get_subgraph(entity=entity, hops=hops, limit=limit))
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "neo4j_subgraph_failed", "message": str(exc)},
        )


@app.post("/graph/neo4j/question")
async def neo4j_question(request: Neo4jQuestionRequest, current_user: dict = Depends(require_admin)):
    try:
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(
                status_code=400,
                content={"error": "neo4j_unavailable", "message": "Neo4j is not configured or the SDK is missing."},
            )
        return JSONResponse(
            content=store.find_relevant_subgraph(
                question=request.question,
                hops=request.hops,
                limit=request.limit,
            )
        )
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "neo4j_question_failed", "message": str(exc)},
        )


@app.get("/graph/causal/backward")
async def causal_backward(entity: str, depth: int = 3, limit: int = 20):
    try:
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(status_code=400, content={"error": "neo4j_unavailable", "message": "Neo4j is not configured or unavailable."})
        return JSONResponse(content=CausalReasoner(store).backward_chain(entity, depth=depth, limit=limit))
    except Exception as exc:
        return JSONResponse(status_code=500, content={"error": "causal_backward_failed", "message": str(exc)})


@app.get("/graph/causal/forward")
async def causal_forward(entity: str, depth: int = 3, limit: int = 20):
    try:
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(status_code=400, content={"error": "neo4j_unavailable", "message": "Neo4j is not configured or unavailable."})
        return JSONResponse(content=CausalReasoner(store).forward_chain(entity, depth=depth, limit=limit))
    except Exception as exc:
        return JSONResponse(status_code=500, content={"error": "causal_forward_failed", "message": str(exc)})


@app.get("/graph/causal/path")
async def causal_path(source: str, target: str, max_depth: int = 5):
    try:
        store = get_neo4j_store()
        if store is None:
            return JSONResponse(status_code=400, content={"error": "neo4j_unavailable", "message": "Neo4j is not configured or unavailable."})
        return JSONResponse(content=CausalReasoner(store).shortest_path(source, target, max_depth=max_depth))
    except Exception as exc:
        return JSONResponse(status_code=500, content={"error": "causal_path_failed", "message": str(exc)})


@app.post("/extract", response_model=EsgExtractionResponse)
async def extract(request: EsgExtractionRequest):
    try:
        result = extract_esg(request.text)
        return EsgExtractionResponse(**result)
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={
                "entities": [],
                "relations": [],
                "error": "extraction_failed",
                "message": str(exc),
            },
        )


@app.post("/rag/ask")
async def rag_ask(
    request: RagAskRequest,
    current_user: Optional[dict] = Depends(get_optional_current_user),
    db: aiosqlite.Connection = Depends(_get_db),
):
    try:
        if not current_user:
            await _enforce_rag_rate_limit(db, current_user, request.reasoning_mode)
        pre_history, pre_memory_backend = _load_request_history_for_rag(request, current_user)
        answer_intent = classify_answer_intent(
            query=request.question,
            history_block=format_history(pre_history, current_query=request.question),
        )
        if str(answer_intent.get("mode") or "") in {"general", "chitchat"}:
            context = _resolve_general_rag_request_context(request, current_user, pre_history, pre_memory_backend)
        else:
            context = _resolve_rag_request_context(request, current_user)
        if context["error_response"] is not None:
            return context["error_response"]
        quota = await _enforce_rag_rate_limit(db, current_user, request.reasoning_mode) if current_user else None
        memory_context, injected_memories = await _load_long_term_memory_context(db, current_user, request.question)
        answer_history = _history_with_long_term_memory(context["history"], memory_context)
        result = answer_question(
            request.question,
            top_k=request.top_k,
            history=answer_history,
            retrieval_filters=context["filters"],
            mode=request.mode or "ask",
            reasoning_mode=request.reasoning_mode or "flash",
            user_id=context["user_id"],
            answer_intent=answer_intent,
        )
        result["memory_backend"] = context["memory_backend"]
        result["long_term_memory"] = {
            "backend": "sqlite+vector",
            "injected": len(injected_memories),
            "auto_extract": bool(current_user),
        }
        if quota and not quota.get("bypassed"):
            result["quota"] = quota
        if request.session_id:
            result["session_id"] = request.session_id
        if current_user:
            _remember_exchange_later(
                user_id=str(current_user.get("id") or ""),
                user_message=request.question,
                assistant_message=str(result.get("answer") or ""),
                source="rag",
            )
        return JSONResponse(content=result)
    except HTTPException:
        raise
    except FileNotFoundError as exc:
        return JSONResponse(
            status_code=400,
            content={"answer": "", "sources": [], "error": "vector_store_missing", "message": str(exc)},
        )
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"answer": "", "sources": [], "error": "rag_failed", "message": str(exc)},
        )


def _encode_sse_event(event: Dict[str, Any]) -> str:
    return f"data: {json.dumps(event, ensure_ascii=False)}\n\n"


def _build_streaming_response(
    *,
    request: RagAskRequest,
    stream_factory,
    fallback_factory,
) -> StreamingResponse:
    event_queue: "queue.Queue[Optional[Dict[str, Any]]]" = queue.Queue()

    def _producer() -> None:
        try:
            try:
                iterator = stream_factory()
                for event in iterator:
                    event_queue.put(event)
            except NotImplementedError:
                event_queue.put({"type": "done", "payload": fallback_factory()})
        except Exception as exc:
            event_queue.put({"type": "error", "message": str(exc)})
        finally:
            event_queue.put(None)

    def _event_stream():
        thread = threading.Thread(target=_producer, daemon=True)
        thread.start()
        while True:
            try:
                item = event_queue.get(timeout=15.0)
            except queue.Empty:
                yield ": heartbeat\n\n"
                continue
            if item is None:
                break
            yield _encode_sse_event(item)

    return StreamingResponse(
        _event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


@app.post("/rag/ask/stream")
async def rag_ask_stream(
    request: RagAskRequest,
    current_user: Optional[dict] = Depends(get_optional_current_user),
    db: aiosqlite.Connection = Depends(_get_db),
):
    try:
        if not current_user:
            await _enforce_rag_rate_limit(db, current_user, request.reasoning_mode)
        pre_history, pre_memory_backend = _load_request_history_for_rag(request, current_user)
        answer_intent = classify_answer_intent(
            query=request.question,
            history_block=format_history(pre_history, current_query=request.question),
        )
        if str(answer_intent.get("mode") or "") in {"general", "chitchat"}:
            context = _resolve_general_rag_request_context(request, current_user, pre_history, pre_memory_backend)
        else:
            context = _resolve_rag_request_context(request, current_user)
        if context["error_response"] is not None:
            return context["error_response"]
        quota = await _enforce_rag_rate_limit(db, current_user, request.reasoning_mode) if current_user else None
        memory_context, injected_memories = await _load_long_term_memory_context(db, current_user, request.question)
        answer_history = _history_with_long_term_memory(context["history"], memory_context)

        def _stream_factory():
            for event in stream_answer_question(
                request.question,
                top_k=request.top_k,
                history=answer_history,
                retrieval_filters=context["filters"],
                mode=request.mode or "ask",
                reasoning_mode=request.reasoning_mode or "flash",
                user_id=context["user_id"],
                answer_intent=answer_intent,
            ):
                if event.get("type") == "done":
                    payload = dict(event.get("payload") or {})
                    payload["memory_backend"] = context["memory_backend"]
                    payload["long_term_memory"] = {
                        "backend": "sqlite+vector",
                        "injected": len(injected_memories),
                        "auto_extract": bool(current_user),
                    }
                    if quota and not quota.get("bypassed"):
                        payload["quota"] = quota
                    if request.session_id:
                        payload["session_id"] = request.session_id
                    if current_user:
                        _remember_exchange_later(
                            user_id=str(current_user.get("id") or ""),
                            user_message=request.question,
                            assistant_message=str(payload.get("answer") or ""),
                            source="rag_stream",
                        )
                    yield {"type": "done", "payload": payload}
                else:
                    if event.get("type") == "meta":
                        payload = dict(event.get("payload") or {})
                        if request.session_id:
                            payload["session_id"] = request.session_id
                        yield {"type": "meta", "payload": payload}
                    else:
                        yield event

        def _fallback_factory():
            result = answer_question(
                request.question,
                top_k=request.top_k,
                history=answer_history,
                retrieval_filters=context["filters"],
                mode=request.mode or "ask",
                reasoning_mode=request.reasoning_mode or "flash",
                user_id=context["user_id"],
                answer_intent=answer_intent,
            )
            result["memory_backend"] = context["memory_backend"]
            result["long_term_memory"] = {
                "backend": "sqlite+vector",
                "injected": len(injected_memories),
                "auto_extract": bool(current_user),
            }
            if quota and not quota.get("bypassed"):
                result["quota"] = quota
            if request.session_id:
                result["session_id"] = request.session_id
            if current_user:
                _remember_exchange_later(
                    user_id=str(current_user.get("id") or ""),
                    user_message=request.question,
                    assistant_message=str(result.get("answer") or ""),
                    source="rag_stream_fallback",
                )
            return result

        return _build_streaming_response(
            request=request,
            stream_factory=_stream_factory,
            fallback_factory=_fallback_factory,
        )
    except HTTPException:
        raise
    except FileNotFoundError as exc:
        return JSONResponse(
            status_code=400,
            content={"answer": "", "sources": [], "error": "vector_store_missing", "message": str(exc)},
        )
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"answer": "", "sources": [], "error": "rag_failed", "message": str(exc)},
        )


@app.post("/desktop/screenshot/summarize")
async def desktop_screenshot_summarize(
    request: DesktopScreenshotSummaryRequest,
    current_user: dict = Depends(get_current_user),
    db: aiosqlite.Connection = Depends(_get_db),
):
    image_data_url = _validate_desktop_image_data_url(request.image_data_url)
    prompt = str(request.prompt or "").strip() or "Summarize this screen."
    client = get_openai_client()
    if client is None:
        return JSONResponse(
            status_code=503,
            content={
                "error": "openai_unavailable",
                "message": "Screenshot summary requires OPENAI_API_KEY to be configured.",
            },
        )

    try:
        memory_context, injected_memories = await _load_long_term_memory_context(db, current_user, prompt)
        system_prompt = (
            "You summarize user-initiated desktop screenshots for a work and academic research assistant. "
            "Prioritize useful document, study, ESG, finance, strategy, data, and error information. "
            "Ignore desktop chrome, wallpaper, window controls, app navigation, decorative UI, casual chat, "
            "and other noisy content unless it directly affects the user's task. "
            "Return concise Markdown with useful evidence, learning points, and next steps. "
            "Do not invent hidden context, and clearly separate visible evidence from inference."
        )
        if memory_context:
            system_prompt = f"{system_prompt}\n\n{memory_context}"
        quota = await _enforce_rag_rate_limit(db, current_user, "deep")
        response = client.chat.completions.create(
            model=_DESKTOP_SCREENSHOT_SUMMARY_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt,
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": image_data_url, "detail": "low"}},
                    ],
                },
            ],
            temperature=0.2,
            **chat_token_kwargs(_DESKTOP_SCREENSHOT_SUMMARY_MODEL, _DESKTOP_SCREENSHOT_SUMMARY_MAX_TOKENS),
        )
        summary = _extract_chat_completion_text(response)
        if not summary:
            return JSONResponse(
                status_code=502,
                content={"error": "empty_screenshot_summary", "message": "The model returned an empty summary."},
            )
        payload = {
            "summary": summary,
            "model": _DESKTOP_SCREENSHOT_SUMMARY_MODEL,
            "long_term_memory": {
                "backend": "sqlite+vector",
                "injected": len(injected_memories),
                "auto_extract": True,
            },
        }
        if quota and not quota.get("bypassed"):
            payload["quota"] = quota
        _remember_exchange_later(
            user_id=str(current_user.get("id") or ""),
            user_message=prompt,
            assistant_message=summary,
            source="desktop_screenshot",
        )
        return JSONResponse(content=payload)
    except HTTPException:
        raise
    except Exception as exc:
        return JSONResponse(
            status_code=502,
            content={"error": "screenshot_summary_failed", "message": str(exc)},
        )


@app.post("/desktop/word/review")
async def desktop_word_review(
    instruction: str = Form(""),
    goal: str = Form("academic"),
    template: str = Form("general"),
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db: aiosqlite.Connection = Depends(_get_db),
):
    try:
        raw_name = Path(unquote(str(file.filename or ""))).name
        if not raw_name.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="Only .docx Word documents are supported for AI editing.")
        file_name = _safe_word_filename(file.filename)
        file_bytes = await file.read()
        if not file_bytes:
            raise HTTPException(status_code=400, detail="Uploaded Word document is empty.")
        if len(file_bytes) > _DESKTOP_WORD_EDIT_MAX_BYTES:
            raise HTTPException(status_code=413, detail="Word document is too large for desktop editing.")

        _, paragraphs = _parse_docx_paragraphs(file_bytes)
        normalized_goal = _normalize_word_edit_goal(goal)
        normalized_template = _normalize_word_edit_template(template)
        quota = await _enforce_rag_rate_limit(db, current_user, "deep")
        evidence_sources = _retrieve_word_review_evidence(
            current_user=current_user,
            paragraphs=paragraphs,
            instruction=instruction,
            goal=normalized_goal,
            template=normalized_template,
        )
        suggestions, model_name = _generate_word_edit_suggestions(
            instruction,
            paragraphs,
            goal=normalized_goal,
            template=normalized_template,
            evidence_sources=evidence_sources,
        )

        session_id = uuid.uuid4().hex
        session_dir = _word_edit_session_dir(current_user, session_id)
        session_dir.mkdir(parents=True, exist_ok=True)
        original_path = session_dir / "original.docx"
        original_path.write_bytes(file_bytes)

        session_payload = {
            "session_id": session_id,
            "user_id": str(current_user.get("id") or ""),
            "file_name": file_name,
            "instruction": str(instruction or "").strip(),
            "goal": normalized_goal,
            "goal_label": _WORD_EDIT_GOAL_LABELS[normalized_goal],
            "template": normalized_template,
            "template_label": _WORD_EDIT_TEMPLATE_LABELS[normalized_template],
            "model": model_name,
            "created_at": _utc_now_iso(),
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "suggestions": suggestions,
            "evidence_sources": evidence_sources,
        }
        (session_dir / "session.json").write_text(json.dumps(session_payload, ensure_ascii=False, indent=2), encoding="utf-8")

        response_payload: Dict[str, Any] = {
            "session_id": session_id,
            "file_name": file_name,
            "goal": normalized_goal,
            "goal_label": _WORD_EDIT_GOAL_LABELS[normalized_goal],
            "template": normalized_template,
            "template_label": _WORD_EDIT_TEMPLATE_LABELS[normalized_template],
            "paragraph_count": len(paragraphs),
            "reviewed_paragraph_count": min(len(_select_word_paragraphs_for_review(paragraphs)), len(paragraphs)),
            "suggestions": suggestions,
            "evidence_sources": evidence_sources,
            "category_labels": _WORD_EDIT_CATEGORY_LABELS,
            "evidence_gap_labels": _WORD_EDIT_EVIDENCE_GAP_LABELS,
            "model": model_name,
        }
        if quota and not quota.get("bypassed"):
            response_payload["quota"] = quota
        return JSONResponse(content=response_payload)
    except HTTPException:
        raise
    except Exception as exc:
        return JSONResponse(
            status_code=400,
            content={"error": "word_review_failed", "message": str(exc)},
        )


@app.post("/desktop/word/export")
async def desktop_word_export(
    request: DesktopWordEditExportRequest,
    current_user: dict = Depends(get_current_user),
):
    try:
        session = _load_word_edit_session(current_user, request.session_id)
        session_dir = _word_edit_session_dir(current_user, request.session_id)
        original_path = session_dir / "original.docx"
        if not original_path.is_file():
            raise HTTPException(status_code=404, detail="Original Word document was not found.")

        accepted_ids = {str(item) for item in (request.accepted_suggestion_ids or [])}
        suggestions = session.get("suggestions") if isinstance(session.get("suggestions"), list) else []
        replacements = {
            str(item.get("paragraph_id") or ""): str(item.get("replacement") or "")
            for item in suggestions
            if isinstance(item, dict) and str(item.get("id") or "") in accepted_ids
        }

        document = DocxDocument(str(original_path))
        applied_count = 0
        visible_index = 0
        for paragraph in document.paragraphs:
            if not str(paragraph.text or "").strip():
                continue
            visible_index += 1
            paragraph_id = f"p_{visible_index:03d}"
            replacement = replacements.get(paragraph_id)
            if replacement:
                _replace_paragraph_text(paragraph, replacement)
                applied_count += 1

        output_name = _unique_output_name(str(session.get("file_name") or "document.docx"))
        output_path = session_dir / output_name
        document.save(str(output_path))
        encoded = base64.b64encode(output_path.read_bytes()).decode("ascii")
        return JSONResponse(
            content={
                "file_name": output_name,
                "mime_type": _WORD_DOCX_MIME,
                "data_base64": encoded,
                "applied_count": applied_count,
            }
        )
    except HTTPException:
        raise
    except Exception as exc:
        return JSONResponse(
            status_code=400,
            content={"error": "word_export_failed", "message": str(exc)},
        )


@app.post("/pipeline/pdf")
async def pipeline_pdf(request: PipelinePdfRequest, current_user: dict = Depends(get_current_user)):
    try:
        result = run_pdf_pipeline(request.pdf_path, request.name)
        return JSONResponse(content=result)
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": str(exc)},
        )


@app.get("/documents")
async def list_documents(current_user: dict = Depends(get_current_user)):
    try:
        audit_index = list_latest_uploads_by_document_id()
        entries = [entry for entry in _collect_document_entries() if _can_access_entry(current_user, entry)]
        documents: List[Dict[str, Any]] = []
        for entry in entries:
            document_id = str(entry.get("document_id") or "").strip()
            if not document_id:
                continue
            audit = audit_index.get(document_id)
            if audit and str(audit.get("status") or "") in {"deleted", "deleted_with_warnings"}:
                continue
            try:
                documents.append(summarize_registered_document(entry, audit=audit))
            except Exception as exc:
                print(f"[documents] Summary load failed for {document_id}: {type(exc).__name__}: {exc}")
        return JSONResponse(content={"documents": documents})
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "document_list_failed", "message": str(exc)},
        )


@app.post("/documents/upload")
async def upload_document(
    title: str = Form(""),
    domain: str = Form("general"),
    source_type: str = Form(""),
    source: str = Form(""),
    content: str = Form(""),
    file: Optional[UploadFile] = File(default=None),
    current_user: dict = Depends(get_current_user),
):
    try:
        file_bytes = await file.read() if file is not None else None
        is_admin = _is_admin_user(current_user)
        document_group = "global_kb" if is_admin else "user_private"
        visibility_scope = "global" if is_admin else "private"
        result = ingest_uploaded_document(
            title=title or (file.filename if file else "Uploaded document"),
            domain=domain,
            source=source,
            source_type=source_type,
            content=content,
            filename=file.filename if file else None,
            file_bytes=file_bytes,
            document_group=document_group,
            owner_user_id=str(current_user.get("id") or ""),
            visibility_scope=visibility_scope,
        )
        return JSONResponse(content=result)
    except Exception as exc:
        return JSONResponse(
            status_code=400,
            content={"error": "document_upload_failed", "message": str(exc)},
        )


@app.post("/documents/upload-async")
async def upload_document_async(
    title: str = Form(""),
    domain: str = Form("general"),
    source_type: str = Form(""),
    source: str = Form(""),
    content: str = Form(""),
    file: Optional[UploadFile] = File(default=None),
    current_user: dict = Depends(get_current_user),
):
    try:
        file_bytes = await file.read() if file is not None else None
        is_admin = _is_admin_user(current_user)
        document_group = "global_kb" if is_admin else "user_private"
        visibility_scope = "global" if is_admin else "private"
        job = start_ingestion_job(
            title=title or (file.filename if file else "Uploaded document"),
            domain=domain,
            source=source,
            source_type=source_type,
            content=content,
            filename=file.filename if file else None,
            file_bytes=file_bytes,
            document_group=document_group,
            uploader=current_user,
            owner_user_id=str(current_user.get("id") or ""),
            visibility_scope=visibility_scope,
        )
        return JSONResponse(content=job)
    except Exception as exc:
        return JSONResponse(
            status_code=400,
            content={"error": "document_upload_failed", "message": str(exc)},
        )


@app.get("/documents/jobs/{job_id}")
async def get_document_job(job_id: str, current_user: dict = Depends(get_current_user)):
    upload = get_upload(job_id)
    if upload is not None and not _is_admin_user(current_user):
        uploader = upload.get("uploader") or {}
        if str(uploader.get("id") or "").strip() != str(current_user.get("id") or "").strip():
            return JSONResponse(
                status_code=403,
                content={"error": "job_forbidden", "message": "You do not have access to this ingestion job."},
            )
    job = get_ingestion_job(job_id)
    if job is None:
        return JSONResponse(
            status_code=404,
            content={"error": "job_not_found", "message": f"No ingestion job found for id {job_id}."},
        )
    return JSONResponse(content=job)


@app.get("/documents/{document_id}")
async def get_document(document_id: str, current_user: dict = Depends(get_current_user)):
    try:
        entry = document_registry.get_entry(document_id, valid_only=True)
        audit = get_latest_upload_by_document_id(document_id)
        if entry is None and audit is not None:
            entry = _entry_from_upload(audit)
        if entry is None:
            return JSONResponse(
                status_code=404,
                content={"error": "document_not_found", "message": f"No document found for id {document_id}."},
            )
        if not _can_access_entry(current_user, entry):
            return JSONResponse(
                status_code=403,
                content={"error": "document_forbidden", "message": "You do not have access to this document."},
            )
        if audit and str(audit.get("status") or "") in {"deleted", "deleted_with_warnings"}:
            return JSONResponse(
                status_code=404,
                content={"error": "document_not_found", "message": f"No document found for id {document_id}."},
            )
        return JSONResponse(content={"document": load_registered_document(entry, audit=audit)})
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "document_load_failed", "message": str(exc)},
        )


@app.delete("/documents/{document_id}")
async def delete_document(document_id: str, current_user: dict = Depends(get_current_user)):
    try:
        entry = document_registry.get_entry(document_id, valid_only=False)
        audit = get_latest_upload_by_document_id(document_id)
        if entry is None and audit is not None:
            entry = _entry_from_upload(audit)
        if entry is None:
            return JSONResponse(
                status_code=404,
                content={"error": "document_not_found", "message": f"No document found for id {document_id}."},
            )
        if not _can_access_entry(current_user, entry):
            return JSONResponse(
                status_code=403,
                content={"error": "document_forbidden", "message": "You do not have access to this document."},
            )
        if not _is_admin_user(current_user):
            entry_owner = str(entry.get("owner_user_id") or "").strip()
            if entry_owner != str(current_user.get("id") or "").strip():
                return JSONResponse(
                    status_code=403,
                    content={"error": "document_forbidden", "message": "Only admins can delete global knowledge base documents."},
                )
        if audit and str(audit.get("status") or "") in {"deleted", "deleted_with_warnings"}:
            return JSONResponse(
                status_code=404,
                content={"error": "document_not_found", "message": f"No document found for id {document_id}."},
            )

        upload = audit or {
            "document_id": document_id,
            "status": "completed",
            "paths": {
                "processed_text_path": str((entry.get("paths") or {}).get("processed_text") or ""),
                "chunks_path": str((entry.get("paths") or {}).get("chunks") or ""),
                "extractions_path": str((entry.get("paths") or {}).get("extractions") or ""),
                "graph_path": str((entry.get("paths") or {}).get("graph") or ""),
                "vector_store_path": str((entry.get("paths") or {}).get("vector_store") or ""),
            },
        }

        if audit:
            mark_upload_deleted(
                audit["job_id"],
                deleted_by=str(current_user.get("email") or current_user.get("username") or ""),
                reason="user_deleted",
                status="deleted",
                cleanup_status="cleanup_pending",
                cleanup_detail="User deletion in progress.",
            )

        cleanup = delete_uploaded_document(upload)
        document_registry.remove(document_id)

        warnings = list(cleanup.get("warnings") or [])
        neo4j_result = cleanup.get("neo4j") or {}
        if neo4j_result.get("enabled") and neo4j_result.get("deleted") is False and neo4j_result.get("reason") not in {
            "missing_document_id",
        }:
            warnings.append(f"neo4j: {neo4j_result.get('reason') or 'delete_not_confirmed'}")

        if audit:
            if warnings:
                record_upload_cleanup(
                    audit["job_id"],
                    cleanup_status="cleanup_failed",
                    cleanup_detail=f"Cleanup warnings: {'; '.join(warnings[:5])}",
                    status="deleted_with_warnings",
                )
            else:
                deleted_count = len(cleanup.get("deleted_paths") or [])
                record_upload_cleanup(
                    audit["job_id"],
                    cleanup_status="cleanup_completed",
                    cleanup_detail=f"Deleted {deleted_count} local path(s). Neo4j: {neo4j_result.get('reason') or neo4j_result.get('deleted')}.",
                    status="deleted",
                )

        return JSONResponse(
            content={
                "deleted": True,
                "document_id": document_id,
                "cleanup": cleanup,
            }
        )
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": "document_delete_failed", "message": str(exc)},
        )


@app.post("/documents/ingest-text")
async def ingest_text_document(request: ManualDocumentRequest, current_user: dict = Depends(get_current_user)):
    try:
        is_admin = _is_admin_user(current_user)
        result = ingest_uploaded_document(
            title=request.title,
            domain=request.domain,
            source=request.source,
            source_type=request.source_type,
            content=request.content,
            document_group="global_kb" if is_admin else "user_private",
            owner_user_id=str(current_user.get("id") or ""),
            visibility_scope="global" if is_admin else "private",
        )
        return JSONResponse(content=result)
    except Exception as exc:
        return JSONResponse(
            status_code=400,
            content={"error": "document_ingest_failed", "message": str(exc)},
        )


@app.post("/documents/rebuild-graph")
async def rebuild_graph(request: RebuildDocumentGraphRequest, current_user: dict = Depends(get_current_user)):
    try:
        result = rebuild_document_graph(request.model_dump())
        return JSONResponse(content=result)
    except Exception as exc:
        return JSONResponse(
            status_code=400,
            content={"error": "document_rebuild_failed", "message": str(exc)},
        )
