"""
CausalGraph Platform - Main API Server
Provides RESTful API endpoints for causal relationship extraction and knowledge graph construction
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import re
import uvicorn
import os
from dotenv import load_dotenv
import PyPDF2
import docx
import io
import time
import uuid
import aiosqlite
from datetime import datetime

from causal_extractor import CausalExtractor
from knowledge_graph_builder import KnowledgeGraphBuilder
from graph_visualizer import GraphVisualizer
from config import Config
from auth import auth_service, get_current_user_dependency
from storage import storage_service
from models import (
    UserResponse,
    DocumentCreate,
    KnowledgeGraphCreate,
    AnalysisResult,
    CDKActivationCreate,
    CDKStatus,
    DocumentIngestRequest,
    RagRequest,
    GraphRagRequest,
)
from database import get_db, db_manager
from cdk_service import cdk_service
from services.graph_service import graph_service
from services.ingest_service import ingest_service
from services.parser_service import parser_service
from services.rag_service import rag_service
from services.router_service import router_service

# Load environment variables
load_dotenv()

app = FastAPI(
    title="CausalGraph Platform API",
    description="Enterprise platform for causal relationship extraction and knowledge graph construction",
    version="1.0.0"
)

# Configure CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=Config.ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize core components
try:
    causal_extractor = CausalExtractor()
except Exception as e:
    print(f"Warning: CausalExtractor initialization failed: {e}")
    print("OpenAI integration will not be available. Please set OPENAI_API_KEY environment variable.")
    causal_extractor = None

kg_builder = KnowledgeGraphBuilder()
graph_viz = GraphVisualizer()

# Data Models
class TextInput(BaseModel):
    """Input model for text processing requests"""
    text: str
    domain: str = "general"  # healthcare, financial, general

class QueryInput(BaseModel):
    """Input model for graph query requests"""
    question: str
    graph_id: str

class DocumentInput(BaseModel):
    """Input model for document processing requests"""
    title: str
    content: str
    domain: str
    source: str = ""

@app.on_event("startup")
async def startup_event():
    """Initialize database on startup"""
    await db_manager.init_database()

@app.get("/")
async def root():
    """Root endpoint providing API information"""
    return {
        "message": "CausalGraph Platform API",
        "version": "1.0.0",
        "description": "Enterprise platform for causal relationship extraction and knowledge graph construction"
    }

# Authentication endpoints
@app.post("/auth/register")
async def register_user(user_data: dict, db: aiosqlite.Connection = Depends(get_db)):
    """Register new user account"""
    try:
        print(f"Received user data: {user_data}")
        from models import UserCreate
        user_create = UserCreate(**user_data)
        user = await auth_service.create_user(user_create, db)
        return {"success": True, "user": user.dict()}
    except Exception as e:
        print(f"Registration error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/auth/login")
async def login_user(user_data: dict, db: aiosqlite.Connection = Depends(get_db)):
    """Login user and return access token"""
    try:
        from models import UserLogin
        user_login = UserLogin(**user_data)
        token = await auth_service.login_user(user_login, db)
        return {"success": True, "token": token.dict()}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/auth/me")
async def get_current_user(current_user: UserResponse = Depends(get_current_user_dependency)):
    """Get current authenticated user information"""
    return {"success": True, "user": current_user.dict()}

# CDK (Activation Code) endpoints
@app.post("/cdk/activate")
async def activate_cdk(
    cdk_data: CDKActivationCreate,
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    """Activate a CDK code for the current user"""
    success, message = await cdk_service.activate_cdk(
        current_user.id, 
        cdk_data.cdk_code, 
        db
    )
    
    if success:
        return {"success": True, "message": message}
    else:
        raise HTTPException(status_code=400, detail=message)

@app.get("/cdk/status")
async def get_cdk_status(
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    """Get the CDK activation status for the current user"""
    status = await cdk_service.get_user_cdk_status(current_user.id, db)
    return {"success": True, "status": status}

@app.get("/cdk/models")
async def get_available_models(
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    """Get available models based on CDK status"""
    status = await cdk_service.get_user_cdk_status(current_user.id, db)
    models = cdk_service.get_available_models(status)
    return {"success": True, "models": models}

# User data management endpoints
@app.get("/documents")
async def get_user_documents(current_user: UserResponse = Depends(get_current_user_dependency), db: aiosqlite.Connection = Depends(get_db)):
    """Get all documents for current user"""
    documents = await storage_service.get_user_documents(current_user.id, db)
    return {"success": True, "documents": [doc.dict() for doc in documents]}

@app.post("/documents/upload")
async def upload_document_file(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    """Upload an ESG document and store its parsed text for later ingestion."""
    file_bytes = await file.read()
    if len(file_bytes) > Config.MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File exceeds maximum allowed size")

    try:
        parsed = parser_service.parse_bytes(file.filename or "document.txt", file_bytes)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Document parsing failed: {e}")

    stored_content = parsed.text
    if parsed.file_type == "pdf" and parsed.pages:
        stored_content = "\n\n".join(
            f"<<<PAGE:{page.page_number}>>>\n{page.text}" for page in parsed.pages
        ).strip()

    document_data = DocumentCreate(
        title=title or parsed.title,
        file_type=parsed.file_type,
        file_size=len(file_bytes),
        content=stored_content,
        original_filename=file.filename or f"{parsed.title}.{parsed.file_type}",
        user_id=current_user.id,
    )
    saved_document = await storage_service.save_document(document_data, db)

    return {
        "success": True,
        "document": saved_document.dict(),
        "processing_stats": parsed.stats,
    }

@app.post("/documents/ingest")
async def ingest_document(
    request: DocumentIngestRequest,
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    """Chunk a stored document, create embeddings, and extract ESG entities/relations."""
    document = await storage_service.get_document_by_id(request.document_id, current_user.id, db)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    summary = await ingest_service.ingest_document(
        db=db,
        document=document,
        user_id=current_user.id,
        company=request.company,
        source=request.source,
        category=request.category,
        chunk_size=request.chunk_size,
        chunk_overlap=request.chunk_overlap,
        run_entity_extraction=request.run_entity_extraction,
        run_relation_extraction=request.run_relation_extraction,
    )
    await storage_service.mark_document_processed(document.id, db)

    return {"success": True, "ingest_summary": summary}

@app.get("/documents/{document_id}")
async def get_document(document_id: str, current_user: UserResponse = Depends(get_current_user_dependency), db: aiosqlite.Connection = Depends(get_db)):
    """Get specific document by ID"""
    document = await storage_service.get_document_by_id(document_id, current_user.id, db)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"success": True, "document": document.dict()}

@app.delete("/documents/{document_id}")
async def delete_document(document_id: str, current_user: UserResponse = Depends(get_current_user_dependency), db: aiosqlite.Connection = Depends(get_db)):
    """Delete document and related data"""
    success = await storage_service.delete_document(document_id, current_user.id, db)
    if not success:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"success": True, "message": "Document deleted successfully"}

@app.get("/knowledge-graphs")
async def get_user_knowledge_graphs(current_user: UserResponse = Depends(get_current_user_dependency), db: aiosqlite.Connection = Depends(get_db)):
    """Get all knowledge graphs for current user"""
    graphs = await storage_service.get_user_knowledge_graphs(current_user.id, db)
    return {"success": True, "knowledge_graphs": [graph.dict() for graph in graphs]}

@app.get("/knowledge-graphs/{graph_id}")
async def get_knowledge_graph(graph_id: str, current_user: UserResponse = Depends(get_current_user_dependency), db: aiosqlite.Connection = Depends(get_db)):
    """Get specific knowledge graph by ID"""
    graph = await storage_service.get_knowledge_graph_by_id(graph_id, current_user.id, db)
    if not graph:
        raise HTTPException(status_code=404, detail="Knowledge graph not found")
    return {"success": True, "knowledge_graph": graph.dict()}

@app.get("/analysis-results")
async def get_user_analysis_results(current_user: UserResponse = Depends(get_current_user_dependency), db: aiosqlite.Connection = Depends(get_db)):
    """Get all analysis results for current user"""
    results = await storage_service.get_user_analysis_results(current_user.id, db)
    return {"success": True, "analysis_results": [result.dict() for result in results]}

@app.post("/chat/rag")
async def chat_rag(
    request: RagRequest,
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    """Answer an ESG question with vector retrieval over document chunks."""
    result = await rag_service.answer_rag(
        db=db,
        user_id=current_user.id,
        question=request.question,
        document_id=request.document_id,
        top_k=request.top_k,
    )
    result["recommended_route"] = router_service.route(request.question)
    return {"success": True, **result}

@app.post("/chat/graph-rag")
async def chat_graph_rag(
    request: GraphRagRequest,
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    """Answer a multi-hop ESG question using both retrieved chunks and a graph subgraph."""
    result = await rag_service.answer_graph_rag(
        db=db,
        user_id=current_user.id,
        question=request.question,
        document_id=request.document_id,
        top_k=request.top_k,
        depth=request.depth,
    )
    return {"success": True, **result}

@app.post("/chat/ask")
async def chat_ask(
    request: GraphRagRequest,
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    """Automatically route a question to RAG or Graph RAG."""
    route = router_service.route(request.question)
    if route == "graph-rag":
        result = await rag_service.answer_graph_rag(
            db=db,
            user_id=current_user.id,
            question=request.question,
            document_id=request.document_id,
            top_k=request.top_k,
            depth=request.depth,
        )
    else:
        result = await rag_service.answer_rag(
            db=db,
            user_id=current_user.id,
            question=request.question,
            document_id=request.document_id,
            top_k=request.top_k,
        )
    return {"success": True, **result}

@app.get("/graph/entities")
async def list_graph_entities(
    document_id: Optional[str] = None,
    company: Optional[str] = None,
    limit: int = 200,
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    entities = await graph_service.list_entities(
        db=db,
        user_id=current_user.id,
        document_id=document_id,
        company=company,
        limit=limit,
    )
    return {"success": True, "entities": [entity.dict() for entity in entities], "count": len(entities)}

@app.get("/graph/relations")
async def list_graph_relations(
    document_id: Optional[str] = None,
    entity_id: Optional[str] = None,
    limit: int = 200,
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    relations = await graph_service.list_relations(
        db=db,
        user_id=current_user.id,
        document_id=document_id,
        entity_id=entity_id,
        limit=limit,
    )
    return {"success": True, "relations": [relation.dict() for relation in relations], "count": len(relations)}

@app.get("/graph/entity/{entity_id}")
async def get_graph_entity(
    entity_id: str,
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    entity = await graph_service.get_entity(db, current_user.id, entity_id)
    if not entity:
        raise HTTPException(status_code=404, detail="Entity not found")

    relations = await graph_service.list_relations(db, current_user.id, entity_id=entity_id, limit=100)
    return {
        "success": True,
        "entity": entity.dict(),
        "relations": [relation.dict() for relation in relations],
    }

@app.get("/graph/subgraph")
async def get_graph_subgraph(
    entity_id: Optional[str] = None,
    question: Optional[str] = None,
    document_id: Optional[str] = None,
    depth: int = 2,
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    subgraph = await graph_service.get_subgraph(
        db=db,
        user_id=current_user.id,
        entity_id=entity_id,
        question=question,
        document_id=document_id,
        depth=depth,
    )
    return {
        "success": True,
        "matched_entity_ids": subgraph["matched_entity_ids"],
        "entities": [entity.dict() for entity in subgraph["entities"]],
        "relations": [relation.dict() for relation in subgraph["relations"]],
    }

@app.post("/extract-causal-relationships")
async def extract_causal_relationships(input_data: TextInput):
    """Extract causal relationships from text input"""
    if not causal_extractor:
        raise HTTPException(
            status_code=503, 
            detail="OpenAI integration not available. Please set OPENAI_API_KEY environment variable."
        )
    
    try:
        relationships = await causal_extractor.extract_relationships(
            input_data.text, 
            input_data.domain
        )
        return {
            "success": True,
            "relationships": relationships,
            "count": len(relationships)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/build-knowledge-graph")
async def build_knowledge_graph(input_data: TextInput):
    """Build a knowledge graph from text input"""
    if not causal_extractor:
        raise HTTPException(
            status_code=503, 
            detail="OpenAI integration not available. Please set OPENAI_API_KEY environment variable."
        )
    
    try:
        # Extract causal relationships
        relationships = await causal_extractor.extract_relationships(
            input_data.text, 
            input_data.domain
        )
        
        # Build enhanced knowledge graph for better quality
        graph_data = kg_builder.build_enhanced_graph(relationships, input_data.domain)
        
        return {
            "success": True,
            "graph": graph_data,
            "relationships_count": len(relationships),
            "nodes_count": len(graph_data["nodes"]),
            "edges_count": len(graph_data["edges"])
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-document")
async def upload_document(
    title: str = Form(...),
    content: str = Form(...),
    domain: str = Form("general"),
    source: str = Form(""),
    current_user: UserResponse = Depends(get_current_user_dependency),
    db: aiosqlite.Connection = Depends(get_db)
):
    """Upload and process document for causal relationship extraction"""
    if not causal_extractor:
        raise HTTPException(
            status_code=503, 
            detail="OpenAI integration not available. Please set OPENAI_API_KEY environment variable."
        )
    
    start_time = time.time()
    
    try:
        # Check CDK status for the user
        cdk_status = await cdk_service.get_user_cdk_status(current_user.id, db)
        print(f"CDK Status for user {current_user.id}: {cdk_status}")
        
        # Set CDK status in causal extractor if available
        if causal_extractor:
            causal_extractor.set_cdk_status(cdk_status)
            print(f"Set CDK status in causal extractor: {cdk_status}")
        else:
            print("Warning: Causal extractor not available")
        
        # Save document to database
        document_data = DocumentCreate(
            title=title,
            file_type="text",
            file_size=len(content.encode('utf-8')),
            content=content,
            original_filename=f"{title}.txt",
            user_id=current_user.id
        )
        
        saved_document = await storage_service.save_document(document_data, db)
        
        # Clean and segment text for better extraction
        cleaned_text = clean_text_for_extraction(content)
        text_segments = smart_text_segmentation(cleaned_text)
        
        all_relationships = []
        
        # Process each segment
        for i, segment in enumerate(text_segments):
            print(f"Processing segment {i+1}/{len(text_segments)}")
            
            try:
                # Try GPT-4 extraction first (if CDK activated)
                if causal_extractor:
                    segment_relationships = await causal_extractor.extract_relationships(segment, domain)
                else:
                    # Fallback to pattern matching
                    segment_relationships = extract_relationships_patterns(segment, domain)
                
                if segment_relationships:
                    all_relationships.extend(segment_relationships)
                
            except Exception as e:
                print(f"Error processing segment {i+1}: {e}")
                # Continue with next segment
                continue
        
        if not all_relationships:
            # Mark document as processed even if no relationships found
            await storage_service.mark_document_processed(saved_document.id, db)
            
            return JSONResponse(
                status_code=200,
                content={
                    "success": True,
                    "message": "No causal relationships found in the document",
                    "relationships": all_relationships,
                    "count": 0,
                    "document_id": saved_document.id,
                    "quality_metrics": {
                        "total_segments": len(text_segments),
                        "processed_segments": len([r for r in all_relationships if r]),
                        "extraction_method": "GPT-4 + Pattern Fallback"
                    }
                }
            )
        
        # Build enhanced knowledge graph
        try:
            graph_data = kg_builder.build_enhanced_graph(all_relationships, domain)
            
            # Get graph statistics from the dictionary
            graph_stats = {
                "node_count": len(graph_data.get("nodes", [])),
                "edge_count": len(graph_data.get("edges", [])),
                "is_directed": graph_data.get("metadata", {}).get("is_directed", True),
                "is_acyclic": graph_data.get("metadata", {}).get("is_acyclic", True),
                "enhanced": graph_data.get("metadata", {}).get("enhanced", True)
            }
            
        except Exception as e:
            print(f"Error building enhanced graph: {e}")
            # Fallback to basic graph
            try:
                graph_data = kg_builder.build_graph(all_relationships, domain)
                graph_stats = kg_builder.get_graph_statistics(graph_data)
            except Exception as fallback_error:
                print(f"Fallback graph building also failed: {fallback_error}")
                # Create empty graph data as last resort
                graph_data = {"nodes": [], "edges": [], "metadata": {}}
                graph_stats = {"node_count": 0, "edge_count": 0, "is_directed": True, "is_acyclic": True, "enhanced": False}
        
        # Determine extraction method based on CDK status
        extraction_method = "Local Pattern Matching"
        if cdk_status.get("is_activated", False):
            extraction_method = "GPT-4 + Pattern Fallback"
        
        # Save knowledge graph to database
        graph_data_to_save = KnowledgeGraphCreate(
            title=f"Knowledge Graph: {title}",
            description=f"Automatically generated knowledge graph from {title}",
            user_id=current_user.id,
            document_id=saved_document.id,
            nodes=graph_data.get("nodes", []),
            edges=graph_data.get("edges", []),
            metadata={
                "domain": domain,
                "source": source,
                "extraction_method": extraction_method,
                "total_segments": len(text_segments),
                "statistics": graph_stats
            }
        )
        
        print(f"About to save knowledge graph")
        print(f"Graph data type: {type(graph_data_to_save)}")
        print(f"Graph data metadata keys: {list(graph_data_to_save.metadata.keys())}")
        try:
            saved_graph = await storage_service.save_knowledge_graph(graph_data_to_save, db)
            print(f"Saved knowledge graph successfully: {saved_graph.id}")
        except Exception as e:
            print(f"Error saving knowledge graph: {e}")
            import traceback
            traceback.print_exc()
            raise
        
        # Save analysis result
        processing_time = time.time() - start_time
        
        # Convert CausalRelationship objects to dictionaries for storage
        relationships_dict = []
        for rel in all_relationships:
            if rel:  # Check if relationship exists
                rel_dict = {
                    "cause": rel.cause,
                    "effect": rel.effect,
                    "relationship_type": rel.relationship_type,
                    "confidence": rel.confidence,
                    "evidence": rel.evidence,
                    "cause_char_span": rel.cause_char_span,
                    "effect_char_span": rel.effect_char_span
                }
                relationships_dict.append(rel_dict)
        
        # Determine extraction method based on CDK status
        extraction_method = "Local Pattern Matching"
        if cdk_status.get("is_activated", False):
            extraction_method = "GPT-4 + Pattern Fallback"
        
        analysis_result = AnalysisResult(
            id=str(uuid.uuid4()),  # Add missing id field
            document_id=saved_document.id,
            user_id=current_user.id,
            relationships=relationships_dict,  # Use converted dictionaries
            quality_metrics={
                "total_segments": len(text_segments),
                "processed_segments": len([r for r in all_relationships if r]),
                "extraction_method": extraction_method,
                "graph_nodes": len(graph_data.get("nodes", [])),
                "graph_edges": len(graph_data.get("edges", []))
            },
            processing_time=processing_time,
            created_at=datetime.utcnow()  # Add missing created_at field
        )
        
        await storage_service.save_analysis_result(analysis_result, db)
        
        # Mark document as processed
        await storage_service.mark_document_processed(saved_document.id, db)
        
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": f"Successfully extracted {len(all_relationships)} causal relationships",
                "relationships": relationships_dict,  # Use converted dictionaries instead of CausalRelationship objects
                "count": len(all_relationships),
                "document_id": saved_document.id,
                "graph_id": saved_graph.id,
                "graph_data": graph_data,
                "graph_statistics": graph_stats,
                "quality_metrics": {
                    "total_segments": len(text_segments),
                    "processed_segments": len([r for r in all_relationships if r]),
                    "extraction_method": extraction_method,
                    "graph_nodes": len(graph_data.get("nodes", [])),
                    "graph_edges": len(graph_data.get("edges", [])),
                    "processing_time": processing_time
                }
            }
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

def smart_text_segmentation(text: str, max_chars: int = 8000) -> list:
    """Smart text segmentation to ensure each segment doesn't exceed GPT-4 token limits"""
    if len(text) <= max_chars:
        return [text]
    
    segments = []
    current_pos = 0
    
    while current_pos < len(text):
        # Calculate end position for current segment
        end_pos = min(current_pos + max_chars, len(text))
        
        # If not the last segment, try to split at sentence boundaries
        if end_pos < len(text):
            # Look backward for nearest sentence ending
            for i in range(end_pos, max(current_pos + max_chars - 500, current_pos), -1):
                if text[i] in '.!?':
                    end_pos = i + 1
                    break
        
        # Extract current segment
        segment = text[current_pos:end_pos].strip()
        if segment:
            segments.append(segment)
        
        current_pos = end_pos
    
    return segments

def clean_text_for_extraction(text: str) -> str:
    """Clean and preprocess text for better causal relationship extraction"""
    import re
    
    # Remove excessive whitespace and normalize
    text = re.sub(r'\s+', ' ', text)
    
    # Remove common document artifacts (less aggressive)
    text = re.sub(r'Page \d+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'ISSN.*?\d+', '', text, flags=re.IGNORECASE)
    
    # Remove citation patterns like [1], [2], etc.
    text = re.sub(r'\[\d+\]', '', text)
    
    # Remove very short lines (likely headers or formatting) - less aggressive
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if len(line) > 5:  # Keep more lines
            cleaned_lines.append(line)
    
    text = ' '.join(cleaned_lines)
    
    # Final cleanup
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def extract_relationships_patterns(text: str, domain: str) -> List[Dict[str, Any]]:
    """Extract causal relationships using pattern matching when OpenAI is unavailable"""
    relationships = []
    
    # Clean the text first
    cleaned_text = clean_text_for_extraction(text)
    print(f"Original text length: {len(text)}, Cleaned text length: {len(cleaned_text)}")
    
    # More specific and meaningful causal patterns
    causal_patterns = [
        # Direct causation
        r'(\b\w+(?:\s+\w+){1,4})\s+(?:causes?|leads?\s+to|results?\s+in|triggers?|induces?|provokes?)\s+(\b\w+(?:\s+\w+){1,4})',
        
        # Influence and impact
        r'(\b\w+(?:\s+\w+){1,4})\s+(?:affects?|influences?|impacts?|modifies?|changes?|alters?)\s+(\b\w+(?:\s+\w+){1,4})',
        
        # Positive effects
        r'(\b\w+(?:\s+\w+){1,4})\s+(?:improves?|enhances?|boosts?|strengthens?|promotes?|increases?)\s+(\b\w+(?:\s+\w+){1,4})',
        
        # Negative effects
        r'(\b\w+(?:\s+\w+){1,4})\s+(?:reduces?|decreases?|lowers?|diminishes?|weakens?|harms?)\s+(\b\w+(?:\s+\w+){1,4})',
        
        # Prevention
        r'(\b\w+(?:\s+\w+){1,4})\s+(?:prevents?|blocks?|inhibits?|suppresses?|stops?|avoids?)\s+(\b\w+(?:\s+\w+){1,4})',
        
        # Risk and probability
        r'(\b\w+(?:\s+\w+){1,4})\s+(?:increases?\s+the\s+risk\s+of|reduces?\s+the\s+chance\s+of)\s+(\b\w+(?:\s+\w+){1,4})',
    ]
    
    import re
    
    for pattern in causal_patterns:
        matches = re.finditer(pattern, cleaned_text, re.IGNORECASE)
        for match in matches:
            cause = match.group(1).strip()
            effect = match.group(2).strip()
            
            # Quality checks for extracted concepts
            if (len(cause) < 3 or len(effect) < 3 or 
                len(cause) > 100 or len(effect) > 100 or
                cause.lower() in ['the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'] or
                effect.lower() in ['the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']):
                continue
            
            # Determine relationship type based on the pattern
            if 'causes' in pattern or 'leads to' in pattern or 'results in' in pattern:
                rel_type = 'causes'
            elif 'improves' in pattern or 'enhances' in pattern or 'increases' in pattern:
                rel_type = 'improves'
            elif 'reduces' in pattern or 'decreases' in pattern or 'lowers' in pattern:
                rel_type = 'reduces'
            elif 'prevents' in pattern or 'blocks' in pattern or 'inhibits' in pattern:
                rel_type = 'prevents'
            elif 'harms' in pattern or 'damages' in pattern:
                rel_type = 'harms'
            elif 'risk' in pattern:
                rel_type = 'increases_risk'
            else:
                rel_type = 'affects'
            
            # Create relationship object
            relationship = {
                "cause": cause,
                "effect": effect,
                "confidence": 0.8,  # Higher confidence for cleaned patterns
                "evidence": cleaned_text[max(0, match.start()-100):match.end()+100],  # More context
                "domain": domain,
                "relationship_type": rel_type
            }
            
            relationships.append(relationship)
    
    print(f"Pattern extraction found {len(relationships)} relationships")
    return relationships

@app.post("/process-file")
async def process_file(
    file: UploadFile = File(...), 
    file_type: str = Form(...),
    current_user: UserResponse = Depends(get_current_user_dependency)
):
    """Process uploaded files and extract text content"""
    try:
        # Read file content
        file_content = await file.read()
        
        if file_type == "pdf":
            # Process PDF files
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                if not text_content.strip():
                    raise HTTPException(status_code=400, detail="PDF appears to be empty or contains no extractable text")
                
                return {
                    "success": True,
                    "content": text_content.strip(),
                    "file_type": "pdf",
                    "pages": len(pdf_reader.pages)
                }
                
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"PDF processing failed: {str(e)}")
        
        elif file_type in ["doc", "docx"]:
            # Process Word documents
            try:
                if file_type == "docx":
                    doc = docx.Document(io.BytesIO(file_content))
                    text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    # For .doc files, we'll need additional libraries
                    raise HTTPException(status_code=400, detail="DOC files require additional processing libraries")
                
                if not text_content.strip():
                    raise HTTPException(status_code=400, detail="Word document appears to be empty")
                
                return {
                    "success": True,
                    "content": text_content.strip(),
                    "file_type": file_type,
                    "paragraphs": len(doc.paragraphs)
                }
                
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Word document processing failed: {str(e)}")
        
        elif file_type == "txt":
            # Process text files
            try:
                text_content = file_content.decode('utf-8')
                
                if not text_content.strip():
                    raise HTTPException(status_code=400, detail="Text file appears to be empty")
                
                return {
                    "success": True,
                    "content": text_content.strip(),
                    "file_type": "txt",
                    "characters": len(text_content)
                }
                
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Text file processing failed: {str(e)}")
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File processing failed: {str(e)}")

@app.post("/query-graph")
async def query_graph(input_data: QueryInput):
    """Query the knowledge graph with natural language questions"""
    try:
        # For MVP, we'll use a simple keyword-based approach
        # In production, this would use more sophisticated NLP
        answer = kg_builder.query_graph(input_data.question, input_data.graph_id)
        
        return {
            "success": True,
            "question": input_data.question,
            "answer": answer
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/get-graph/{graph_id}")
async def get_graph(graph_id: str):
    """Retrieve a specific knowledge graph"""
    try:
        if graph_id in kg_builder.documents:
            return {
                "success": True,
                "document": kg_builder.documents[graph_id]
            }
        else:
            raise HTTPException(status_code=404, detail="Graph not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/list-graphs")
async def list_graphs():
    """List all available knowledge graphs"""
    try:
        graphs = []
        for doc_id, doc in kg_builder.documents.items():
            graphs.append({
                "id": doc_id,
                "title": doc["title"],
                "domain": doc["domain"],
                "source": doc["source"],
                "relationships_count": len(doc["relationships"]),
                "nodes_count": len(doc["graph"]["nodes"]),
                "edges_count": len(doc["graph"]["edges"])
            })
        
        return {
            "success": True,
            "graphs": graphs,
            "total": len(graphs)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    try:
        # Test database connection
        async with aiosqlite.connect("causalgraph.db") as conn:
            cursor = await conn.execute("SELECT name FROM sqlite_master WHERE type='table'")
            tables = await cursor.fetchall()
            table_names = [table[0] for table in tables]
        
        return {
            "status": "healthy", 
            "service": "CausalGraph.AI API",
            "database": "connected",
            "tables": table_names
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "service": "CausalGraph.AI API", 
            "database": "error",
            "error": str(e)
        }

@app.post("/build-enhanced-graph")
async def build_enhanced_graph(input_data: DocumentInput):
    """Build an enhanced knowledge graph with advanced normalization and merging"""
    try:
        relationships = []
        
        # Try OpenAI first if available
        if causal_extractor:
            try:
                relationships = await causal_extractor.extract_relationships(
                    input_data.content, 
                    input_data.domain
                )
                print(f"OpenAI extracted {len(relationships)} relationships")
            except Exception as e:
                print(f"OpenAI extraction failed: {e}")
                # Fall back to pattern-based extraction
                relationships = extract_relationships_patterns(input_data.content, input_data.domain)
        else:
            # Use pattern-based extraction when OpenAI is not available
            relationships = extract_relationships_patterns(input_data.content, input_data.domain)
        
        # Build enhanced knowledge graph
        graph_data = kg_builder.build_enhanced_graph(relationships, input_data.domain)
        
        # Store document and graph
        document_id = f"enhanced_doc_{len(kg_builder.documents) + 1}"
        kg_builder.documents[document_id] = {
            "title": input_data.title,
            "domain": input_data.domain,
            "source": input_data.source,
            "graph": graph_data,
            "relationships": relationships,
            "enhanced": True
        }
        
        return {
            "success": True,
            "document_id": document_id,
            "graph": graph_data,
            "relationships": relationships,
            "relationships_count": len(relationships),
            "enhanced": True,
            "message": "Enhanced knowledge graph built with advanced normalization and merging"
        }
        
    except Exception as e:
        print(f"Enhanced graph building error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host=Config.HOST, port=Config.PORT)
