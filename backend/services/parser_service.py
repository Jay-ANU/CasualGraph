"""Document parsing helpers for PDF, DOCX, TXT, and Markdown."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional
import io

import PyPDF2
import docx


@dataclass
class ParsedPage:
    page_number: int
    text: str


@dataclass
class ParsedDocument:
    title: str
    file_type: str
    text: str
    pages: List[ParsedPage] = field(default_factory=list)
    stats: dict = field(default_factory=dict)


class ParserService:
    """Parse uploaded bytes into plain text with lightweight metadata."""

    @staticmethod
    def infer_file_type(filename: str, declared_type: Optional[str] = None) -> str:
        if declared_type:
            return declared_type.lower().strip(".")
        suffix = Path(filename).suffix.lower().strip(".")
        if suffix == "markdown":
            return "md"
        return suffix or "txt"

    def parse_bytes(self, filename: str, content: bytes, declared_type: Optional[str] = None) -> ParsedDocument:
        file_type = self.infer_file_type(filename, declared_type)
        title = Path(filename).stem or "Untitled Document"

        if file_type == "pdf":
            return self._parse_pdf(title, content)
        if file_type == "docx":
            return self._parse_docx(title, content)
        if file_type in {"txt", "md", "markdown", "rtf"}:
            return self._parse_text(title, file_type, content)

        raise ValueError(f"Unsupported file type: {file_type}")

    def _parse_pdf(self, title: str, content: bytes) -> ParsedDocument:
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        pages: List[ParsedPage] = []
        page_texts: List[str] = []

        for idx, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(ParsedPage(page_number=idx, text=text))
                page_texts.append(text)

        full_text = "\n\n".join(page_texts).strip()
        if not full_text:
            raise ValueError("PDF contains no extractable text")

        return ParsedDocument(
            title=title,
            file_type="pdf",
            text=full_text,
            pages=pages,
            stats={
                "pages": len(reader.pages),
                "extractable_pages": len(pages),
                "characters": len(full_text),
            },
        )

    def _parse_docx(self, title: str, content: bytes) -> ParsedDocument:
        document = docx.Document(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs).strip()
        if not full_text:
            raise ValueError("DOCX contains no extractable text")

        return ParsedDocument(
            title=title,
            file_type="docx",
            text=full_text,
            pages=[ParsedPage(page_number=1, text=full_text)],
            stats={
                "paragraphs": len(paragraphs),
                "characters": len(full_text),
            },
        )

    def _parse_text(self, title: str, file_type: str, content: bytes) -> ParsedDocument:
        text = None
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if text is None:
            raise ValueError("Unable to decode text content")

        full_text = text.strip()
        if not full_text:
            raise ValueError("Text document is empty")

        return ParsedDocument(
            title=title,
            file_type=file_type,
            text=full_text,
            pages=[ParsedPage(page_number=1, text=full_text)],
            stats={"characters": len(full_text), "lines": len(full_text.splitlines())},
        )


parser_service = ParserService()
